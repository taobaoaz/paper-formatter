"""
封面和声明页配置器
Cover and Declaration Page Configurator
用于配置和管理论文的封面、独创性声明、版权授权书等前置页面
"""

from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QGridLayout,
    QLabel, QPushButton, QLineEdit, QTextEdit, QComboBox,
    QGroupBox, QFormLayout, QCheckBox, QFileDialog, QMessageBox
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont
import json
import os


class CoverDeclarationConfigDialog(QDialog):
    """封面和声明页配置对话框"""
    
    def __init__(self, config_data=None, parent=None):
        super().__init__(parent)
        self.config_data = config_data or self._get_default_config()
        self.template_path = None
        self.page_selection = None
        self.init_ui()
    
    def _get_default_config(self):
        """获取默认配置"""
        return {
            "cover": {
                "enabled": True,
                "type": "simple",
                "school_name": "XX 大学",
                "thesis_type": "本科毕业论文",
                "show_logo": True,
                "fields": {
                    "论文题目": "",
                    "作者姓名": "",
                    "学号": "",
                    "专业": "",
                    "指导教师": "",
                    "完成日期": ""
                }
            },
            "declaration": {
                "enabled": True,
                "type": "独创性声明",
                "show_signature": True,
                "custom_content": ""
            }
        }
    
    def init_ui(self):
        """初始化界面"""
        self.setWindowTitle('封面和声明页配置')
        self.setMinimumSize(700, 600)
        
        layout = QVBoxLayout(self)
        
        # 0. 模板上传和页面选择
        template_group = QGroupBox('📄 上传学校模板（可选）')
        template_layout = QVBoxLayout()
        
        template_info = QLabel(
            '如果学校提供了 Word 模板，可以上传后选择封面和声明页\n'
            '软件会自动提取并应用到你的论文中'
        )
        template_info.setWordWrap(True)
        template_layout.addWidget(template_info)
        
        template_btn_layout = QHBoxLayout()
        
        self.upload_template_btn = QPushButton('📂 上传模板并选择页面')
        self.upload_template_btn.clicked.connect(self.upload_and_select_pages)
        template_btn_layout.addWidget(self.upload_template_btn)
        
        self.template_info_label = QLabel('未上传模板')
        self.template_info_label.setStyleSheet('color: gray;')
        template_btn_layout.addWidget(self.template_info_label)
        
        template_layout.addLayout(template_btn_layout)
        template_group.setLayout(template_layout)
        layout.addWidget(template_group)
        
        # 1. 封面配置
        cover_group = QGroupBox('📘 封面配置')
        cover_layout = QFormLayout()
        
        # 启用封面
        self.cover_enabled = QCheckBox('启用封面')
        self.cover_enabled.setChecked(self.config_data['cover']['enabled'])
        cover_layout.addRow(self.cover_enabled)
        
        # 学校名称
        self.school_name = QLineEdit()
        self.school_name.setText(self.config_data['cover']['school_name'])
        cover_layout.addRow('学校名称:', self.school_name)
        
        # 论文类型
        self.thesis_type = QComboBox()
        self.thesis_type.addItems(['本科毕业论文', '硕士毕业论文', '博士毕业论文', '课程论文'])
        index = self.thesis_type.findText(self.config_data['cover']['thesis_type'])
        if index >= 0:
            self.thesis_type.setCurrentIndex(index)
        cover_layout.addRow('论文类型:', self.thesis_type)
        
        # 显示校徽
        self.show_logo = QCheckBox('显示校徽（如有）')
        self.show_logo.setChecked(self.config_data['cover']['show_logo'])
        cover_layout.addRow(self.show_logo)
        
        cover_group.setLayout(cover_layout)
        layout.addWidget(cover_group)
        
        # 2. 封面字段
        fields_group = QGroupBox('📝 封面字段信息')
        fields_layout = QGridLayout()
        
        self.field_entries = {}
        field_names = ['论文题目', '作者姓名', '学号', '专业', '指导教师', '完成日期']
        
        for i, field_name in enumerate(field_names):
            label = QLabel(f'{field_name}:')
            entry = QLineEdit()
            value = self.config_data['cover']['fields'].get(field_name, '')
            entry.setText(value)
            self.field_entries[field_name] = entry
            
            row = i // 2
            col = (i % 2) * 2
            fields_layout.addWidget(label, row, col)
            fields_layout.addWidget(entry, row, col + 1)
        
        fields_group.setLayout(fields_layout)
        layout.addWidget(fields_group)
        
        # 3. 声明页配置
        declaration_group = QGroupBox('📋 声明页配置')
        declaration_layout = QFormLayout()
        
        # 启用声明页
        self.declaration_enabled = QCheckBox('启用声明页')
        self.declaration_enabled.setChecked(self.config_data['declaration']['enabled'])
        declaration_layout.addRow(self.declaration_enabled)
        
        # 声明类型
        self.declaration_type = QComboBox()
        self.declaration_type.addItems([
            '独创性声明',
            '版权使用授权书',
            '独创性声明 + 授权书',
            '保密论文声明'
        ])
        index = self.declaration_type.findText(self.config_data['declaration']['type'])
        if index >= 0:
            self.declaration_type.setCurrentIndex(index)
        declaration_layout.addRow('声明类型:', self.declaration_type)
        
        # 显示签名区域
        self.show_signature = QCheckBox('显示签名区域')
        self.show_signature.setChecked(self.config_data['declaration']['show_signature'])
        declaration_layout.addRow(self.show_signature)
        
        declaration_group.setLayout(declaration_layout)
        layout.addWidget(declaration_group)
        
        # 4. 自定义声明内容
        custom_group = QGroupBox('📄 自定义声明内容（可选）')
        custom_layout = QVBoxLayout()
        
        self.custom_content = QTextEdit()
        self.custom_content.setPlaceholderText('如果不填写，将使用标准声明模板')
        self.custom_content.setText(self.config_data['declaration']['custom_content'])
        self.custom_content.setMaximumHeight(100)
        custom_layout.addWidget(self.custom_content)
        
        custom_group.setLayout(custom_layout)
        layout.addWidget(custom_group)
        
        # 5. 按钮
        button_layout = QHBoxLayout()
        
        self.save_btn = QPushButton('💾 保存配置')
        self.save_btn.clicked.connect(self.save_config)
        self.save_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 8px;')
        
        self.load_btn = QPushButton('📂 加载配置')
        self.load_btn.clicked.connect(self.load_config)
        
        self.cancel_btn = QPushButton('❌ 取消')
        self.cancel_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(self.save_btn)
        button_layout.addWidget(self.load_btn)
        button_layout.addWidget(self.cancel_btn)
        
        layout.addLayout(button_layout)
    
    def upload_and_select_pages(self):
        """上传模板并选择页面"""
        from template_page_selector import select_template_pages
        
        # 打开页面选择对话框
        selection = select_template_pages(self.template_path, self)
        
        if selection:
            self.template_path = selection['template_path']
            self.page_selection = selection
            
            # 更新界面显示
            filename = os.path.basename(self.template_path)
            self.template_info_label.setText(
                f'✓ 已上传：{filename}\n'
                f'封面：第 {selection["cover_page_index"] + 1} 页 | '
                f'声明：第 {selection["declaration_page_index"] + 1} 页'
            )
            self.template_info_label.setStyleSheet('color: green;')
            
            QMessageBox.information(
                self,
                '成功',
                f'模板页面选择成功！\n\n'
                f'封面：第 {selection["cover_page_index"] + 1} 页\n'
                f'声明页：第 {selection["declaration_page_index"] + 1} 页\n\n'
                f'软件会自动提取这两页并应用到论文中。'
            )
    
    def save_config(self):
        """保存配置"""
        try:
            # 收集配置
            config = {
                "template": {
                    "enabled": self.template_path is not None,
                    "template_path": self.template_path,
                    "page_selection": self.page_selection
                },
                "cover": {
                    "enabled": self.cover_enabled.isChecked(),
                    "type": "from_template" if self.template_path else "simple",
                    "school_name": self.school_name.text(),
                    "thesis_type": self.thesis_type.currentText(),
                    "show_logo": self.show_logo.isChecked(),
                    "fields": {
                        name: entry.text()
                        for name, entry in self.field_entries.items()
                    }
                },
                "declaration": {
                    "enabled": self.declaration_enabled.isChecked(),
                    "type": self.declaration_type.currentText(),
                    "show_signature": self.show_signature.isChecked(),
                    "custom_content": self.custom_content.toPlainText()
                }
            }
            
            # 保存文件
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                '保存封面配置',
                '',
                'JSON 文件 (*.json)'
            )
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(config, f, ensure_ascii=False, indent=2)
                
                QMessageBox.information(
                    self,
                    '成功',
                    f'配置已保存！\n\n{file_path}'
                )
                
                # 更新内部配置
                self.config_data = config
                self.accept()
        
        except Exception as e:
            QMessageBox.critical(self, '错误', f'保存失败：{str(e)}')
    
    def load_config(self):
        """加载配置"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            '加载封面配置',
            '',
            'JSON 文件 (*.json)'
        )
        
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                
                # 加载模板信息（如果有）
                if 'template' in config and config['template']['enabled']:
                    self.template_path = config['template']['template_path']
                    self.page_selection = config['template']['page_selection']
                    
                    if self.template_path and os.path.exists(self.template_path):
                        filename = os.path.basename(self.template_path)
                        self.template_info_label.setText(
                            f'✓ 已加载：{filename}\n'
                            f'封面：第 {self.page_selection["cover_page_index"] + 1} 页 | '
                            f'声明：第 {self.page_selection["declaration_page_index"] + 1} 页'
                        )
                        self.template_info_label.setStyleSheet('color: green;')
                
                # 填充界面
                self.cover_enabled.setChecked(config['cover']['enabled'])
                self.school_name.setText(config['cover']['school_name'])
                self.thesis_type.setCurrentText(config['cover']['thesis_type'])
                self.show_logo.setChecked(config['cover']['show_logo'])
                
                for name, value in config['cover']['fields'].items():
                    if name in self.field_entries:
                        self.field_entries[name].setText(value)
                
                self.declaration_enabled.setChecked(config['declaration']['enabled'])
                self.declaration_type.setCurrentText(config['declaration']['type'])
                self.show_signature.setChecked(config['declaration']['show_signature'])
                self.custom_content.setPlainText(config['declaration'].get('custom_content', ''))
                
                self.config_data = config
                
                QMessageBox.information(self, '成功', '配置加载成功！')
            
            except Exception as e:
                QMessageBox.critical(self, '错误', f'加载失败：{str(e)}')
    
    def get_config(self):
        """获取当前配置"""
        return self.config_data


class CoverDeclarationManager:
    """封面和声明页管理器"""
    
    def __init__(self, config=None):
        self.config = config or {}
    
    def generate_cover_page(self, doc, cover_config):
        """生成封面页"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # 添加学校名称
        school_para = doc.add_paragraph()
        school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        school_run = school_para.add_run(cover_config.get('school_name', 'XX 大学'))
        school_run.font.size = Pt(16)
        school_run.font.bold = True
        school_run.font.name = '黑体'
        
        # 添加论文类型
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_run = type_para.add_run(cover_config.get('thesis_type', '毕业论文'))
        type_run.font.size = Pt(14)
        type_run.font.bold = True
        
        # 添加封面字段
        fields = cover_config.get('fields', {})
        for field_name, value in fields.items():
            if value:  # 只添加有值的字段
                field_para = doc.add_paragraph()
                field_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                label_run = field_para.add_run(f'{field_name}：')
                label_run.font.size = Pt(12)
                label_run.font.bold = True
                
                value_run = field_para.add_run(value)
                value_run.font.size = Pt(12)
        
        # 添加分节符
        doc.add_section()
    
    def generate_declaration_page(self, doc, declaration_config):
        """生成声明页"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        declaration_type = declaration_config.get('type', '独创性声明')
        
        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(declaration_type)
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # 声明内容
        if declaration_config.get('custom_content'):
            # 使用自定义内容
            content_para = doc.add_paragraph(declaration_config['custom_content'])
            content_para.paragraph_format.line_spacing = 1.5
        else:
            # 使用标准模板
            if declaration_type == '独创性声明':
                content = """本人郑重声明：所呈交的学位论文，是本人在导师指导下进行研究工作所取得的成果。
除文中已经注明引用的内容外，本学位论文的研究成果不包含任何他人创作的、已公开发表或者没有公开发表的作品的内容。
对本论文所涉及的研究工作做出贡献的其他个人和集体，均已在文中以明确方式标明。
本学位论文独创性声明的法律责任由本人承担。"""
                
                content_para = doc.add_paragraph(content)
                content_para.paragraph_format.line_spacing = 1.5
        
        # 签名区域
        if declaration_config.get('show_signature', True):
            doc.add_paragraph('\n\n')
            
            sign_para = doc.add_paragraph()
            sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sign_run = sign_para.add_run('学位论文作者签名：____________')
            
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run('日期：______年____月____日')
        
        # 添加分节符
        doc.add_section()
    
    def apply_cover_and_declaration(self, doc, config):
        """应用封面和声明页配置"""
        if config.get('cover', {}).get('enabled', True):
            self.generate_cover_page(doc, config['cover'])
        
        if config.get('declaration', {}).get('enabled', True):
            self.generate_declaration_page(doc, config['declaration'])
        
        return doc
