import os
import re
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config
from parsers import ParsedDocument, DocumentElement
from template_manager import Template


@dataclass
class OptimizationResult:
    success: bool
    output_path: str
    changes: List[Dict[str, Any]] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    error_code: int = 0


class BaseFormatter(ABC):
    @abstractmethod
    def format(self, input_path: str, output_path: str, template: Template) -> OptimizationResult:
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        pass


class DocxFormatter(BaseFormatter):
    def get_supported_extensions(self) -> List[str]:
        return ['.docx']
    
    def format(self, input_path: str, output_path: str, template: Template) -> OptimizationResult:
        result = OptimizationResult(success=False, output_path=output_path)
        
        try:
            doc = Document(input_path)
            rules = template.rules
            
            self._apply_page_settings(doc, rules, result)
            self._apply_styles(doc, rules, result)
            self._format_paragraphs(doc, rules, result)
            self._format_headings(doc, rules, result)
            self._format_abstract(doc, rules, result)
            self._format_references(doc, rules, result)
            self._format_figures(doc, rules, result)
            self._format_tables(doc, rules, result)
            self._format_equations(doc, rules, result)
            self._format_page_numbers(doc, rules, result)
            self._format_headers_footers(doc, rules, result)
            
            doc.save(output_path)
            result.success = True
            
        except FileNotFoundError:
            result.errors.append("文件不存在")
            result.error_code = 1001
        except PermissionError:
            result.errors.append("文件被占用或无权限访问")
            result.error_code = 1003
        except Exception as e:
            result.errors.append(f"格式化失败: {str(e)}")
            result.error_code = 5001
        
        return result
    
    def _apply_page_settings(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        for section in doc.sections:
            if 'margin_top' in rules:
                section.top_margin = Cm(rules['margin_top'])
                result.changes.append({'type': '页面设置', 'detail': f'上边距设置为 {rules["margin_top"]}cm'})
            
            if 'margin_bottom' in rules:
                section.bottom_margin = Cm(rules['margin_bottom'])
                result.changes.append({'type': '页面设置', 'detail': f'下边距设置为 {rules["margin_bottom"]}cm'})
            
            if 'margin_left' in rules:
                section.left_margin = Cm(rules['margin_left'])
                result.changes.append({'type': '页面设置', 'detail': f'左边距设置为 {rules["margin_left"]}cm'})
            
            if 'margin_right' in rules:
                section.right_margin = Cm(rules['margin_right'])
                result.changes.append({'type': '页面设置', 'detail': f'右边距设置为 {rules["margin_right"]}cm'})
    
    def _apply_styles(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        try:
            normal_style = doc.styles['Normal']
            
            if 'font_family' in rules:
                self._set_style_font(normal_style, rules['font_family'])
            
            if 'font_size' in rules:
                normal_style.font.size = Pt(rules['font_size'])
            
            if 'line_spacing' in rules:
                normal_style.paragraph_format.line_spacing = rules['line_spacing']
            
            result.changes.append({'type': '样式设置', 'detail': '正文样式已更新'})
            
        except Exception as e:
            result.warnings.append(f'样式设置警告: {str(e)}')
    
    def _set_style_font(self, style, font_name: str):
        style.font.name = font_name
        r = style.element.rPr
        if r is None:
            r = OxmlElement('w:rPr')
            style.element.insert(0, r)
        rFonts = r.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            r.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    
    def _set_run_font(self, run, font_name: str):
        run.font.name = font_name
        r = run._element.rPr
        if r is None:
            r = OxmlElement('w:rPr')
            run._element.insert(0, r)
        rFonts = r.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            r.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)
    
    def _format_paragraphs(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        paragraph_count = 0
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            style_name = para.style.name if para.style else ''
            
            if 'Title' in style_name or 'Heading' in style_name or '标题' in style_name:
                continue
            
            if self._is_special_paragraph(para.text):
                continue
            
            if 'line_spacing' in rules:
                para.paragraph_format.line_spacing = rules['line_spacing']
            
            if 'font_size' in rules:
                for run in para.runs:
                    run.font.size = Pt(rules['font_size'])
            
            if 'font_family' in rules:
                for run in para.runs:
                    self._set_run_font(run, rules['font_family'])
            
            para.paragraph_format.first_line_indent = Cm(0.74)
            
            paragraph_count += 1
        
        if paragraph_count > 0:
            result.changes.append({'type': '段落格式', 'detail': f'已格式化 {paragraph_count} 个段落'})
    
    def _is_special_paragraph(self, text: str) -> bool:
        special_patterns = [
            r'^摘要',
            r'^Abstract',
            r'^关键词',
            r'^Keywords',
            r'^参考文献',
            r'^References',
            r'^\[\d+\]',
            r'^[1-9]\.\s',
        ]
        
        for pattern in special_patterns:
            if re.match(pattern, text.strip(), re.IGNORECASE):
                return True
        return False
    
    def _format_headings(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        heading_count = 0
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            style_name = para.style.name if para.style else ''
            level = 0
            
            if 'Heading 1' in style_name or '标题 1' in style_name:
                level = 1
            elif 'Heading 2' in style_name or '标题 2' in style_name:
                level = 2
            elif 'Heading 3' in style_name or '标题 3' in style_name:
                level = 3
            elif 'Title' in style_name:
                level = 0
            else:
                text = para.text.strip()
                if re.match(r'^第[一二三四五六七八九十]+[章节部分]', text):
                    level = 1
                elif re.match(r'^[1-9]\s', text):
                    level = 1
                elif re.match(r'^[1-9]\.[0-9]+\s', text):
                    level = 2
                elif re.match(r'^[1-9]\.[0-9]+\.[0-9]+\s', text):
                    level = 3
            
            if level > 0:
                font_key = f'heading{level}_font'
                size_key = f'heading{level}_size'
                
                if font_key in rules:
                    for run in para.runs:
                        self._set_run_font(run, rules[font_key])
                        run.font.bold = rules.get(f'heading{level}_bold', True)
                
                if size_key in rules:
                    for run in para.runs:
                        run.font.size = Pt(rules[size_key])
                
                para.paragraph_format.first_line_indent = Cm(0)
                heading_count += 1
        
        if heading_count > 0:
            result.changes.append({'type': '标题格式', 'detail': f'已格式化 {heading_count} 个标题'})
    
    def _format_abstract(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        abstract_found = False
        keywords_found = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if re.match(r'^摘要|^Abstract', text, re.IGNORECASE):
                abstract_found = True
                if 'abstract_font' in rules:
                    for run in para.runs:
                        self._set_run_font(run, rules['abstract_font'])
                if 'abstract_size' in rules:
                    for run in para.runs:
                        run.font.size = Pt(rules['abstract_size'])
            
            if re.match(r'^关键词|^Keywords|^关键字', text, re.IGNORECASE):
                keywords_found = True
                if 'keywords_font' in rules:
                    for run in para.runs:
                        self._set_run_font(run, rules['keywords_font'])
                if 'keywords_size' in rules:
                    for run in para.runs:
                        run.font.size = Pt(rules['keywords_size'])
        
        if abstract_found:
            result.changes.append({'type': '摘要格式', 'detail': '已格式化摘要'})
        if keywords_found:
            result.changes.append({'type': '关键词格式', 'detail': '已格式化关键词'})
    
    def _format_references(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        ref_count = 0
        in_references = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if re.match(r'^参考文献|^References', text, re.IGNORECASE):
                in_references = True
                if 'reference_font' in rules:
                    for run in para.runs:
                        self._set_run_font(run, rules['reference_font'])
                continue
            
            if in_references and text:
                if re.match(r'^\[\d+\]', text) or re.match(r'^\d+\.', text):
                    if 'reference_font' in rules:
                        for run in para.runs:
                            self._set_run_font(run, rules['reference_font'])
                    if 'reference_size' in rules:
                        for run in para.runs:
                            run.font.size = Pt(rules['reference_size'])
                    ref_count += 1
                
                if para.text and not re.match(r'^\[\d+\]', text) and not re.match(r'^\d+\.', text):
                    break
        
        if ref_count > 0:
            result.changes.append({'type': '参考文献格式', 'detail': f'已格式化 {ref_count} 条参考文献'})
    
    def _format_figures(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        figure_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if re.match(r'^图\s*\d', text) or re.match(r'^Figure\s*\d', text, re.IGNORECASE):
                if 'figure_caption_font' in rules:
                    for run in para.runs:
                        self._set_run_font(run, rules['figure_caption_font'])
                if 'figure_caption_size' in rules:
                    for run in para.runs:
                        run.font.size = Pt(rules['figure_caption_size'])
                
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                figure_count += 1
        
        if figure_count > 0:
            result.changes.append({'type': '图片标题格式', 'detail': f'已格式化 {figure_count} 个图片标题'})
    
    def _format_tables(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        table_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if re.match(r'^表\s*\d', text) or re.match(r'^Table\s*\d', text, re.IGNORECASE):
                if 'table_caption_font' in rules:
                    for run in para.runs:
                        self._set_run_font(run, rules['table_caption_font'])
                if 'table_caption_size' in rules:
                    for run in para.runs:
                        run.font.size = Pt(rules['table_caption_size'])
                
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                table_count += 1
        
        for table in doc.tables:
            if 'table_content_font' in rules or 'table_content_size' in rules:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if 'table_content_font' in rules:
                                    self._set_run_font(run, rules['table_content_font'])
                                if 'table_content_size' in rules:
                                    run.font.size = Pt(rules['table_content_size'])
        
        if table_count > 0:
            result.changes.append({'type': '表格格式', 'detail': f'已格式化 {table_count} 个表格'})
    
    def _format_equations(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        equation_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if re.match(r'^\([一二三四五六七八九十\d]+[-–]\d+\)$', text) or \
               re.match(r'^公式\s*\d', text):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                equation_count += 1
        
        if equation_count > 0:
            result.changes.append({'type': '公式格式', 'detail': f'已格式化 {equation_count} 个公式'})
    
    def _format_headers_footers(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        for section in doc.sections:
            header = section.header
            footer = section.footer
            
            if 'header_font' in rules or 'header_size' in rules:
                for para in header.paragraphs:
                    for run in para.runs:
                        if 'header_font' in rules:
                            self._set_run_font(run, rules['header_font'])
                        if 'header_size' in rules:
                            run.font.size = Pt(rules['header_size'])
            
            if 'page_number_font' in rules or 'page_number_size' in rules:
                for para in footer.paragraphs:
                    for run in para.runs:
                        if 'page_number_font' in rules:
                            self._set_run_font(run, rules['page_number_font'])
                        if 'page_number_size' in rules:
                            run.font.size = Pt(rules['page_number_size'])
        
        result.changes.append({'type': '页眉页脚', 'detail': '已格式化页眉页脚'})
    
    def _format_toc(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        toc_found = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            
            if re.match(r'^目录$|^Contents$', text, re.IGNORECASE):
                toc_found = True
                for run in para.runs:
                    run.font.bold = True
                    if 'heading1_size' in rules:
                        run.font.size = Pt(rules['heading1_size'])
        
        if toc_found:
            result.changes.append({'type': '目录格式', 'detail': '已格式化目录标题'})
    
    def _format_page_numbers(self, doc: Document, rules: Dict[str, Any], result: OptimizationResult):
        for section in doc.sections:
            footer = section.footer
            
            if rules.get('page_number_position') == '页脚居中':
                for para in footer.paragraphs:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif rules.get('page_number_position') == '页脚居右':
                for para in footer.paragraphs:
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        result.changes.append({'type': '页码格式', 'detail': '已设置页码位置'})


class TexFormatter(BaseFormatter):
    def get_supported_extensions(self) -> List[str]:
        return ['.tex']
    
    def format(self, input_path: str, output_path: str, template: Template) -> OptimizationResult:
        result = OptimizationResult(success=False, output_path=output_path)
        
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            rules = template.rules
            modified_content = content
            
            modified_content = self._apply_document_class(modified_content, rules, result)
            modified_content = self._apply_packages(modified_content, rules, result)
            modified_content = self._apply_geometry(modified_content, rules, result)
            modified_content = self._apply_font_settings(modified_content, rules, result)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(modified_content)
            
            result.success = True
            
        except Exception as e:
            result.errors.append(f"LaTeX格式化失败: {str(e)}")
        
        return result
    
    def _apply_document_class(self, content: str, rules: Dict[str, Any], result: OptimizationResult) -> str:
        doc_class_match = re.search(r'\\documentclass\[.*?\]\{(\w+)\}', content)
        if not doc_class_match:
            doc_class_match = re.search(r'\\documentclass\{(\w+)\}', content)
        
        if doc_class_match:
            result.changes.append({'type': '文档类', 'detail': f'检测到文档类: {doc_class_match.group(1)}'})
        
        return content
    
    def _apply_packages(self, content: str, rules: Dict[str, Any], result: OptimizationResult) -> str:
        required_packages = [
            '\\usepackage{geometry}',
            '\\usepackage{setspace}',
            '\\usepackage{fontspec}',
        ]
        
        for pkg in required_packages:
            if pkg not in content:
                if '\\begin{document}' in content:
                    content = content.replace('\\begin{document}', f'{pkg}\n\\begin{{document}}')
                    result.changes.append({'type': '宏包', 'detail': f'添加宏包: {pkg}'})
        
        return content
    
    def _apply_geometry(self, content: str, rules: Dict[str, Any], result: OptimizationResult) -> str:
        geometry_settings = []
        
        if 'margin_top' in rules:
            geometry_settings.append(f'top={rules["margin_top"]}cm')
        if 'margin_bottom' in rules:
            geometry_settings.append(f'bottom={rules["margin_bottom"]}cm')
        if 'margin_left' in rules:
            geometry_settings.append(f'left={rules["margin_left"]}cm')
        if 'margin_right' in rules:
            geometry_settings.append(f'right={rules["margin_right"]}cm')
        
        if geometry_settings:
            geometry_cmd = f'\\geometry{{{", ".join(geometry_settings)}}}'
            
            existing_geometry = re.search(r'\\geometry\{.*?\}', content)
            if existing_geometry:
                content = content.replace(existing_geometry.group(0), geometry_cmd)
            else:
                if '\\begin{document}' in content:
                    content = content.replace('\\begin{document}', f'{geometry_cmd}\n\\begin{{document}}')
            
            result.changes.append({'type': '页面设置', 'detail': '已设置页边距'})
        
        return content
    
    def _apply_font_settings(self, content: str, rules: Dict[str, Any], result: OptimizationResult) -> str:
        if 'font_size' in rules:
            fontsize = rules['font_size']
            content = re.sub(
                r'\\documentclass\[(.*?)\]',
                lambda m: f'\\documentclass[{m.group(1)}, {fontsize}pt]' if m.group(1) else f'\\documentclass[{fontsize}pt]',
                content,
                count=1
            )
            result.changes.append({'type': '字号设置', 'detail': f'正文字号设置为 {fontsize}pt'})
        
        if 'line_spacing' in rules:
            linespread = rules['line_spacing']
            linespread_cmd = f'\\linespread{{{linespread}}}'
            
            if '\\linespread' not in content:
                if '\\begin{document}' in content:
                    content = content.replace('\\begin{document}', f'{linespread_cmd}\n\\begin{{document}}')
                    result.changes.append({'type': '行距设置', 'detail': f'行距设置为 {linespread} 倍'})
        
        return content


class DocumentFormatter:
    _formatters: Dict[str, BaseFormatter] = {}
    
    @classmethod
    def register_formatter(cls, formatter: BaseFormatter):
        for ext in formatter.get_supported_extensions():
            cls._formatters[ext.lower()] = formatter
    
    @classmethod
    def format(cls, input_path: str, output_path: str, template: Template) -> OptimizationResult:
        ext = os.path.splitext(input_path)[1].lower()
        
        if ext not in cls._formatters:
            return OptimizationResult(
                success=False,
                output_path=output_path,
                errors=[f"不支持的文件格式: {ext}"]
            )
        
        return cls._formatters[ext].format(input_path, output_path, template)
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        return list(cls._formatters.keys())


DocumentFormatter.register_formatter(DocxFormatter())
DocumentFormatter.register_formatter(TexFormatter())
