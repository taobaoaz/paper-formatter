import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
    QLabel, QPushButton, QLineEdit, QTextEdit, QComboBox, QGroupBox,
    QFileDialog, QMessageBox, QProgressBar, QTabWidget, QTableWidget,
    QTableWidgetItem, QHeaderView, QSplitter, QFrame, QScrollArea,
    QDialog, QDialogButtonBox, QFormLayout, QSpinBox, QDoubleSpinBox,
    QCheckBox, QListWidget, QListWidgetItem, QStatusBar, QAction,
    QToolBar, QMenu, QMenuBar, QStackedWidget, QSizePolicy
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QSize, QTimer
from PyQt5.QtGui import QFont, QIcon, QColor, QPalette
from typing import Dict, Any, Optional, List
import json
from datetime import datetime

PYWINSTYLES_AVAILABLE = False
try:
    import pywinstyles
    PYWINSTYLES_AVAILABLE = True
except ImportError:
    pass

try:
    from qfluentwidgets import (
        setTheme, Theme, setThemeColor, FluentWindow,
        PushButton, CardWidget, PrimaryPushButton,
        ComboBox, LineEdit, TextEdit, SpinBox, DoubleSpinBox,
        CheckBox, ListWidget, TableWidget, TabWidget, ScrollArea,
        TitleLabel, SubtitleLabel, BodyLabel, CaptionLabel
    )
    FLUENT_WIDGETS_AVAILABLE = True
except ImportError as e:
    FLUENT_WIDGETS_AVAILABLE = False
    import traceback
    FLUENT_IMPORT_ERROR = str(e)
    FLUENT_IMPORT_TRACEBACK = traceback.format_exc()

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from parsers import DocumentParserFactory, ParsedDocument
from ai_recognizer import AIRecognizer, TemplateRecognitionResult
from template_manager import TemplateManager, Template
from formatter import DocumentFormatter, OptimizationResult
from ai_prompt import generate_ai_prompt, parse_ai_response, ParsedTemplateData, get_quick_prompt_for_template_file
from document_generator import DocumentGenerator, get_fields_by_group, get_all_fields, DEFAULT_FIELDS, SECTION_TEMPLATES
from reference_formatter import ReferenceManager, ReferenceParser, ReferenceFormatter, ReferenceFormat
from error_handler import ErrorHandler, ErrorCode, create_error, get_error_message
from template_spec_parser import parse_template_spec
from word_format_extractor import extract_format_from_docx, WordFormatExtractor
import config


class WorkerThread(QThread):
    progress = pyqtSignal(str)
    finished = pyqtSignal(object)
    error = pyqtSignal(str)
    
    def __init__(self, task_func, *args, **kwargs):
        super().__init__()
        self.task_func = task_func
        self.args = args
        self.kwargs = kwargs
    
    def run(self):
        try:
            result = self.task_func(*self.args, **self.kwargs)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))


class TemplateEditDialog(QDialog):
    def __init__(self, template: Template = None, parent=None):
        super().__init__(parent)
        self.template = template
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('编辑模板')
        self.setMinimumSize(600, 500)
        
        layout = QVBoxLayout(self)
        
        form_layout = QFormLayout()
        
        self.name_edit = QLineEdit()
        self.name_edit.setText(self.template.template_name if self.template else '')
        form_layout.addRow('模板名称:', self.name_edit)
        
        self.type_combo = QComboBox()
        self.type_combo.addItems(['中文期刊', '国际期刊', '学位论文', '自定义'])
        if self.template:
            index = self.type_combo.findText(self.template.template_type)
            if index >= 0:
                self.type_combo.setCurrentIndex(index)
        form_layout.addRow('模板类型:', self.type_combo)
        
        self.desc_edit = QTextEdit()
        self.desc_edit.setMaximumHeight(60)
        self.desc_edit.setPlainText(self.template.description if self.template else '')
        form_layout.addRow('描述:', self.desc_edit)
        
        layout.addLayout(form_layout)
        
        rules_group = QGroupBox('格式规则')
        rules_layout = QFormLayout(rules_group)
        
        self.font_family = QLineEdit()
        self.font_family.setText(str(self.template.rules.get('font_family', '宋体')) if self.template else '宋体')
        rules_layout.addRow('正文字体:', self.font_family)
        
        self.font_size = QDoubleSpinBox()
        self.font_size.setRange(5, 72)
        self.font_size.setValue(self.template.rules.get('font_size', 12) if self.template else 12)
        rules_layout.addRow('正文字号(pt):', self.font_size)
        
        self.line_spacing = QDoubleSpinBox()
        self.line_spacing.setRange(0.5, 3.0)
        self.line_spacing.setSingleStep(0.25)
        self.line_spacing.setValue(self.template.rules.get('line_spacing', 1.5) if self.template else 1.5)
        rules_layout.addRow('行距倍数:', self.line_spacing)
        
        margin_layout = QHBoxLayout()
        self.margin_top = QDoubleSpinBox()
        self.margin_top.setRange(0, 10)
        self.margin_top.setValue(self.template.rules.get('margin_top', 2.54) if self.template else 2.54)
        margin_layout.addWidget(QLabel('上:'))
        margin_layout.addWidget(self.margin_top)
        
        self.margin_bottom = QDoubleSpinBox()
        self.margin_bottom.setRange(0, 10)
        self.margin_bottom.setValue(self.template.rules.get('margin_bottom', 2.54) if self.template else 2.54)
        margin_layout.addWidget(QLabel('下:'))
        margin_layout.addWidget(self.margin_bottom)
        
        self.margin_left = QDoubleSpinBox()
        self.margin_left.setRange(0, 10)
        self.margin_left.setValue(self.template.rules.get('margin_left', 3.17) if self.template else 3.17)
        margin_layout.addWidget(QLabel('左:'))
        margin_layout.addWidget(self.margin_left)
        
        self.margin_right = QDoubleSpinBox()
        self.margin_right.setRange(0, 10)
        self.margin_right.setValue(self.template.rules.get('margin_right', 3.17) if self.template else 3.17)
        margin_layout.addWidget(QLabel('右:'))
        margin_layout.addWidget(self.margin_right)
        
        rules_layout.addRow('页边距(cm):', margin_layout)
        
        self.title_font = QLineEdit()
        self.title_font.setText(str(self.template.rules.get('title_font', '黑体')) if self.template else '黑体')
        rules_layout.addRow('标题字体:', self.title_font)
        
        self.title_size = QDoubleSpinBox()
        self.title_size.setRange(5, 72)
        self.title_size.setValue(self.template.rules.get('title_size', 22) if self.template else 22)
        rules_layout.addRow('标题字号(pt):', self.title_size)
        
        self.h1_size = QDoubleSpinBox()
        self.h1_size.setRange(5, 72)
        self.h1_size.setValue(self.template.rules.get('heading1_size', 16) if self.template else 16)
        rules_layout.addRow('一级标题字号(pt):', self.h1_size)
        
        self.h2_size = QDoubleSpinBox()
        self.h2_size.setRange(5, 72)
        self.h2_size.setValue(self.template.rules.get('heading2_size', 14) if self.template else 14)
        rules_layout.addRow('二级标题字号(pt):', self.h2_size)
        
        self.h3_size = QDoubleSpinBox()
        self.h3_size.setRange(5, 72)
        self.h3_size.setValue(self.template.rules.get('heading3_size', 12) if self.template else 12)
        rules_layout.addRow('三级标题字号(pt):', self.h3_size)
        
        self.ref_format = QComboBox()
        self.ref_format.addItems(['GB/T 7714', 'IEEE', 'APA', 'MLA', 'Springer'])
        ref_format_text = self.template.rules.get('reference_format', 'GB/T 7714') if self.template else 'GB/T 7714'
        index = self.ref_format.findText(ref_format_text)
        if index >= 0:
            self.ref_format.setCurrentIndex(index)
        rules_layout.addRow('参考文献格式:', self.ref_format)
        
        layout.addWidget(rules_group)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def get_template_data(self) -> Dict[str, Any]:
        return {
            'template_name': self.name_edit.text(),
            'template_type': self.type_combo.currentText(),
            'description': self.desc_edit.toPlainText(),
            'rules': {
                'font_family': self.font_family.text(),
                'font_size': self.font_size.value(),
                'line_spacing': self.line_spacing.value(),
                'margin_top': self.margin_top.value(),
                'margin_bottom': self.margin_bottom.value(),
                'margin_left': self.margin_left.value(),
                'margin_right': self.margin_right.value(),
                'title_font': self.title_font.text(),
                'title_size': self.title_size.value(),
                'heading1_size': self.h1_size.value(),
                'heading2_size': self.h2_size.value(),
                'heading3_size': self.h3_size.value(),
                'reference_format': self.ref_format.currentText()
            }
        }


class AIDataImportDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parsed_data = None
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('导入AI返回数据')
        self.setMinimumSize(700, 500)
        
        layout = QVBoxLayout(self)
        
        help_label = QLabel(
            '使用说明:\n'
            '1. 将模板文件内容复制到下方"模板内容"区域\n'
            '2. 点击"生成提示词"按钮，复制生成的提示词\n'
            '3. 将提示词发送给豆包/千问等AI APP\n'
            '4. 将AI返回的JSON数据粘贴到"AI返回数据"区域\n'
            '5. 点击"解析并导入"'
        )
        help_label.setStyleSheet('background-color: #f0f0f0; padding: 10px; border-radius: 5px;')
        layout.addWidget(help_label)
        
        tabs = QTabWidget()
        
        prompt_tab = QWidget()
        prompt_layout = QVBoxLayout(prompt_tab)
        
        prompt_layout.addWidget(QLabel('模板文件内容（可选，用于生成提示词）:'))
        self.template_content = QTextEdit()
        self.template_content.setPlaceholderText('将模板文件的内容粘贴到这里...\n\n支持Word文档的文字内容或LaTeX源码')
        prompt_layout.addWidget(self.template_content)
        
        generate_btn = QPushButton('生成提示词')
        generate_btn.clicked.connect(self.generate_prompt)
        prompt_layout.addWidget(generate_btn)
        
        prompt_layout.addWidget(QLabel('生成的提示词（复制发送给AI）:'))
        self.generated_prompt = QTextEdit()
        self.generated_prompt.setReadOnly(True)
        prompt_layout.addWidget(self.generated_prompt)
        
        copy_prompt_btn = QPushButton('复制提示词')
        copy_prompt_btn.clicked.connect(self.copy_prompt)
        prompt_layout.addWidget(copy_prompt_btn)
        
        tabs.addTab(prompt_tab, '生成提示词')
        
        import_tab = QWidget()
        import_layout = QVBoxLayout(import_tab)
        
        import_layout.addWidget(QLabel('AI返回的JSON数据:'))
        self.ai_response = QTextEdit()
        self.ai_response.setPlaceholderText('将AI返回的JSON数据粘贴到这里...\n\n例如:\n{\n  "template_name": "xxx期刊模板",\n  "rules": {...}\n}')
        import_layout.addWidget(self.ai_response)
        
        parse_btn = QPushButton('解析并导入')
        parse_btn.clicked.connect(self.parse_and_import)
        import_layout.addWidget(parse_btn)
        
        import_layout.addWidget(QLabel('解析预览:'))
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        import_layout.addWidget(self.preview_text)
        
        tabs.addTab(import_tab, '导入数据')
        
        layout.addWidget(tabs)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def generate_prompt(self):
        content = self.template_content.toPlainText()
        if not content.strip():
            prompt = get_quick_prompt_for_template_file()
        else:
            prompt = generate_ai_prompt(content)
        self.generated_prompt.setPlainText(prompt)
    
    def copy_prompt(self):
        text = self.generated_prompt.toPlainText()
        if text:
            QApplication.clipboard().setText(text)
            QMessageBox.information(self, '成功', '提示词已复制到剪贴板')
    
    def parse_and_import(self):
        response_text = self.ai_response.toPlainText()
        if not response_text.strip():
            QMessageBox.warning(self, '提示', '请先粘贴AI返回的数据')
            return
        
        self.parsed_data = parse_ai_response(response_text)
        
        if not self.parsed_data.is_valid:
            QMessageBox.warning(self, '解析失败', self.parsed_data.error_message)
            return
        
        preview = f"""模板名称: {self.parsed_data.template_name}
模板类型: {self.parsed_data.template_type}
描述: {self.parsed_data.description}

规则预览:
- 正文字体: {self.parsed_data.rules.get('font_family', '未设置')}
- 正文字号: {self.parsed_data.rules.get('font_size', '未设置')}pt
- 行距: {self.parsed_data.rules.get('line_spacing', '未设置')}倍
- 页边距: 上{self.parsed_data.rules.get('margin_top', '-')}cm / 下{self.parsed_data.rules.get('margin_bottom', '-')}cm
- 参考文献格式: {self.parsed_data.rules.get('reference_format', '未设置')}"""
        
        self.preview_text.setPlainText(preview)
        QMessageBox.information(self, '成功', '数据解析成功，点击确定导入')


class StructureItemDialog(QDialog):
    def __init__(self, parent=None, item_type: str = '部分', title: str = '新建', data: dict = None):
        super().__init__(parent)
        self.item_type = item_type
        self.data = data or {}
        self.setWindowTitle(title)
        self.setMinimumWidth(400)
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        form_layout = QFormLayout()
        
        self.type_combo = QComboBox()
        self.type_combo.addItems(['部分', '章节'])
        self.type_combo.setCurrentText(self.item_type)
        form_layout.addRow('类型:', self.type_combo)
        
        self.title_edit = QLineEdit()
        self.title_edit.setText(self.data.get('title', ''))
        self.title_edit.setPlaceholderText('输入标题')
        form_layout.addRow('标题:', self.title_edit)
        
        self.desc_edit = QTextEdit()
        self.desc_edit.setPlaceholderText('输入描述（可选）')
        self.desc_edit.setText(self.data.get('description', ''))
        self.desc_edit.setMaximumHeight(80)
        form_layout.addRow('描述:', self.desc_edit)
        
        layout.addLayout(form_layout)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def get_data(self) -> dict:
        return {
            'type': self.type_combo.currentText(),
            'title': self.title_edit.text(),
            'description': self.desc_edit.toPlainText()
        }


class SectionFormatDialog(QDialog):
    def __init__(self, parent=None, current_settings: dict = None):
        super().__init__(parent)
        self.setWindowTitle('章节排版设置')
        self.setMinimumWidth(450)
        self.settings = current_settings or {}
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        desc_label = QLabel('为此章节单独设置排版格式（留空则使用全局默认）')
        desc_label.setStyleSheet('color: #666; font-size: 11px; margin-bottom: 10px;')
        layout.addWidget(desc_label)
        
        format_group = QGroupBox('段落格式')
        format_layout = QFormLayout(format_group)
        
        self.indent_check = QCheckBox('自定义首行缩进')
        self.indent_check.setChecked(bool(self.settings.get('first_line_indent')))
        format_layout.addRow(self.indent_check)
        
        self.indent_spin = QSpinBox()
        self.indent_spin.setRange(0, 10)
        self.indent_spin.setValue(self.settings.get('first_line_indent', 2))
        self.indent_spin.setSuffix(' 字符')
        self.indent_spin.setEnabled(bool(self.settings.get('first_line_indent')))
        self.indent_check.stateChanged.connect(lambda s: self.indent_spin.setEnabled(s))
        format_layout.addRow('缩进:', self.indent_spin)
        
        self.line_spacing_check = QCheckBox('自定义行距')
        self.line_spacing_check.setChecked(bool(self.settings.get('line_spacing')))
        format_layout.addRow(self.line_spacing_check)
        
        self.line_spacing_spin = QDoubleSpinBox()
        self.line_spacing_spin.setRange(0.5, 3.0)
        self.line_spacing_spin.setSingleStep(0.1)
        self.line_spacing_spin.setValue(self.settings.get('line_spacing', 1.5))
        self.line_spacing_spin.setSuffix(' 倍')
        self.line_spacing_spin.setEnabled(bool(self.settings.get('line_spacing')))
        self.line_spacing_check.stateChanged.connect(lambda s: self.line_spacing_spin.setEnabled(s))
        format_layout.addRow('行距:', self.line_spacing_spin)
        
        self.space_before_check = QCheckBox('自定义段前间距')
        self.space_before_check.setChecked(bool(self.settings.get('space_before')))
        format_layout.addRow(self.space_before_check)
        
        self.space_before_spin = QSpinBox()
        self.space_before_spin.setRange(0, 50)
        self.space_before_spin.setValue(self.settings.get('space_before', 0))
        self.space_before_spin.setSuffix(' 磅')
        self.space_before_spin.setEnabled(bool(self.settings.get('space_before')))
        self.space_before_check.stateChanged.connect(lambda s: self.space_before_spin.setEnabled(s))
        format_layout.addRow('段前:', self.space_before_spin)
        
        self.space_after_check = QCheckBox('自定义段后间距')
        self.space_after_check.setChecked(bool(self.settings.get('space_after')))
        format_layout.addRow(self.space_after_check)
        
        self.space_after_spin = QSpinBox()
        self.space_after_spin.setRange(0, 50)
        self.space_after_spin.setValue(self.settings.get('space_after', 0))
        self.space_after_spin.setSuffix(' 磅')
        self.space_after_spin.setEnabled(bool(self.settings.get('space_after')))
        self.space_after_check.stateChanged.connect(lambda s: self.space_after_spin.setEnabled(s))
        format_layout.addRow('段后:', self.space_after_spin)
        
        layout.addWidget(format_group)
        
        align_group = QGroupBox('对齐方式')
        align_layout = QVBoxLayout(align_group)
        
        self.alignment_combo = QComboBox()
        self.alignment_combo.addItems(['默认', '两端对齐', '左对齐', '居中', '右对齐', '分散对齐'])
        if self.settings.get('alignment'):
            index = self.alignment_combo.findText(self.settings['alignment'])
            if index >= 0:
                self.alignment_combo.setCurrentIndex(index)
        align_layout.addWidget(self.alignment_combo)
        
        layout.addWidget(align_group)
        
        font_group = QGroupBox('字体设置')
        font_layout = QFormLayout(font_group)
        
        self.font_name_edit = QLineEdit()
        self.font_name_edit.setText(self.settings.get('font_name', ''))
        self.font_name_edit.setPlaceholderText('留空使用默认')
        font_layout.addRow('字体:', self.font_name_edit)
        
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(5, 72)
        self.font_size_spin.setValue(self.settings.get('font_size', 12))
        self.font_size_spin.setSuffix(' pt')
        font_layout.addRow('字号:', self.font_size_spin)
        
        self.font_size_check = QCheckBox('自定义字号')
        self.font_size_check.setChecked(bool(self.settings.get('font_size')))
        self.font_size_spin.setEnabled(bool(self.settings.get('font_size')))
        self.font_size_check.stateChanged.connect(lambda s: self.font_size_spin.setEnabled(s))
        font_layout.addRow(self.font_size_check)
        
        layout.addWidget(font_group)
        
        btn_layout = QHBoxLayout()
        
        reset_btn = QPushButton('重置为默认')
        reset_btn.clicked.connect(self.reset_to_default)
        btn_layout.addWidget(reset_btn)
        
        btn_layout.addStretch()
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        btn_layout.addWidget(buttons)
        
        layout.addLayout(btn_layout)
    
    def reset_to_default(self):
        self.indent_check.setChecked(False)
        self.line_spacing_check.setChecked(False)
        self.space_before_check.setChecked(False)
        self.space_after_check.setChecked(False)
        self.alignment_combo.setCurrentIndex(0)
        self.font_name_edit.clear()
        self.font_size_check.setChecked(False)
    
    def get_settings(self) -> dict:
        settings = {}
        
        if self.indent_check.isChecked():
            settings['first_line_indent'] = self.indent_spin.value()
        if self.line_spacing_check.isChecked():
            settings['line_spacing'] = self.line_spacing_spin.value()
        if self.space_before_check.isChecked():
            settings['space_before'] = self.space_before_spin.value()
        if self.space_after_check.isChecked():
            settings['space_after'] = self.space_after_spin.value()
        
        alignment = self.alignment_combo.currentText()
        if alignment != '默认':
            settings['alignment'] = alignment
        
        if self.font_name_edit.text().strip():
            settings['font_name'] = self.font_name_edit.text().strip()
        if self.font_size_check.isChecked():
            settings['font_size'] = self.font_size_spin.value()
        
        return settings


class QuickSpecParseDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parsed_rules = None
        self.parse_summary = []
        self.extracted_formats = []
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('创建模板 - 选择方式')
        self.setMinimumSize(900, 700)
        
        layout = QVBoxLayout(self)
        
        help_label = QLabel(
            '选择模板创建方式:\n'
            '• 方式一：粘贴文字说明 - 直接粘贴模板说明文件的文字内容\n'
            '• 方式二：解析标注文档 - 打开带格式标注的模板文档\n'
            '• 方式三：提取文档格式 - 直接从Word模板文档提取格式设置'
        )
        help_label.setStyleSheet('background-color: #e3f2fd; padding: 12px; border-radius: 5px; font-size: 12px;')
        layout.addWidget(help_label)
        
        mode_tabs = QTabWidget()
        
        ai_tab = self._create_ai_tab()
        mode_tabs.addTab(ai_tab, '🤖 AI识别创建')
        
        text_tab = self._create_text_tab()
        mode_tabs.addTab(text_tab, '📝 粘贴文字说明')
        
        annotated_tab = self._create_annotated_tab()
        mode_tabs.addTab(annotated_tab, '📋 解析标注文档')
        
        format_tab = self._create_format_extract_tab()
        mode_tabs.addTab(format_tab, '📄 提取文档格式')
        
        layout.addWidget(mode_tabs)
        
        result_group = QGroupBox('解析结果')
        result_layout = QVBoxLayout(result_group)
        
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(4)
        self.result_table.setHorizontalHeaderLabels(['规则项', '识别值', '来源', '置信度'])
        self.result_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.result_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.result_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.result_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.result_table.setMaximumHeight(180)
        result_layout.addWidget(self.result_table)
        
        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel('模板名称:'))
        self.template_name_edit = QLineEdit()
        self.template_name_edit.setPlaceholderText('输入模板名称（如：兰州文理学院本科论文）')
        name_layout.addWidget(self.template_name_edit)
        result_layout.addLayout(name_layout)
        
        layout.addWidget(result_group)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def _create_ai_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc = QLabel('适用于：使用豆包/千问等AI识别模板文件，获取完整格式要求')
        desc.setStyleSheet('color: #666; font-style: italic;')
        layout.addWidget(desc)
        
        steps_group = QGroupBox('使用步骤')
        steps_layout = QVBoxLayout(steps_group)
        
        steps_text = QLabel(
            '第一步：复制下方提示词\n'
            '第二步：将模板文件发送给AI（豆包/千问等）\n'
            '第三步：将提示词发送给AI\n'
            '第四步：复制AI返回的JSON结果，粘贴到下方\n'
            '第五步：点击"解析JSON"完成创建'
        )
        steps_text.setStyleSheet('padding: 10px;')
        steps_layout.addWidget(steps_text)
        layout.addWidget(steps_group)
        
        prompt_group = QGroupBox('提示词（复制发送给AI）')
        prompt_layout = QVBoxLayout(prompt_group)
        
        self.ai_prompt_text = QTextEdit()
        from ai_prompt import AI_PROMPT_TEMPLATE
        self.ai_prompt_text.setPlainText(AI_PROMPT_TEMPLATE.replace('{template_content}', '[在此处粘贴模板文件内容]'))
        self.ai_prompt_text.setMaximumHeight(200)
        prompt_layout.addWidget(self.ai_prompt_text)
        
        copy_prompt_btn = QPushButton('📋 复制提示词')
        copy_prompt_btn.setStyleSheet('background-color: #2196F3; color: white; padding: 8px;')
        copy_prompt_btn.clicked.connect(self.copy_ai_prompt)
        prompt_layout.addWidget(copy_prompt_btn)
        
        layout.addWidget(prompt_group)
        
        result_group = QGroupBox('粘贴AI返回的JSON结果')
        result_layout = QVBoxLayout(result_group)
        
        self.ai_result_text = QTextEdit()
        self.ai_result_text.setPlaceholderText('将AI返回的JSON结果粘贴到这里...')
        self.ai_result_text.setMaximumHeight(150)
        result_layout.addWidget(self.ai_result_text)
        
        parse_json_btn = QPushButton('🔍 解析JSON')
        parse_json_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 10px; font-weight: bold;')
        parse_json_btn.clicked.connect(self.parse_ai_json)
        result_layout.addWidget(parse_json_btn)
        
        layout.addWidget(result_group)
        
        return tab
    
    def copy_ai_prompt(self):
        prompt = self.ai_prompt_text.toPlainText()
        if prompt:
            QApplication.clipboard().setText(prompt)
            QMessageBox.information(self, '成功', '提示词已复制到剪贴板\n\n请将模板文件内容粘贴到提示词中\n然后发送给AI')
        else:
            QMessageBox.warning(self, '提示', '提示词为空')
    
    def parse_ai_json(self):
        json_text = self.ai_result_text.toPlainText()
        if not json_text.strip():
            QMessageBox.warning(self, '提示', '请先粘贴AI返回的JSON结果')
            return
        
        try:
            from ai_prompt import parse_ai_response
            parsed_data = parse_ai_response(json_text)
            
            if not parsed_data.is_valid:
                QMessageBox.warning(self, '解析失败', parsed_data.error_message)
                return
            
            self.parsed_rules = parsed_data.rules
            self.template_name_edit.setText(parsed_data.template_name)
            
            self.parse_summary = []
            for key, value in parsed_data.rules.items():
                if value is not None:
                    self.parse_summary.append((key, value, 'AI识别', 0.95))
            
            self._display_parse_results()
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'解析失败: {str(e)}')
    
    def _create_text_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc = QLabel('适用于：学校提供的是文字形式的格式说明文件')
        desc.setStyleSheet('color: #666; font-style: italic;')
        layout.addWidget(desc)
        
        self.spec_text = QTextEdit()
        self.spec_text.setPlaceholderText(
            '请粘贴模板说明文件的内容...\n\n'
            '示例:\n'
            '正文采用宋体，小四号字，行距1.5倍。\n'
            '页边距：上下2.54cm，左右3.17cm。\n'
            '论文标题使用黑体二号字。\n'
            '一级标题使用黑体三号字...'
        )
        layout.addWidget(self.spec_text)
        
        parse_btn = QPushButton('🔍 解析文字说明')
        parse_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 10px; font-weight: bold;')
        parse_btn.clicked.connect(self.parse_text_spec)
        layout.addWidget(parse_btn)
        
        return tab
    
    def _create_annotated_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc = QLabel('适用于：文档中包含格式标注（如【标题：黑体二号】）')
        desc.setStyleSheet('color: #666; font-style: italic;')
        layout.addWidget(desc)
        
        btn_layout = QHBoxLayout()
        open_btn = QPushButton('📂 打开标注文档')
        open_btn.setStyleSheet('background-color: #2196F3; color: white; padding: 8px;')
        open_btn.clicked.connect(self.open_annotated_file)
        btn_layout.addWidget(open_btn)
        
        self.annotated_file_label = QLabel('未选择文件')
        self.annotated_file_label.setStyleSheet('color: gray;')
        btn_layout.addWidget(self.annotated_file_label)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        self.annotated_text = QTextEdit()
        self.annotated_text.setReadOnly(True)
        self.annotated_text.setPlaceholderText(
            '打开文档后，内容将显示在这里...\n\n'
            '支持的标注格式:\n'
            '【标题：黑体二号，居中】论文标题示例\n'
            '【正文：宋体小四，1.5倍行距】正文内容...\n'
            '【一级标题：黑体三号】第一章 绪论'
        )
        layout.addWidget(self.annotated_text)
        
        parse_btn = QPushButton('🔍 解析标注内容')
        parse_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 10px; font-weight: bold;')
        parse_btn.clicked.connect(self.parse_annotated)
        layout.addWidget(parse_btn)
        
        return tab
    
    def _create_format_extract_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc = QLabel('适用于：学校提供的是已设置好格式的Word模板文档（直接提取文档中的格式设置）')
        desc.setStyleSheet('color: #666; font-style: italic;')
        layout.addWidget(desc)
        
        btn_layout = QHBoxLayout()
        open_btn = QPushButton('📂 打开模板文档')
        open_btn.setStyleSheet('background-color: #FF9800; color: white; padding: 8px;')
        open_btn.clicked.connect(self.open_format_file)
        btn_layout.addWidget(open_btn)
        
        self.format_file_label = QLabel('未选择文件')
        self.format_file_label.setStyleSheet('color: gray;')
        btn_layout.addWidget(self.format_file_label)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        self.format_preview_table = QTableWidget()
        self.format_preview_table.setColumnCount(5)
        self.format_preview_table.setHorizontalHeaderLabels(['元素类型', '内容示例', '字体', '字号', '行距'])
        self.format_preview_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.format_preview_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.format_preview_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.format_preview_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        self.format_preview_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeToContents)
        layout.addWidget(self.format_preview_table)
        
        extract_btn = QPushButton('🔍 提取格式设置')
        extract_btn.setStyleSheet('background-color: #FF9800; color: white; padding: 10px; font-weight: bold;')
        extract_btn.clicked.connect(self.extract_format)
        layout.addWidget(extract_btn)
        
        return tab
    
    def open_annotated_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '选择标注文档', '',
            'Word文档 (*.docx);;文本文件 (*.txt);;所有文件 (*.*)'
        )
        
        if file_path:
            self.annotated_file_label.setText(os.path.basename(file_path))
            self.annotated_file_label.setStyleSheet('color: black;')
            
            try:
                if file_path.endswith('.docx'):
                    text = self._extract_text_from_docx(file_path)
                else:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                
                self.annotated_text.setPlainText(text)
                self.spec_text.setPlainText(text)
                
            except Exception as e:
                QMessageBox.critical(self, '错误', f'读取文件失败: {str(e)}')
    
    def open_format_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '选择模板文档', '',
            'Word文档 (*.docx)'
        )
        
        if file_path:
            self.format_file_label.setText(os.path.basename(file_path))
            self.format_file_label.setStyleSheet('color: black;')
            self._current_format_file = file_path
            
            try:
                rules, formats = extract_format_from_docx(file_path)
                self.extracted_formats = formats
                
                self.format_preview_table.setRowCount(len(formats))
                
                for row, fmt in enumerate(formats):
                    self.format_preview_table.setItem(row, 0, QTableWidgetItem(fmt.element_type))
                    self.format_preview_table.setItem(row, 1, QTableWidgetItem(fmt.content_sample[:30]))
                    self.format_preview_table.setItem(row, 2, QTableWidgetItem(fmt.font_name))
                    self.format_preview_table.setItem(row, 3, QTableWidgetItem(f'{fmt.font_size}pt'))
                    self.format_preview_table.setItem(row, 4, QTableWidgetItem(f'{fmt.line_spacing}倍'))
                
                if not self.template_name_edit.text():
                    self.template_name_edit.setText(os.path.splitext(os.path.basename(file_path))[0])
                
            except Exception as e:
                QMessageBox.critical(self, '错误', f'读取文件失败: {str(e)}')
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts)
    
    def parse_text_spec(self):
        text = self.spec_text.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, '提示', '请先粘贴模板说明内容')
            return
        
        self._do_parse_text(text)
    
    def parse_annotated(self):
        text = self.annotated_text.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, '提示', '请先打开标注文档')
            return
        
        self._do_parse_text(text)
    
    def _do_parse_text(self, text: str):
        try:
            self.parsed_rules, self.parse_summary = parse_template_spec(text)
            self._display_parse_results()
        except Exception as e:
            QMessageBox.critical(self, '错误', f'解析失败: {str(e)}')
    
    def extract_format(self):
        if not hasattr(self, '_current_format_file') or not self._current_format_file:
            QMessageBox.warning(self, '提示', '请先打开模板文档')
            return
        
        try:
            self.parsed_rules, self.extracted_formats = extract_format_from_docx(self._current_format_file)
            
            self.parse_summary = []
            for fmt in self.extracted_formats:
                self.parse_summary.append((
                    fmt.element_type,
                    fmt.font_size,
                    fmt.content_sample[:30],
                    fmt.confidence
                ))
            
            self._display_parse_results()
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'提取失败: {str(e)}')
    
    def _display_parse_results(self):
        if not self.parse_summary:
            QMessageBox.warning(self, '提示', '未能识别出格式规则')
            return
        
        self.result_table.setRowCount(len(self.parse_summary))
        
        for row, item in enumerate(self.parse_summary):
            rule_name = item[0]
            value = item[1]
            source = item[2] if len(item) > 2 else ''
            confidence = item[3] if len(item) > 3 else 0.8
            
            self.result_table.setItem(row, 0, QTableWidgetItem(str(rule_name)))
            self.result_table.setItem(row, 1, QTableWidgetItem(str(value)))
            self.result_table.setItem(row, 2, QTableWidgetItem(str(source)[:40]))
            
            conf_item = QTableWidgetItem(f'{confidence:.0%}' if isinstance(confidence, float) else str(confidence))
            if isinstance(confidence, float):
                if confidence >= 0.9:
                    conf_item.setBackground(QColor(200, 230, 201))
                elif confidence >= 0.7:
                    conf_item.setBackground(QColor(255, 243, 224))
                else:
                    conf_item.setBackground(QColor(255, 224, 178))
            self.result_table.setItem(row, 3, conf_item)
        
        QMessageBox.information(self, '解析完成', f'成功识别 {len(self.parse_summary)} 项格式规则')
    
    def get_template_data(self) -> Dict[str, Any]:
        return {
            'template_name': self.template_name_edit.text() or '自定义模板',
            'template_type': '自定义',
            'description': '从模板文件解析创建',
            'rules': self.parsed_rules or {}
        }


class SettingsDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('设置')
        self.setMinimumSize(500, 300)
        
        layout = QVBoxLayout(self)
        
        api_group = QGroupBox('AI API 设置（可选）')
        api_layout = QFormLayout(api_group)
        
        self.api_url = QLineEdit()
        self.api_url.setText(config.AI_API_CONFIG.get('api_url', ''))
        self.api_url.setPlaceholderText('例如: https://api.openai.com/v1/chat/completions')
        api_layout.addRow('API URL:', self.api_url)
        
        self.api_key = QLineEdit()
        self.api_key.setEchoMode(QLineEdit.Password)
        self.api_key.setText(config.AI_API_CONFIG.get('api_key', ''))
        api_layout.addRow('API Key:', self.api_key)
        
        self.model_name = QLineEdit()
        self.model_name.setText(config.AI_API_CONFIG.get('model', ''))
        self.model_name.setPlaceholderText('例如: gpt-4, gpt-3.5-turbo')
        api_layout.addRow('模型名称:', self.model_name)
        
        layout.addWidget(api_group)
        
        note_label = QLabel(
            '提示:\n'
            '• 方式一: 配置AI API后，程序可直接调用API识别模板\n'
            '• 方式二: 使用"导入AI数据"功能，手动粘贴豆包/千问等AI返回的数据'
        )
        note_label.setStyleSheet('color: gray; font-style: italic;')
        layout.addWidget(note_label)
        
        layout.addStretch()
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self.save_settings)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def save_settings(self):
        config.AI_API_CONFIG['api_url'] = self.api_url.text()
        config.AI_API_CONFIG['api_key'] = self.api_key.text()
        config.AI_API_CONFIG['model'] = self.model_name.text()
        self.accept()


class DocumentCreatorDialog(QDialog):
    def __init__(self, parent=None, template_manager=None):
        super().__init__(parent)
        self.template_manager = template_manager
        self.doc_generator = DocumentGenerator()
        self.field_widgets = {}
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('论文快速生成')
        self.setMinimumSize(900, 700)
        
        layout = QVBoxLayout(self)
        
        top_layout = QHBoxLayout()
        
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel('文档类型:'))
        self.doc_type_combo = QComboBox()
        self.doc_type_combo.addItems(['中文期刊', '国际期刊', '学位论文', '自定义'])
        self.doc_type_combo.currentTextChanged.connect(self.on_doc_type_changed)
        type_layout.addWidget(self.doc_type_combo)
        top_layout.addLayout(type_layout)
        
        top_layout.addStretch()
        
        template_layout = QHBoxLayout()
        template_layout.addWidget(QLabel('格式模板:'))
        self.format_template_combo = QComboBox()
        self.load_format_templates()
        self.format_template_combo.currentIndexChanged.connect(self.on_format_template_changed)
        template_layout.addWidget(self.format_template_combo)
        top_layout.addLayout(template_layout)
        
        layout.addLayout(top_layout)
        
        self.field_tabs = QTabWidget()
        
        groups = get_fields_by_group()
        for group_name, fields in groups.items():
            tab = self.create_field_tab(fields)
            self.field_tabs.addTab(tab, group_name)
        
        layout.addWidget(self.field_tabs)
        
        btn_layout = QHBoxLayout()
        
        preview_btn = QPushButton('预览')
        preview_btn.clicked.connect(self.preview_document)
        btn_layout.addWidget(preview_btn)
        
        generate_btn = QPushButton('生成文档')
        generate_btn.clicked.connect(self.generate_document)
        btn_layout.addWidget(generate_btn)
        
        save_data_btn = QPushButton('保存草稿')
        save_data_btn.clicked.connect(self.save_draft)
        btn_layout.addWidget(save_data_btn)
        
        load_data_btn = QPushButton('加载草稿')
        load_data_btn.clicked.connect(self.load_draft)
        btn_layout.addWidget(load_data_btn)
        
        clear_btn = QPushButton('清空')
        clear_btn.clicked.connect(self.clear_fields)
        btn_layout.addWidget(clear_btn)
        
        layout.addLayout(btn_layout)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Close)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
    
    def create_field_tab(self, fields: List) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_layout = QFormLayout(scroll_widget)
        scroll_layout.setSpacing(10)
        
        for field in fields:
            label_text = f"{field.field_name}{'*' if field.required else ''}:"
            
            if field.field_type == 'textarea':
                widget = QTextEdit()
                widget.setPlaceholderText(field.placeholder)
                widget.setMinimumHeight(100)
            else:
                widget = QLineEdit()
                widget.setPlaceholderText(field.placeholder)
                if field.default_value:
                    widget.setText(field.default_value)
            
            self.field_widgets[field.field_id] = widget
            scroll_layout.addRow(label_text, widget)
        
        scroll.setWidget(scroll_widget)
        layout.addWidget(scroll)
        
        return tab
    
    def load_format_templates(self):
        self.format_template_combo.clear()
        self.format_template_combo.addItem('默认格式', None)
        
        if self.template_manager:
            templates = self.template_manager.get_all_templates()
            for template in templates:
                self.format_template_combo.addItem(template.template_name, template)
    
    def on_doc_type_changed(self, doc_type: str):
        pass
    
    def on_format_template_changed(self, index: int):
        pass
    
    def get_field_values(self) -> Dict[str, str]:
        values = {}
        for field_id, widget in self.field_widgets.items():
            if isinstance(widget, QTextEdit):
                values[field_id] = widget.toPlainText()
            else:
                values[field_id] = widget.text()
        return values
    
    def set_field_values(self, values: Dict[str, str]):
        for field_id, value in values.items():
            if field_id in self.field_widgets:
                widget = self.field_widgets[field_id]
                if isinstance(widget, QTextEdit):
                    widget.setPlainText(value)
                else:
                    widget.setText(value)
    
    def get_format_rules(self) -> Dict[str, Any]:
        template = self.format_template_combo.currentData()
        if template:
            return template.rules
        return config.DEFAULT_TEMPLATE_RULES.copy()
    
    def validate_fields(self) -> tuple:
        values = self.get_field_values()
        errors = []
        
        for field in get_all_fields():
            if field.required:
                value = values.get(field.field_id, '').strip()
                if not value:
                    errors.append(f"{field.field_name} 为必填项")
        
        return len(errors) == 0, errors
    
    def preview_document(self):
        valid, errors = self.validate_fields()
        if not valid:
            QMessageBox.warning(self, '验证失败', '\n'.join(errors))
            return
        
        values = self.get_field_values()
        rules = self.get_format_rules()
        
        preview_text = f"""【论文预览】

标题: {values.get('title', '')}
英文标题: {values.get('title_en', '')}

作者: {values.get('author', '')}
单位: {values.get('affiliation', '')}

摘要:
{values.get('abstract', '')}

关键词: {values.get('keywords', '')}

引言:
{values.get('introduction', '')[:500]}{'...' if len(values.get('introduction', '')) > 500 else ''}

...

结论:
{values.get('conclusion', '')[:300]}{'...' if len(values.get('conclusion', '')) > 300 else ''}

参考文献:
{values.get('references', '')[:500]}{'...' if len(values.get('references', '')) > 500 else ''}
"""
        
        preview_dialog = QDialog(self)
        preview_dialog.setWindowTitle('文档预览')
        preview_dialog.setMinimumSize(600, 500)
        
        layout = QVBoxLayout(preview_dialog)
        
        preview_text_edit = QTextEdit()
        preview_text_edit.setReadOnly(True)
        preview_text_edit.setPlainText(preview_text)
        layout.addWidget(preview_text_edit)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Ok)
        buttons.accepted.connect(preview_dialog.accept)
        layout.addWidget(buttons)
        
        preview_dialog.exec_()
    
    def generate_document(self):
        valid, errors = self.validate_fields()
        if not valid:
            QMessageBox.warning(self, '验证失败', '\n'.join(errors))
            return
        
        values = self.get_field_values()
        rules = self.get_format_rules()
        doc_type = self.doc_type_combo.currentText()
        
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            '保存文档',
            f"{values.get('title', 'document')[:30]}.docx",
            'Word文档 (*.docx)'
        )
        
        if not save_path:
            return
        
        try:
            output_path = self.doc_generator.generate_document(
                field_values=values,
                format_rules=rules,
                doc_type=doc_type,
                output_path=save_path
            )
            
            QMessageBox.information(
                self, '成功',
                f'文档已生成!\n保存位置: {output_path}\n\n是否打开文件?',
                QMessageBox.Yes | QMessageBox.No
            )
            
            if QMessageBox.question(self, '打开文件', '是否打开生成的文档?', 
                                   QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
                os.startfile(output_path)
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'生成文档失败: {str(e)}')
    
    def save_draft(self):
        values = self.get_field_values()
        
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            '保存草稿',
            f"draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            'JSON文件 (*.json)'
        )
        
        if save_path:
            with open(save_path, 'w', encoding='utf-8') as f:
                json.dump(values, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, '成功', f'草稿已保存到: {save_path}')
    
    def load_draft(self):
        load_path, _ = QFileDialog.getOpenFileName(
            self,
            '加载草稿',
            '',
            'JSON文件 (*.json)'
        )
        
        if load_path:
            try:
                with open(load_path, 'r', encoding='utf-8') as f:
                    values = json.load(f)
                self.set_field_values(values)
                QMessageBox.information(self, '成功', '草稿已加载')
            except Exception as e:
                QMessageBox.warning(self, '错误', f'加载草稿失败: {str(e)}')
    
    def clear_fields(self):
        reply = QMessageBox.question(
            self, '确认清空',
            '确定要清空所有填写内容吗?',
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            for widget in self.field_widgets.values():
                if isinstance(widget, QTextEdit):
                    widget.clear()
                else:
                    widget.clear()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.template_manager = TemplateManager()
        self.ai_recognizer = AIRecognizer()
        self.current_file = None
        self.parsed_document = None
        self.recognized_template = None
        self.selected_template = None
        
        self.init_ui()
        
        # 延迟检查更新（5 秒后）
        from PyQt5.QtCore import QTimer
        QTimer.singleShot(5000, self.auto_check_update)
    
    def auto_check_update(self):
        """自动检查更新（不显示对话框，只在后台检查）"""
        try:
            from auto_updater import AutoUpdater
            updater = AutoUpdater()
            update_info = updater.check_update()
            
            if update_info.get('has_update'):
                # 显示托盘通知（如果有）
                if hasattr(self, 'tray_icon') and self.tray_icon:
                    self.tray_icon.showMessage(
                        '发现新版本',
                        f'发现新版本 {update_info["latest_version"]}，点击"工具"菜单检查更新',
                        QSystemTrayIcon.Information,
                        3000
                    )
        except:
            pass  # 静默失败，不影响主程序
    
    def init_ui(self):
        self.setWindowTitle('论文排版优化工具')
        self.setMinimumSize(1200, 750)
        
        self.setStyleSheet('''
            QMainWindow {
                background: transparent;
            }
        ''')
        
        self.create_menu_bar()
        self.create_sidebar_layout()
    
    def create_sidebar_layout(self):
        central_widget = QWidget()
        central_widget.setStyleSheet('''
            QWidget {
                background: transparent;
            }
        ''')
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout(central_widget)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)
        
        sidebar = self.create_sidebar()
        main_layout.addWidget(sidebar)
        
        self.content_area = QStackedWidget()
        self.content_area.setStyleSheet('''
            QStackedWidget {
                background: transparent;
            }
        ''')
        
        self.home_page = self.create_home_page_widget()
        self.content_area.addWidget(self.home_page)
        
        main_layout.addWidget(self.content_area, 1)
    
    def create_sidebar(self):
        sidebar = QWidget()
        sidebar.setFixedWidth(260)
        sidebar.setStyleSheet('''
            QWidget {
                background: rgba(44, 62, 80, 0.65);
                border-right: 1px solid rgba(255, 255, 255, 0.15);
                color: #ecf0f1;
            }
        ''')
        
        layout = QVBoxLayout(sidebar)
        layout.setSpacing(0)
        layout.setContentsMargins(0, 0, 0, 0)
        
        logo_area = QWidget()
        logo_area.setFixedHeight(120)
        logo_area.setStyleSheet('''
            QWidget {
                background: rgba(52, 73, 94, 0.7);
                border-bottom: 1px solid rgba(255, 255, 255, 0.15);
            }
        ''')
        logo_layout = QVBoxLayout(logo_area)
        logo_layout.setContentsMargins(20, 20, 20, 20)
        
        logo_title = QLabel('📄 论文排版')
        logo_title.setStyleSheet('font-size: 20px; font-weight: bold; color: #ecf0f1;')
        logo_title.setAlignment(Qt.AlignCenter)
        logo_layout.addWidget(logo_title)
        
        logo_subtitle = QLabel('智能排版优化工具')
        logo_subtitle.setStyleSheet('font-size: 12px; color: #95a5a6;')
        logo_subtitle.setAlignment(Qt.AlignCenter)
        logo_layout.addWidget(logo_subtitle)
        
        layout.addWidget(logo_area)
        
        nav_label = QLabel('  功能导航')
        nav_label.setStyleSheet('''
            QLabel {
                padding: 15px 20px;
                font-size: 11px;
                font-weight: bold;
                color: #95a5a6;
                text-transform: uppercase;
                letter-spacing: 1px;
            }
        ''')
        layout.addWidget(nav_label)
        
        self.nav_buttons = []
        
        nav_items = [
            ('🏠', '首页', '#3498db', self.show_home),
            ('🚀', '快速生成', '#27ae60', self.open_quick_generate),
            ('🔍', '识别生成', '#3498db', self.open_recognize_page),
            ('⚙️', '模板设定', '#9b59b6', self.open_template_settings),
            ('📚', '参考文献', '#f39c12', self.open_reference_formatter),
            ('📝', '分部分排版', '#1abc9c', self.open_section_formatter),
        ]
        
        for icon, title, color, callback in nav_items:
            btn = self.create_nav_button(icon, title, color, callback)
            self.nav_buttons.append(btn)
            layout.addWidget(btn)
        
        layout.addStretch()
        
        help_label = QLabel('  帮助')
        help_label.setStyleSheet('''
            QLabel {
                padding: 15px 20px;
                font-size: 11px;
                font-weight: bold;
                color: #95a5a6;
                text-transform: uppercase;
                letter-spacing: 1px;
                border-top: 1px solid #34495e;
            }
        ''')
        layout.addWidget(help_label)
        
        help_btn = self.create_nav_button('❓', '系统诊断', '#e74c3c', self.show_system_check)
        self.nav_buttons.append(help_btn)
        layout.addWidget(help_btn)
        
        about_btn = self.create_nav_button('ℹ️', '关于', '#95a5a6', self.show_about)
        self.nav_buttons.append(about_btn)
        layout.addWidget(about_btn)
        
        return sidebar
    
    def create_nav_button(self, icon, title, color, callback):
        btn = QPushButton(f'  {icon}  {title}')
        btn.setFixedHeight(50)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setStyleSheet(f'''
            QPushButton {{
                background-color: transparent;
                border: none;
                color: #bdc3c7;
                text-align: left;
                padding: 0 20px;
                font-size: 14px;
            }}
            QPushButton:hover {{
                background-color: #34495e;
                color: #ecf0f1;
                border-left: 4px solid {color};
            }}
            QPushButton:pressed {{
                background-color: {color};
                color: white;
                border-left: 4px solid {color};
            }}
        ''')
        btn.clicked.connect(callback)
        return btn
    
    def create_home_page_widget(self):
        page = QWidget()
        page.setStyleSheet('''
            QWidget {
                background: transparent;
            }
            QLabel {
                background: transparent;
            }
        ''')
        layout = QVBoxLayout(page)
        layout.setSpacing(20)
        layout.setContentsMargins(30, 30, 30, 30)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet('''
            QScrollArea {
                border: none;
                background: transparent;
            }
        ''')
        
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setSpacing(20)
        
        welcome_frame = QFrame()
        welcome_frame.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        welcome_frame.setStyleSheet('''
            QFrame {
                background: rgba(255, 255, 255, 0.55);
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 20px;
                padding: 30px;
            }
        ''')
        welcome_layout = QVBoxLayout(welcome_frame)
        
        welcome_title = QLabel('欢迎使用论文排版优化工具！')
        welcome_title.setStyleSheet('font-size: 28px; font-weight: bold; color: #667eea; margin-bottom: 8px;')
        welcome_title.setAlignment(Qt.AlignCenter)
        welcome_layout.addWidget(welcome_title)
        
        welcome_subtitle = QLabel('选择左侧功能开始您的工作')
        welcome_subtitle.setStyleSheet('font-size: 16px; color: #555;')
        welcome_subtitle.setAlignment(Qt.AlignCenter)
        welcome_layout.addWidget(welcome_subtitle)
        
        scroll_layout.addWidget(welcome_frame)
        
        quick_access_frame = QFrame()
        quick_access_frame.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        quick_access_frame.setStyleSheet('''
            QFrame {
                background: rgba(255, 255, 255, 0.55);
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 20px;
                padding: 25px;
            }
        ''')
        quick_layout = QVBoxLayout(quick_access_frame)
        
        quick_title = QLabel('🚀 快速开始')
        quick_title.setStyleSheet('font-size: 20px; font-weight: bold; color: #667eea; margin-bottom: 15px;')
        quick_layout.addWidget(quick_title)
        
        buttons_layout = QGridLayout()
        buttons_layout.setSpacing(12)
        
        quick_items = [
            ('🚀', '快速生成', '#27ae60', self.open_quick_generate),
            ('🔍', '识别生成', '#3498db', self.open_recognize_page),
            ('⚙️', '模板设定', '#9b59b6', self.open_template_settings),
            ('📚', '参考文献', '#f39c12', self.open_reference_formatter),
            ('📝', '分部分排版', '#1abc9c', self.open_section_formatter),
        ]
        
        for i, (icon, title, color, callback) in enumerate(quick_items):
            btn = QPushButton(f'{icon}  {title}')
            btn.setMinimumHeight(60)
            btn.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
            btn.setStyleSheet(f'''
                QPushButton {{
                    background: qlineargradient(
                        x1: 0, y1: 0, x2: 1, y2: 0,
                        stop: 0 {color},
                        stop: 1 {color}dd
                    );
                    color: white;
                    border: none;
                    border-radius: 12px;
                    font-size: 15px;
                    font-weight: 600;
                    padding: 12px;
                }}
                QPushButton:hover {{
                    background: qlineargradient(
                        x1: 0, y1: 0, x2: 1, y2: 0,
                        stop: 0 {color}ee,
                        stop: 1 {color}
                    );
                }}
                QPushButton:pressed {{
                    background: qlineargradient(
                        x1: 0, y1: 0, x2: 1, y2: 0,
                        stop: 0 {color}cc,
                        stop: 1 {color}ee
                    );
                }}
            ''')
            btn.clicked.connect(callback)
            row = i // 2
            col = i % 2
            buttons_layout.addWidget(btn, row, col)
        
        quick_layout.addLayout(buttons_layout)
        scroll_layout.addWidget(quick_access_frame)
        
        bottom_layout = QHBoxLayout()
        bottom_layout.setSpacing(15)
        
        features_frame = QFrame()
        features_frame.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        features_frame.setStyleSheet('''
            QFrame {
                background: rgba(255, 255, 255, 0.55);
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 20px;
                padding: 25px;
            }
        ''')
        features_layout = QVBoxLayout(features_frame)
        
        features_title = QLabel('✨ 功能特点')
        features_title.setStyleSheet('font-size: 18px; font-weight: bold; color: #667eea; margin-bottom: 12px;')
        features_layout.addWidget(features_title)
        
        features = [
            ('📄', '支持 Word (.docx) 文档'),
            ('📋', '内置多种期刊和学位论文模板'),
            ('🤖', '支持 AI 自动识别文档模板'),
            ('⚙️', '可自定义模板规则'),
            ('✨', '一键优化文档排版格式'),
        ]
        
        for icon, desc in features:
            feature_label = QLabel(f'{icon}  {desc}')
            feature_label.setStyleSheet('font-size: 13px; color: #555; padding: 6px 0;')
            feature_label.setWordWrap(True)
            features_layout.addWidget(feature_label)
        
        features_layout.addStretch()
        bottom_layout.addWidget(features_frame, 1)
        
        help_frame = QFrame()
        help_frame.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        help_frame.setStyleSheet('''
            QFrame {
                background: rgba(255, 255, 255, 0.55);
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 20px;
                padding: 25px;
            }
        ''')
        help_layout = QVBoxLayout(help_frame)
        
        help_title = QLabel('💡 使用提示')
        help_title.setStyleSheet('font-size: 18px; font-weight: bold; color: #667eea; margin-bottom: 12px;')
        help_layout.addWidget(help_title)
        
        help_text = QLabel(
            '• 快速生成：从零开始撰写论文\n'
            '• 识别生成：调整已有论文格式\n'
            '• 模板设定：管理论文格式模板\n'
            '• 参考文献：一键转换格式\n'
            '• 分部分排版：单独设置章节排版'
        )
        help_text.setStyleSheet('font-size: 13px; color: #555; line-height: 1.8;')
        help_text.setWordWrap(True)
        help_layout.addWidget(help_text)
        
        help_layout.addStretch()
        bottom_layout.addWidget(help_frame, 1)
        
        scroll_layout.addLayout(bottom_layout)
        scroll_layout.addStretch()
        
        scroll.setWidget(scroll_content)
        layout.addWidget(scroll)
        
        return page
    
    def show_home(self):
        self.content_area.setCurrentWidget(self.home_page)
    
    def get_global_stylesheet(self):
        return ""
    
    def create_menu_bar(self):
        menubar = self.menuBar()
        menubar.setStyleSheet('''
            QMenuBar {
                background-color: transparent;
                border: none;
                color: #333;
            }
            QMenuBar::item {
                background-color: transparent;
                padding: 6px 12px;
                border-radius: 4px;
            }
            QMenuBar::item:selected {
                background-color: rgba(102, 126, 234, 0.2);
            }
            QMenuBar::item:pressed {
                background-color: rgba(102, 126, 234, 0.3);
            }
            QMenu {
                background-color: rgba(255, 255, 255, 0.9);
                border: 1px solid rgba(255, 255, 255, 0.5);
                border-radius: 8px;
                padding: 8px;
            }
            QMenu::item {
                padding: 6px 24px;
                border-radius: 4px;
            }
            QMenu::item:selected {
                background-color: rgba(102, 126, 234, 0.2);
            }
        ''')
        
        home_action = QAction('主页', self)
        home_action.setShortcut('Home')
        home_action.triggered.connect(self.show_home)
        menubar.addAction(home_action)
        
        file_menu = menubar.addMenu('文件(&F)')
        
        exit_action = QAction('退出', self)
        exit_action.setShortcut('Ctrl+Q')
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        view_menu = menubar.addMenu('视图(&V)')
        
        self.always_on_top_action = QAction('窗口置顶', self, checkable=True)
        self.always_on_top_action.setShortcut('Ctrl+T')
        self.always_on_top_action.triggered.connect(self.toggle_always_on_top)
        view_menu.addAction(self.always_on_top_action)
        
        settings_menu = menubar.addMenu('设置(&S)')
        
        settings_action = QAction('偏好设置', self)
        settings_action.triggered.connect(self.show_settings)
        settings_menu.addAction(settings_action)
        
        help_menu = menubar.addMenu('帮助(&H)')
        # 工具菜单
        tools_menu = menubar.addMenu('工具 (&T)')
        
        # 文件预览
        preview_action = QAction('👁️ 文件预览', self)
        preview_action.setShortcut('Ctrl+Alt+P')
        preview_action.triggered.connect(self.open_file_preview)
        tools_menu.addAction(preview_action)
        
        tools_menu.addSeparator()
        
        # 模板管理
        template_action = QAction('📋 模板管理', self)
        template_action.setShortcut('Ctrl+Alt+T')
        template_action.triggered.connect(self.open_template_manager)
        tools_menu.addAction(template_action)
        
        tools_menu.addSeparator()
        
        # 格式化配置
        format_config_action = QAction('⚙️ 格式化配置', self)
        format_config_action.setShortcut('Ctrl+Alt+F')
        format_config_action.triggered.connect(self.open_format_config)
        tools_menu.addAction(format_config_action)
        
        tools_menu.addSeparator()
        
        cover_config_action = QAction('📘 封面和声明页配置', self)
        cover_config_action.setShortcut('Ctrl+Shift+C')
        cover_config_action.triggered.connect(self.open_cover_declaration_config)
        tools_menu.addAction(cover_config_action)
        
        # 批量处理
        batch_action = QAction('📦 批量处理', self)
        batch_action.setShortcut('Ctrl+B')
        batch_action.triggered.connect(self.open_batch_processor)
        tools_menu.addAction(batch_action)
        
        tools_menu.addSeparator()
        
        # 检查更新
        update_action = QAction('🔄 检查更新', self)
        update_action.setShortcut('Ctrl+U')
        update_action.triggered.connect(self.check_update)
        tools_menu.addAction(update_action)
        
        system_check_action = QAction('🔧 系统诊断', self)
        system_check_action.triggered.connect(self.show_system_check)
        help_menu.addAction(system_check_action)
        
        about_action = QAction('关于', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def toggle_always_on_top(self, checked):
        if checked:
            self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)
        else:
            self.setWindowFlags(self.windowFlags() & ~Qt.WindowStaysOnTopHint)
        self.show()
    
    def create_status_bar(self):
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage('就绪')
    
    def show_settings(self):
        dialog = SettingsDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.ai_recognizer = AIRecognizer()
            QMessageBox.information(self, '成功', '设置已保存')
    
    def show_about(self):
        QMessageBox.about(
            self,
            '关于',
            '''论文排版优化工具 v1.0

功能特点:
• 支持 Word (.docx) 和 LaTeX (.tex) 文档
• 内置多种期刊和学位论文模板
• 支持 AI 自动识别文档模板
• 可自定义模板规则
• 一键优化文档排版格式
• 论文快速生成（填写式）

作者: 夕岸摇
'''
        )
    

    
    def open_file_preview(self):
        """打开文件预览"""
        try:
            from file_preview_dialog import FilePreviewDialog
            
            dialog = FilePreviewDialog(self)
            dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, '错误', f'打开文件预览失败：{e}')
    
    def open_template_manager(self):
        """打开模板管理器"""
        try:
            from template_editor_dialog import TemplateEditorDialog
            
            dialog = TemplateEditorDialog(self)
            dialog.exec_()
        except Exception as e:
            QMessageBox.critical(self, '错误', f'打开模板管理器失败：{e}')
    
    def open_format_config(self):
        """打开格式化配置对话框"""
        try:
            from format_config import FormatConfig
            from format_config_dialog import FormatConfigDialog
            
            config = FormatConfig()
            dialog = FormatConfigDialog(config, self)
            if dialog.exec_() == QDialog.Accepted:
                self.statusBar().showMessage('格式化配置已保存', 3000)
        except Exception as e:
            QMessageBox.critical(self, '错误', f'打开配置失败：{e}')
    
    def check_update(self):
        """检查更新"""
        try:
            from auto_updater import check_for_updates
            check_for_updates(self)
        except Exception as e:
            QMessageBox.critical(self, '错误', f'检查更新失败：{e}')
    
    def open_batch_processor(self):
        """打开批量处理对话框"""
        try:
            from batch_processor import BatchProcessor, BatchReport
            from file_backup import FileBackup
            from PyQt5.QtWidgets import QFileDialog
            
            dialog = BatchProcessorDialog(self)
            if dialog.exec_() == QDialog.Accepted:
                # 处理完成
                self.statusBar().showMessage('批量处理完成', 5000)
        except Exception as e:
            QMessageBox.critical(self, '错误', f'批量处理失败：{e}')
    
    def open_cover_declaration_config(self):
        """打开封面和声明页配置对话框"""
        try:
            from cover_declaration_config import CoverDeclarationConfigDialog
            dialog = CoverDeclarationConfigDialog(parent=self)
            if dialog.exec_() == QDialog.Accepted:
                config = dialog.get_config()
                school = config.get('cover', {}).get('school_name', '未设置')
                self.statusBar().showMessage(f'封面配置已更新：{school}')
                QMessageBox.information(self, '成功', '封面配置已保存！\n\n可以在格式化时自动应用。')
        except ImportError as e:
            QMessageBox.critical(self, '错误', f'无法加载封面配置模块：{str(e)}\n\n请确保 cover_declaration_config.py 文件存在。')
        except Exception as e:
            QMessageBox.critical(self, '错误', f'打开封面配置失败：{str(e)}')

    def open_quick_generate(self):
        if not hasattr(self, 'quick_gen_page'):
            self.quick_gen_page = QuickGeneratePage(self)
            self.content_area.addWidget(self.quick_gen_page)
        self.content_area.setCurrentWidget(self.quick_gen_page)
    
    def open_recognize_page(self):
        if not hasattr(self, 'recognize_page'):
            self.recognize_page = RecognizePage(self)
            self.content_area.addWidget(self.recognize_page)
        self.content_area.setCurrentWidget(self.recognize_page)
    
    def open_template_settings(self):
        if not hasattr(self, 'template_page'):
            self.template_page = TemplateSettingsPage(self)
            self.content_area.addWidget(self.template_page)
        self.content_area.setCurrentWidget(self.template_page)
    
    def open_reference_formatter(self):
        if not hasattr(self, 'ref_page'):
            self.ref_page = ReferenceFormatterPage(self)
            self.content_area.addWidget(self.ref_page)
        self.content_area.setCurrentWidget(self.ref_page)
    
    def open_section_formatter(self):
        if not hasattr(self, 'section_formatter_page'):
            self.section_formatter_page = SectionFormatterPage(self)
            self.content_area.addWidget(self.section_formatter_page)
        self.content_area.setCurrentWidget(self.section_formatter_page)
    
    def show_system_check(self):
        from PyQt5.QtWidgets import QDialog, QVBoxLayout, QTextEdit, QPushButton, QHBoxLayout, QMessageBox
        
        dialog = QDialog(self)
        dialog.setWindowTitle('🔧 系统诊断')
        dialog.setMinimumSize(700, 600)
        
        layout = QVBoxLayout(dialog)
        
        result_text = QTextEdit()
        result_text.setReadOnly(True)
        layout.addWidget(result_text)
        
        btn_layout = QHBoxLayout()
        
        refresh_btn = QPushButton('重新检查')
        refresh_btn.clicked.connect(lambda: self._run_system_check(result_text))
        btn_layout.addWidget(refresh_btn)
        
        copy_btn = QPushButton('复制报告')
        copy_btn.clicked.connect(lambda: self._copy_to_clipboard(result_text.toPlainText()))
        btn_layout.addWidget(copy_btn)
        
        save_btn = QPushButton('保存报告')
        save_btn.clicked.connect(lambda: self._save_diagnosis_report(result_text.toPlainText()))
        btn_layout.addWidget(save_btn)
        
        btn_layout.addStretch()
        
        close_btn = QPushButton('关闭')
        close_btn.clicked.connect(dialog.accept)
        btn_layout.addWidget(close_btn)
        
        layout.addLayout(btn_layout)
        
        self._run_system_check(result_text)
        dialog.exec_()
    
    def _run_system_check(self, text_widget):
        import sys
        import platform
        from datetime import datetime
        
        report = []
        report.append('=' * 60)
        report.append('系统诊断报告')
        report.append('=' * 60)
        report.append(f'检查时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        report.append('')
        
        report.append('【1. Python环境】')
        report.append('-' * 40)
        report.append(f'Python版本: {sys.version}')
        report.append(f'Python路径: {sys.executable}')
        report.append(f'操作系统: {platform.platform()}')
        report.append('')
        
        report.append('【2. PyQt5状态】')
        report.append('-' * 40)
        try:
            from PyQt5.QtCore import PYQT_VERSION_STR
            report.append(f'✓ PyQt5可用 - 版本: {PYQT_VERSION_STR}')
        except Exception as e:
            report.append(f'✗ PyQt5不可用: {str(e)}')
        report.append('')
        
        report.append('【3. PyQt-Fluent-Widgets状态】')
        report.append('-' * 40)
        if FLUENT_WIDGETS_AVAILABLE:
            report.append('✓ PyQt-Fluent-Widgets可用')
            try:
                from qfluentwidgets import __version__ as fw_version
                report.append(f'  版本: {fw_version}')
            except:
                report.append('  版本: 无法获取')
        else:
            report.append('✗ PyQt-Fluent-Widgets不可用')
            if 'FLUENT_IMPORT_ERROR' in globals():
                report.append(f'  错误: {FLUENT_IMPORT_ERROR}')
        report.append('')
        
        report.append('【4. 核心模块检查】')
        report.append('-' * 40)
        modules_to_check = [
            ('parsers', '文档解析模块'),
            ('template_manager', '模板管理模块'),
            ('formatter', '文档格式化模块'),
            ('document_generator', '文档生成模块'),
            ('reference_formatter', '参考文献格式化模块'),
        ]
        for module_name, description in modules_to_check:
            try:
                __import__(module_name)
                report.append(f'✓ {description} ({module_name})')
            except Exception as e:
                report.append(f'✗ {description} ({module_name}): {str(e)}')
        report.append('')
        
        report.append('【5. 依赖库检查】')
        report.append('-' * 40)
        deps_to_check = [
            ('docx', 'python-docx (Word文档处理)'),
            ('openai', 'OpenAI API (AI识别)'),
        ]
        for module_name, description in deps_to_check:
            try:
                __import__(module_name)
                report.append(f'✓ {description}')
            except ImportError:
                report.append(f'○ {description} - 可选依赖，部分功能可能受限')
            except Exception as e:
                report.append(f'✗ {description}: {str(e)}')
        report.append('')
        
        report.append('【6. 配置和目录】')
        report.append('-' * 40)
        try:
            import config
            report.append(f'✓ 配置模块可用')
            report.append(f'  输出目录: {config.OUTPUT_DIR}')
            report.append(f'  模板目录: {config.TEMPLATES_DIR}')
            os.makedirs(config.OUTPUT_DIR, exist_ok=True)
            os.makedirs(config.TEMPLATES_DIR, exist_ok=True)
            report.append(f'✓ 目录已创建或已存在')
        except Exception as e:
            report.append(f'✗ 配置检查失败: {str(e)}')
        report.append('')
        
        report.append('=' * 60)
        report.append('诊断完成')
        report.append('=' * 60)
        
        text_widget.setPlainText('\n'.join(report))
    
    def _copy_to_clipboard(self, text):
        QApplication.clipboard().setText(text)
        QMessageBox.information(self, '成功', '诊断报告已复制到剪贴板')
    
    def _save_diagnosis_report(self, report_text):
        from PyQt5.QtWidgets import QFileDialog
        from datetime import datetime
        
        default_name = f'diagnosis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.txt'
        file_path, _ = QFileDialog.getSaveFileName(
            self, '保存诊断报告', default_name, '文本文件 (*.txt)'
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(report_text)
                QMessageBox.information(self, '成功', f'诊断报告已保存到:\n{file_path}')
            except Exception as e:
                QMessageBox.critical(self, '错误', f'保存失败: {str(e)}')


class GlassEffectPage(QWidget):
    def __init__(self, parent):
        super().__init__(parent)
        self.main_window = parent
    
    def create_glass_container(self):
        container = QWidget()
        container.setStyleSheet('''
            QWidget {
                background: rgba(255, 255, 255, 0.55);
                border-radius: 20px;
                border: 1px solid rgba(255, 255, 255, 0.3);
            }
        ''')
        return container
    
    def init_page_layout(self, title_text, title_color, splitter_sizes):
        self.setStyleSheet(self.get_page_style())
        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        layout.setSpacing(20)
        
        header_layout = QHBoxLayout()
        
        back_btn = QPushButton('← 返回主页')
        back_btn.clicked.connect(self.main_window.show_home)
        header_layout.addWidget(back_btn)
        
        header_layout.addStretch()
        
        title_label = QLabel(title_text)
        title_label.setStyleSheet(f'font-size: 24px; font-weight: bold; color: {title_color};')
        header_layout.addWidget(title_label)
        
        header_layout.addStretch()
        
        layout.addLayout(header_layout)
        
        content_container = self.create_glass_container()
        content_layout = QVBoxLayout(content_container)
        content_layout.setContentsMargins(25, 25, 25, 25)
        content_layout.setSpacing(20)
        
        content_splitter = QSplitter(Qt.Horizontal)
        
        left_panel = self.create_left_panel()
        content_splitter.addWidget(left_panel)
        
        right_panel = self.create_right_panel()
        content_splitter.addWidget(right_panel)
        
        content_splitter.setSizes(splitter_sizes)
        
        content_layout.addWidget(content_splitter)
        layout.addWidget(content_container)
        layout.addStretch()
    
    def get_page_style(self):
        return '''
            QWidget {
                background: transparent;
            }
            QGroupBox {
                background: rgba(255, 255, 255, 0.65);
                border: 1px solid rgba(255, 255, 255, 0.35);
                border-radius: 16px;
                margin-top: 12px;
                font-weight: bold;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 5px;
            }
            QPushButton {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 #667eea,
                    stop: 1 #764ba2
                );
                color: white;
                border: none;
                border-radius: 12px;
                padding: 12px 24px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 #764ba2,
                    stop: 1 #667eea
                );
            }
            QPushButton:pressed {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 0, y2: 1,
                    stop: 0 #5a6fd8,
                    stop: 1 #6a4190
                );
            }
            QLineEdit, QTextEdit, QComboBox {
                background: rgba(255, 255, 255, 0.7);
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 12px;
                padding: 10px;
            }
            QLineEdit:focus, QTextEdit:focus, QComboBox:focus {
                border: 2px solid #667eea;
                background: rgba(255, 255, 255, 0.8);
            }
            QListWidget {
                background: rgba(255, 255, 255, 0.7);
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 12px;
            }
            QTabWidget::pane {
                border: 1px solid rgba(255, 255, 255, 0.3);
                border-radius: 16px;
                background: rgba(255, 255, 255, 0.65);
            }
            QTabBar::tab {
                background: rgba(255, 255, 255, 0.2);
                border: 1px solid rgba(255, 255, 255, 0.25);
                border-bottom: none;
                border-top-left-radius: 12px;
                border-top-right-radius: 12px;
                padding: 10px 24px;
                margin-right: 4px;
                color: #333;
            }
            QTabBar::tab:selected {
                background: rgba(255, 255, 255, 0.7);
                border-bottom: 1px solid rgba(255, 255, 255, 0.7);
                color: #667eea;
                font-weight: bold;
            }
            QTabBar::tab:hover:!selected {
                background: rgba(255, 255, 255, 0.4);
            }
        '''


class QuickGeneratePage(GlassEffectPage):
    def __init__(self, parent):
        super().__init__(parent)
        self.main_window = parent
        self.doc_generator = DocumentGenerator()
        self.field_widgets = {}
        self.sections = []
        self.section_widgets = []
        self.attachments = []
        self.init_ui()
    
    def init_ui(self):
        self.init_page_layout('快速生成论文', '#27ae60', [700, 300])
    
    def create_left_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        basic_group = QGroupBox('基本信息')
        basic_layout = QFormLayout(basic_group)
        
        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText('请输入论文标题')
        basic_layout.addRow('论文标题*:', self.title_edit)
        
        self.title_en_edit = QLineEdit()
        self.title_en_edit.setPlaceholderText('English Title')
        basic_layout.addRow('英文标题:', self.title_en_edit)
        
        self.author_edit = QLineEdit()
        self.author_edit.setPlaceholderText('作者姓名')
        basic_layout.addRow('作者*:', self.author_edit)
        
        self.affiliation_edit = QLineEdit()
        self.affiliation_edit.setPlaceholderText('作者单位')
        basic_layout.addRow('单位*:', self.affiliation_edit)
        
        self.email_edit = QLineEdit()
        self.email_edit.setPlaceholderText('email@example.com')
        basic_layout.addRow('邮箱:', self.email_edit)
        
        layout.addWidget(basic_group)
        
        content_tabs = QTabWidget()
        
        abstract_tab = self.create_abstract_tab()
        content_tabs.addTab(abstract_tab, '摘要关键词')
        
        body_tab = self.create_body_tab()
        content_tabs.addTab(body_tab, '正文内容')
        
        reference_tab = self.create_reference_tab()
        content_tabs.addTab(reference_tab, '参考文献')
        
        acknowledgement_tab = self.create_acknowledgement_tab()
        content_tabs.addTab(acknowledgement_tab, '致谢')
        
        attachment_tab = self.create_attachment_tab()
        content_tabs.addTab(attachment_tab, '附件')
        
        layout.addWidget(content_tabs)
        
        return panel
    
    def create_abstract_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        form_layout = QFormLayout()
        
        self.abstract_edit = QTextEdit()
        self.abstract_edit.setPlaceholderText('请输入中文摘要（200-300字）')
        form_layout.addRow('中文摘要*:', self.abstract_edit)
        
        self.abstract_en_edit = QTextEdit()
        self.abstract_en_edit.setPlaceholderText('English Abstract')
        form_layout.addRow('英文摘要:', self.abstract_en_edit)
        
        self.keywords_edit = QLineEdit()
        self.keywords_edit.setPlaceholderText('关键词1; 关键词2; 关键词3')
        form_layout.addRow('中文关键词*:', self.keywords_edit)
        
        self.keywords_en_edit = QLineEdit()
        self.keywords_en_edit.setPlaceholderText('keyword1; keyword2; keyword3')
        form_layout.addRow('英文关键词:', self.keywords_en_edit)
        
        layout.addLayout(form_layout)
        layout.addStretch()
        
        return tab
    
    def create_body_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        body_tabs = QTabWidget()
        
        structure_tab = self._create_structure_tab()
        body_tabs.addTab(structure_tab, '结构管理')
        
        section_tab = self._create_section_edit_tab()
        body_tabs.addTab(section_tab, '章节编辑')
        
        paragraph_tab = self._create_paragraph_tab()
        body_tabs.addTab(paragraph_tab, '段落排版')
        
        layout.addWidget(body_tabs)
        
        return tab
    
    def _create_structure_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc_label = QLabel('管理论文整体结构，拖拽调整顺序')
        desc_label.setStyleSheet('color: #666; font-size: 11px;')
        layout.addWidget(desc_label)
        
        self.structure_list = QListWidget()
        self.structure_list.setDragDropMode(QListWidget.InternalMove)
        layout.addWidget(self.structure_list)
        
        btn_layout = QHBoxLayout()
        
        add_part_btn = QPushButton('+ 添加部分')
        add_part_btn.setStyleSheet('background-color: #4CAF50; color: white;')
        add_part_btn.clicked.connect(self.add_structure_part)
        btn_layout.addWidget(add_part_btn)
        
        add_chapter_btn = QPushButton('+ 添加章节')
        add_chapter_btn.setStyleSheet('background-color: #2196F3; color: white;')
        add_chapter_btn.clicked.connect(self.add_structure_chapter)
        btn_layout.addWidget(add_chapter_btn)
        
        edit_btn = QPushButton('编辑选中')
        edit_btn.clicked.connect(self.edit_structure_item)
        btn_layout.addWidget(edit_btn)
        
        remove_btn = QPushButton('删除选中')
        remove_btn.clicked.connect(self.remove_structure_item)
        btn_layout.addWidget(remove_btn)
        
        layout.addLayout(btn_layout)
        
        self._init_default_structure()
        
        return tab
    
    def _init_default_structure(self):
        default_structure = [
            ('部分', '绪论', '包含研究背景、研究目的、研究意义等'),
            ('部分', '文献综述', '国内外研究现状综述'),
            ('部分', '研究方法', '研究设计与方法'),
            ('部分', '研究结果', '数据分析与结果'),
            ('部分', '结论与展望', '研究结论与未来展望'),
        ]
        
        for item_type, title, desc in default_structure:
            item = QListWidgetItem(f'📁 {title}')
            item.setData(Qt.UserRole, {'type': item_type, 'title': title, 'description': desc})
            self.structure_list.addItem(item)
    
    def add_structure_part(self):
        dialog = StructureItemDialog(self, '部分', '新建部分')
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            item = QListWidgetItem(f'📁 {data["title"]}')
            item.setData(Qt.UserRole, data)
            self.structure_list.addItem(item)
    
    def add_structure_chapter(self):
        dialog = StructureItemDialog(self, '章节', '新建章节')
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            item = QListWidgetItem(f'  📄 {data["title"]}')
            item.setData(Qt.UserRole, data)
            self.structure_list.addItem(item)
    
    def edit_structure_item(self):
        current = self.structure_list.currentItem()
        if not current:
            QMessageBox.warning(self, '提示', '请先选择要编辑的项目')
            return
        
        data = current.data(Qt.UserRole)
        dialog = StructureItemDialog(self, data['type'], '编辑', data)
        if dialog.exec_() == QDialog.Accepted:
            new_data = dialog.get_data()
            prefix = '📁 ' if new_data['type'] == '部分' else '  📄 '
            current.setText(f'{prefix}{new_data["title"]}')
            current.setData(Qt.UserRole, new_data)
    
    def remove_structure_item(self):
        current = self.structure_list.currentItem()
        if current:
            self.structure_list.takeItem(self.structure_list.row(current))
    
    def _create_section_edit_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        self.sections_scroll = QScrollArea()
        self.sections_scroll.setWidgetResizable(True)
        self.sections_container = QWidget()
        self.sections_container_layout = QVBoxLayout(self.sections_container)
        self.sections_container_layout.setSpacing(10)
        self.sections_scroll.setWidget(self.sections_container)
        layout.addWidget(self.sections_scroll)
        
        add_section_layout = QHBoxLayout()
        
        add_section_layout.addWidget(QLabel('添加章节:'))
        
        self.level_combo = QComboBox()
        self.level_combo.addItems([
            '一级标题 (章)', 
            '二级标题 (节)', 
            '三级标题', 
            '四级标题', 
            '五级标题',
            '六级标题',
            '七级标题',
            '八级标题',
            '九级标题',
            '十级标题'
        ])
        add_section_layout.addWidget(self.level_combo)
        
        add_btn = QPushButton('+ 添加')
        add_btn.setStyleSheet('background-color: #27ae60; color: white;')
        add_btn.clicked.connect(self.add_section)
        add_section_layout.addWidget(add_btn)
        
        clear_sections_btn = QPushButton('清空全部')
        clear_sections_btn.clicked.connect(self.clear_sections)
        add_section_layout.addWidget(clear_sections_btn)
        
        layout.addLayout(add_section_layout)
        
        return tab
    
    def _create_paragraph_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        desc_label = QLabel('设置段落排版格式，应用到所有段落')
        desc_label.setStyleSheet('color: #666; font-size: 11px;')
        layout.addWidget(desc_label)
        
        format_group = QGroupBox('段落格式设置')
        format_layout = QFormLayout(format_group)
        
        self.para_indent_spin = QSpinBox()
        self.para_indent_spin.setRange(0, 10)
        self.para_indent_spin.setValue(2)
        self.para_indent_spin.setSuffix(' 字符')
        format_layout.addRow('首行缩进:', self.para_indent_spin)
        
        self.para_line_spacing_spin = QDoubleSpinBox()
        self.para_line_spacing_spin.setRange(0.5, 3.0)
        self.para_line_spacing_spin.setValue(1.5)
        self.para_line_spacing_spin.setSingleStep(0.1)
        self.para_line_spacing_spin.setSuffix(' 倍')
        format_layout.addRow('行距:', self.para_line_spacing_spin)
        
        self.para_space_before_spin = QSpinBox()
        self.para_space_before_spin.setRange(0, 50)
        self.para_space_before_spin.setValue(0)
        self.para_space_before_spin.setSuffix(' 磅')
        format_layout.addRow('段前间距:', self.para_space_before_spin)
        
        self.para_space_after_spin = QSpinBox()
        self.para_space_after_spin.setRange(0, 50)
        self.para_space_after_spin.setValue(0)
        self.para_space_after_spin.setSuffix(' 磅')
        format_layout.addRow('段后间距:', self.para_space_after_spin)
        
        self.para_alignment_combo = QComboBox()
        self.para_alignment_combo.addItems(['两端对齐', '左对齐', '居中', '右对齐', '分散对齐'])
        format_layout.addRow('对齐方式:', self.para_alignment_combo)
        
        layout.addWidget(format_group)
        
        special_group = QGroupBox('特殊段落设置')
        special_layout = QFormLayout(special_group)
        
        self.first_para_indent_check = QCheckBox('首段不缩进')
        special_layout.addRow(self.first_para_indent_check)
        
        self.last_para_space_check = QCheckBox('最后一段后增加间距')
        special_layout.addRow(self.last_para_space_check)
        
        layout.addWidget(special_group)
        
        preview_group = QGroupBox('预览效果')
        preview_layout = QVBoxLayout(preview_group)
        
        self.para_preview_text = QTextEdit()
        self.para_preview_text.setReadOnly(True)
        self.para_preview_text.setMaximumHeight(120)
        self.para_preview_text.setPlainText(
            '这是段落排版预览效果。\n\n'
            '首行缩进2字符，行距1.5倍，两端对齐。'
        )
        preview_layout.addWidget(self.para_preview_text)
        
        update_preview_btn = QPushButton('更新预览')
        update_preview_btn.clicked.connect(self.update_paragraph_preview)
        preview_layout.addWidget(update_preview_btn)
        
        layout.addWidget(preview_group)
        
        layout.addStretch()
        
        return tab
    
    def update_paragraph_preview(self):
        indent = self.para_indent_spin.value()
        line_spacing = self.para_line_spacing_spin.value()
        alignment = self.para_alignment_combo.currentText()
        
        preview_text = f'''段落排版设置:
- 首行缩进: {indent} 字符
- 行距: {line_spacing} 倍
- 对齐方式: {alignment}

预览效果:
    这是设置了首行缩进{indent}字符的段落文本。行距为{line_spacing}倍，文本将按照{alignment}的方式排列。'''
        
        self.para_preview_text.setPlainText(preview_text)
    
    def create_reference_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        help_label = QLabel('每行一条参考文献，格式示例：')
        help_label.setStyleSheet('color: gray; font-size: 11px;')
        layout.addWidget(help_label)
        
        example_label = QLabel('[1] 作者. 论文标题[J]. 期刊名, 年份, 卷(期): 页码.')
        example_label.setStyleSheet('color: #666; font-size: 11px; margin-bottom: 10px;')
        layout.addWidget(example_label)
        
        self.references_edit = QTextEdit()
        self.references_edit.setPlaceholderText(
            '[1] 张三, 李四. 论文标题[J]. 期刊名称, 2023, 10(2): 1-10.\n'
            '[2] Wang J, Smith M. Paper Title[J]. Journal Name, 2023, 5: 100-120.'
        )
        layout.addWidget(self.references_edit)
        
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel('参考文献格式:'))
        self.ref_format_combo = QComboBox()
        self.ref_format_combo.addItems(['GB/T 7714', 'IEEE', 'APA', 'MLA', 'Springer'])
        format_layout.addWidget(self.ref_format_combo)
        format_layout.addStretch()
        layout.addLayout(format_layout)
        
        return tab
    
    def create_acknowledgement_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        self.acknowledgement_edit = QTextEdit()
        self.acknowledgement_edit.setPlaceholderText('请输入致谢内容...\n\n例如：\n感谢xxx基金项目的资助...\n感谢导师的悉心指导...')
        layout.addWidget(self.acknowledgement_edit)
        
        return tab
    
    def create_attachment_tab(self) -> QWidget:
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        help_label = QLabel('添加附件文件（如图片、数据文件等），生成文档时可选择插入位置')
        help_label.setStyleSheet('color: gray; font-size: 11px;')
        layout.addWidget(help_label)
        
        self.attachment_list = QListWidget()
        self.attachment_list.setMaximumHeight(150)
        layout.addWidget(self.attachment_list)
        
        btn_layout = QHBoxLayout()
        
        add_attach_btn = QPushButton('添加附件')
        add_attach_btn.clicked.connect(self.add_attachment)
        btn_layout.addWidget(add_attach_btn)
        
        remove_attach_btn = QPushButton('移除选中')
        remove_attach_btn.clicked.connect(self.remove_attachment)
        btn_layout.addWidget(remove_attach_btn)
        
        clear_attach_btn = QPushButton('清空全部')
        clear_attach_btn.clicked.connect(self.clear_attachments)
        btn_layout.addWidget(clear_attach_btn)
        
        layout.addLayout(btn_layout)
        layout.addStretch()
        
        return tab
    
    def add_attachment(self):
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, '选择附件', '',
            '所有文件 (*.*);;图片文件 (*.png *.jpg *.jpeg *.bmp *.gif);;数据文件 (*.xlsx *.xls *.csv *.dat)'
        )
        
        for file_path in file_paths:
            if file_path and file_path not in self.attachments:
                self.attachments.append(file_path)
                file_name = os.path.basename(file_path)
                item = QListWidgetItem(f"📎 {file_name}")
                item.setData(Qt.UserRole, file_path)
                self.attachment_list.addItem(item)
    
    def remove_attachment(self):
        current_item = self.attachment_list.currentItem()
        if current_item:
            file_path = current_item.data(Qt.UserRole)
            if file_path in self.attachments:
                self.attachments.remove(file_path)
            self.attachment_list.takeItem(self.attachment_list.row(current_item))
    
    def clear_attachments(self):
        self.attachments.clear()
        self.attachment_list.clear()
    
    def create_right_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        format_group = QGroupBox('格式设置')
        format_layout = QFormLayout(format_group)
        
        self.doc_type_combo = QComboBox()
        self.doc_type_combo.addItems(['中文期刊', '国际期刊', '学位论文', '自定义'])
        format_layout.addRow('文档类型:', self.doc_type_combo)
        
        self.format_template_combo = QComboBox()
        self.load_format_templates()
        format_layout.addRow('格式模板:', self.format_template_combo)
        
        layout.addWidget(format_group)
        
        action_group = QGroupBox('操作')
        action_layout = QVBoxLayout(action_group)
        
        preview_btn = QPushButton('预览文档结构')
        preview_btn.clicked.connect(self.preview_document)
        action_layout.addWidget(preview_btn)
        
        generate_btn = QPushButton('生成文档')
        generate_btn.setStyleSheet('background-color: #27ae60; color: white; padding: 12px; font-weight: bold;')
        generate_btn.clicked.connect(self.generate_document)
        action_layout.addWidget(generate_btn)
        
        draft_layout = QHBoxLayout()
        save_draft_btn = QPushButton('保存草稿')
        save_draft_btn.clicked.connect(self.save_draft)
        draft_layout.addWidget(save_draft_btn)
        
        load_draft_btn = QPushButton('加载草稿')
        load_draft_btn.clicked.connect(self.load_draft)
        draft_layout.addWidget(load_draft_btn)
        action_layout.addLayout(draft_layout)
        
        layout.addWidget(action_group)
        
        self.preview_display = QTextEdit()
        self.preview_display.setReadOnly(True)
        self.preview_display.setPlaceholderText('预览将显示在这里...')
        
        preview_group = QGroupBox('结构预览')
        preview_layout = QVBoxLayout(preview_group)
        preview_layout.addWidget(self.preview_display)
        layout.addWidget(preview_group)
        
        return panel
    
    def add_section(self):
        level = self.level_combo.currentIndex() + 1
        
        section_widget = self.create_section_widget(level)
        self.sections_container_layout.addWidget(section_widget)
        self.section_widgets.append(section_widget)
        
        self.sections_scroll.verticalScrollBar().setValue(
            self.sections_scroll.verticalScrollBar().maximum()
        )
    
    def create_section_widget(self, level: int) -> QWidget:
        widget = QWidget()
        widget.setStyleSheet(f'''
            QWidget {{
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 5px;
                padding: 5px;
            }}
        ''')
        
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(10, 5, 10, 5)
        layout.setSpacing(5)
        
        header_layout = QHBoxLayout()
        
        level_label = QLabel(f'【{level}级标题】')
        level_label.setStyleSheet(f'color: {"#e74c3c" if level == 1 else "#3498db" if level == 2 else "#9b59b6" if level == 3 else "#95a5a6"}; font-weight: bold;')
        header_layout.addWidget(level_label)
        
        header_layout.addStretch()
        
        format_btn = QPushButton('排版设置')
        format_btn.setStyleSheet('background-color: #9C27B0; color: white; padding: 2px 8px;')
        format_btn.clicked.connect(lambda: self.show_section_format_dialog(widget))
        header_layout.addWidget(format_btn)
        
        move_up_btn = QPushButton('↑')
        move_up_btn.setMaximumWidth(30)
        move_up_btn.clicked.connect(lambda: self.move_section_up(widget))
        header_layout.addWidget(move_up_btn)
        
        move_down_btn = QPushButton('↓')
        move_down_btn.setMaximumWidth(30)
        move_down_btn.clicked.connect(lambda: self.move_section_down(widget))
        header_layout.addWidget(move_down_btn)
        
        delete_btn = QPushButton('×')
        delete_btn.setMaximumWidth(30)
        delete_btn.setStyleSheet('color: #e74c3c;')
        delete_btn.clicked.connect(lambda: self.remove_section(widget))
        header_layout.addWidget(delete_btn)
        
        layout.addLayout(header_layout)
        
        title_layout = QHBoxLayout()
        title_layout.addWidget(QLabel('标题:'))
        
        title_edit = QLineEdit()
        title_edit.setPlaceholderText('输入章节标题')
        title_layout.addWidget(title_edit)
        
        layout.addLayout(title_layout)
        
        content_label = QLabel('内容:')
        layout.addWidget(content_label)
        
        content_edit = QTextEdit()
        content_edit.setPlaceholderText('输入章节内容...')
        content_edit.setMaximumHeight(100)
        layout.addWidget(content_edit)
        
        format_preview = QLabel('排版: 默认')
        format_preview.setStyleSheet('color: #666; font-size: 11px;')
        layout.addWidget(format_preview)
        
        widget.level = level
        widget.title_edit = title_edit
        widget.content_edit = content_edit
        widget.format_preview = format_preview
        widget.format_settings = {}
        
        return widget
    
    def show_section_format_dialog(self, widget):
        dialog = SectionFormatDialog(self, widget.format_settings)
        if dialog.exec_() == QDialog.Accepted:
            widget.format_settings = dialog.get_settings()
            preview_text = self._format_settings_preview(widget.format_settings)
            widget.format_preview.setText(f'排版: {preview_text}')
    
    def _format_settings_preview(self, settings: dict) -> str:
        if not settings:
            return '默认'
        parts = []
        if settings.get('first_line_indent'):
            parts.append(f'缩进{settings["first_line_indent"]}字')
        if settings.get('line_spacing'):
            parts.append(f'行距{settings["line_spacing"]}倍')
        if settings.get('alignment'):
            parts.append(settings['alignment'])
        return ', '.join(parts) if parts else '自定义'
    
    def remove_section(self, widget):
        if widget in self.section_widgets:
            self.section_widgets.remove(widget)
            self.sections_container_layout.removeWidget(widget)
            widget.deleteLater()
    
    def move_section_up(self, widget):
        index = self.section_widgets.index(widget)
        if index > 0:
            self.section_widgets.remove(widget)
            self.section_widgets.insert(index - 1, widget)
            self.refresh_section_layout()
    
    def move_section_down(self, widget):
        index = self.section_widgets.index(widget)
        if index < len(self.section_widgets) - 1:
            self.section_widgets.remove(widget)
            self.section_widgets.insert(index + 1, widget)
            self.refresh_section_layout()
    
    def refresh_section_layout(self):
        for i in reversed(range(self.sections_container_layout.count())):
            self.sections_container_layout.itemAt(i).widget().setParent(None)
        
        for widget in self.section_widgets:
            self.sections_container_layout.addWidget(widget)
    
    def clear_sections(self):
        for widget in self.section_widgets[:]:
            self.sections_container_layout.removeWidget(widget)
            widget.deleteLater()
        self.section_widgets.clear()
    
    def load_format_templates(self):
        self.format_template_combo.clear()
        self.format_template_combo.addItem('默认格式', None)
        
        templates = self.main_window.template_manager.get_all_templates()
        for template in templates:
            self.format_template_combo.addItem(template.template_name, template)
    
    def get_basic_values(self) -> Dict[str, str]:
        return {
            'title': self.title_edit.text(),
            'title_en': self.title_en_edit.text(),
            'author': self.author_edit.text(),
            'affiliation': self.affiliation_edit.text(),
            'email': self.email_edit.text(),
            'abstract': self.abstract_edit.toPlainText(),
            'abstract_en': self.abstract_en_edit.toPlainText(),
            'keywords': self.keywords_edit.text(),
            'keywords_en': self.keywords_en_edit.text(),
            'references': self.references_edit.toPlainText(),
            'acknowledgement': self.acknowledgement_edit.toPlainText(),
            'ref_format': self.ref_format_combo.currentText(),
        }
    
    def get_paragraph_settings(self) -> Dict[str, Any]:
        return {
            'first_line_indent': self.para_indent_spin.value(),
            'line_spacing': self.para_line_spacing_spin.value(),
            'space_before': self.para_space_before_spin.value(),
            'space_after': self.para_space_after_spin.value(),
            'alignment': self.para_alignment_combo.currentText(),
            'first_para_no_indent': self.first_para_indent_check.isChecked(),
            'last_para_extra_space': self.last_para_space_check.isChecked(),
        }
    
    def get_structure(self) -> List[Dict[str, Any]]:
        structure = []
        for i in range(self.structure_list.count()):
            item = self.structure_list.item(i)
            data = item.data(Qt.UserRole)
            structure.append(data)
        return structure
    
    def get_sections(self) -> List[Dict[str, Any]]:
        sections = []
        for widget in self.section_widgets:
            sections.append({
                'level': widget.level,
                'title': widget.title_edit.text(),
                'content': widget.content_edit.toPlainText(),
                'format_settings': widget.format_settings
            })
        return sections
    
    def get_format_rules(self) -> Dict[str, Any]:
        template = self.format_template_combo.currentData()
        if template:
            return template.rules
        return config.DEFAULT_TEMPLATE_RULES.copy()
    
    def validate_fields(self) -> tuple:
        errors = []
        
        if not self.title_edit.text().strip():
            errors.append('论文标题为必填项')
        
        return len(errors) == 0, errors
    
    def preview_document(self):
        valid, errors = self.validate_fields()
        if not valid:
            QMessageBox.warning(self, '验证失败', '\n'.join(errors))
            return
        
        basic = self.get_basic_values()
        sections = self.get_sections()
        
        preview_lines = [
            f"【论文标题】{basic['title']}",
            f"【作者】{basic['author']}",
            f"【单位】{basic['affiliation']}",
            "",
            f"【摘要】{basic['abstract'][:100]}...",
            f"【关键词】{basic['keywords']}",
            "",
            "【正文结构】"
        ]
        
        for section in sections:
            indent = "  " * (section['level'] - 1)
            preview_lines.append(f"{indent}{'└─' if section['level'] > 1 else ''}{section['title']}")
            if section['content']:
                preview_lines.append(f"{indent}    {section['content'][:50]}...")
        
        if basic.get('references'):
            ref_count = len(basic['references'].strip().split('\n'))
            preview_lines.append(f"\n【参考文献】{ref_count} 条")
        
        if basic.get('acknowledgement'):
            preview_lines.append(f"\n【致谢】{basic['acknowledgement'][:50]}...")
        
        if self.attachments:
            preview_lines.append(f"\n【附件】{len(self.attachments)} 个文件")
        
        self.preview_display.setPlainText('\n'.join(preview_lines))
    
    def generate_document(self):
        valid, errors = self.validate_fields()
        if not valid:
            QMessageBox.warning(self, '验证失败', '\n'.join(errors))
            return
        
        basic = self.get_basic_values()
        sections = self.get_sections()
        rules = self.get_format_rules()
        doc_type = self.doc_type_combo.currentText()
        
        if not sections:
            reply = QMessageBox.question(
                self, '提示',
                '正文内容为空，是否继续生成？',
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.No:
                return
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '保存文档',
            f"{basic['title'][:30]}.docx",
            'Word文档 (*.docx)'
        )
        
        if not save_path:
            return
        
        try:
            output_path = self.doc_generator.generate_document_with_sections(
                basic_values=basic,
                sections=sections,
                format_rules=rules,
                doc_type=doc_type,
                output_path=save_path,
                attachments=self.attachments
            )
            
            reply = QMessageBox.question(
                self, '成功',
                f'文档已生成!\n保存位置: {output_path}\n\n是否打开文件?',
                QMessageBox.Yes | QMessageBox.No
            )
            
            if reply == QMessageBox.Yes:
                os.startfile(output_path)
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'生成文档失败: {str(e)}')
    
    def save_draft(self):
        data = {
            'basic': self.get_basic_values(),
            'sections': self.get_sections(),
            'attachments': self.attachments
        }
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '保存草稿',
            f"draft_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            'JSON文件 (*.json)'
        )
        
        if save_path:
            with open(save_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            QMessageBox.information(self, '成功', '草稿已保存')
    
    def load_draft(self):
        load_path, _ = QFileDialog.getOpenFileName(
            self, '加载草稿', '', 'JSON文件 (*.json)'
        )
        
        if load_path:
            try:
                with open(load_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                basic = data.get('basic', {})
                self.title_edit.setText(basic.get('title', ''))
                self.title_en_edit.setText(basic.get('title_en', ''))
                self.author_edit.setText(basic.get('author', ''))
                self.affiliation_edit.setText(basic.get('affiliation', ''))
                self.email_edit.setText(basic.get('email', ''))
                self.abstract_edit.setPlainText(basic.get('abstract', ''))
                self.abstract_en_edit.setPlainText(basic.get('abstract_en', ''))
                self.keywords_edit.setText(basic.get('keywords', ''))
                self.keywords_en_edit.setText(basic.get('keywords_en', ''))
                self.references_edit.setPlainText(basic.get('references', ''))
                self.acknowledgement_edit.setPlainText(basic.get('acknowledgement', ''))
                
                ref_format = basic.get('ref_format', 'GB/T 7714')
                index = self.ref_format_combo.findText(ref_format)
                if index >= 0:
                    self.ref_format_combo.setCurrentIndex(index)
                
                self.clear_sections()
                
                for section in data.get('sections', []):
                    level = section.get('level', 1)
                    self.level_combo.setCurrentIndex(level - 1)
                    widget = self.create_section_widget(level)
                    widget.title_edit.setText(section.get('title', ''))
                    widget.content_edit.setPlainText(section.get('content', ''))
                    self.sections_container_layout.addWidget(widget)
                    self.section_widgets.append(widget)
                
                self.clear_attachments()
                for attach_path in data.get('attachments', []):
                    if os.path.exists(attach_path):
                        self.attachments.append(attach_path)
                        file_name = os.path.basename(attach_path)
                        item = QListWidgetItem(f"📎 {file_name}")
                        item.setData(Qt.UserRole, attach_path)
                        self.attachment_list.addItem(item)
                
                QMessageBox.information(self, '成功', '草稿已加载')
            except Exception as e:
                QMessageBox.warning(self, '错误', f'加载失败: {str(e)}')


class RecognizePage(GlassEffectPage):
    def __init__(self, parent):
        super().__init__(parent)
        self.main_window = parent
        self.ai_recognizer = AIRecognizer()
        self.current_file = None
        self.parsed_document = None
        self.recognized_template = None
        self.selected_template = None
        self.init_ui()
    
    def init_ui(self):
        self.init_page_layout('识别生成', '#3498db', [700, 300])
    
    def create_left_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        file_group = QGroupBox('文档信息')
        file_layout = QVBoxLayout(file_group)
        
        open_btn = QPushButton('打开文档')
        open_btn.setStyleSheet('background-color: #3498db; color: white; padding: 8px;')
        open_btn.clicked.connect(self.open_file)
        file_layout.addWidget(open_btn)
        
        self.file_path_label = QLabel('未选择文件')
        self.file_path_label.setStyleSheet('color: gray;')
        self.file_path_label.setWordWrap(True)
        file_layout.addWidget(self.file_path_label)
        
        self.file_info_text = QTextEdit()
        self.file_info_text.setReadOnly(True)
        self.file_info_text.setMaximumHeight(100)
        file_layout.addWidget(self.file_info_text)
        
        layout.addWidget(file_group)
        
        template_group = QGroupBox('模板选择')
        template_layout = QVBoxLayout(template_group)
        
        self.template_type_combo = QComboBox()
        self.template_type_combo.addItems(['全部', '中文期刊', '国际期刊', '学位论文', '自定义'])
        self.template_type_combo.currentTextChanged.connect(self.filter_templates)
        template_layout.addWidget(self.template_type_combo)
        
        self.template_list = QListWidget()
        self.template_list.itemClicked.connect(self.on_template_selected)
        self.refresh_template_list()
        template_layout.addWidget(self.template_list)
        
        layout.addWidget(template_group)
        
        ai_group = QGroupBox('AI识别')
        ai_layout = QVBoxLayout(ai_group)
        
        btn_layout = QHBoxLayout()
        
        auto_btn = QPushButton('自动识别')
        auto_btn.clicked.connect(self.auto_recognize)
        btn_layout.addWidget(auto_btn)
        
        import_btn = QPushButton('导入AI数据')
        import_btn.clicked.connect(self.import_ai_data)
        btn_layout.addWidget(import_btn)
        
        ai_layout.addLayout(btn_layout)
        
        self.ai_result_text = QTextEdit()
        self.ai_result_text.setReadOnly(True)
        self.ai_result_text.setMaximumHeight(120)
        self.ai_result_text.setPlaceholderText('点击按钮进行模板识别...')
        ai_layout.addWidget(self.ai_result_text)
        
        apply_btn = QPushButton('应用识别结果')
        apply_btn.clicked.connect(self.apply_recognition)
        ai_layout.addWidget(apply_btn)
        
        layout.addWidget(ai_group)
        
        return panel
    
    def create_right_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        tabs = QTabWidget()
        
        preview_tab = QWidget()
        preview_layout = QVBoxLayout(preview_tab)
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setPlaceholderText('打开文档后显示预览...')
        preview_layout.addWidget(self.preview_text)
        tabs.addTab(preview_tab, '文档预览')
        
        rules_tab = QWidget()
        rules_layout = QVBoxLayout(rules_tab)
        self.rules_table = QTableWidget()
        self.rules_table.setColumnCount(2)
        self.rules_table.setHorizontalHeaderLabels(['规则项', '当前值'])
        self.rules_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.rules_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        rules_layout.addWidget(self.rules_table)
        tabs.addTab(rules_tab, '模板规则')
        
        result_tab = QWidget()
        result_layout = QVBoxLayout(result_tab)
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setPlaceholderText('处理完成后显示结果...')
        result_layout.addWidget(self.result_text)
        
        result_btn_layout = QHBoxLayout()
        self.save_as_btn = QPushButton('另存为')
        self.save_as_btn.clicked.connect(self.save_formatted)
        self.save_as_btn.setEnabled(False)
        result_btn_layout.addWidget(self.save_as_btn)
        
        open_folder_btn = QPushButton('打开输出目录')
        open_folder_btn.clicked.connect(self.open_output_folder)
        result_btn_layout.addWidget(open_folder_btn)
        result_layout.addLayout(result_btn_layout)
        tabs.addTab(result_tab, '处理结果')
        
        layout.addWidget(tabs)
        
        format_btn = QPushButton('开始排版')
        format_btn.setStyleSheet('background-color: #27ae60; color: white; padding: 12px; font-weight: bold; font-size: 14px;')
        format_btn.clicked.connect(self.format_document)
        layout.addWidget(format_btn)
        
        return panel
    
    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '选择文档', '',
            'Word文档 (*.docx);;LaTeX文档 (*.tex);;所有文件 (*.*)'
        )
        
        if file_path:
            self.current_file = file_path
            self.file_path_label.setText(file_path)
            self.file_path_label.setStyleSheet('color: black;')
            
            try:
                self.parsed_document = DocumentParserFactory.parse(file_path)
                self.update_file_info()
                self.update_preview()
            except Exception as e:
                QMessageBox.critical(self, '错误', f'无法解析文档: {str(e)}')
    
    def update_file_info(self):
        if not self.parsed_document:
            return
        
        info = [
            f"类型: {self.parsed_document.file_type}",
            f"元素: {len(self.parsed_document.elements)}个"
        ]
        
        meta = self.parsed_document.metadata
        if meta.get('title'):
            info.append(f"标题: {meta['title']}")
        
        self.file_info_text.setPlainText('\n'.join(info))
    
    def update_preview(self):
        if not self.parsed_document:
            return
        
        lines = []
        for elem in self.parsed_document.elements[:30]:
            prefix = ""
            if elem.element_type == 'title':
                prefix = "【标题】"
            elif elem.element_type == 'heading':
                prefix = f"【标题{elem.level}】"
            elif elem.element_type == 'abstract':
                prefix = "【摘要】"
            
            content = elem.content[:150] + "..." if len(elem.content) > 150 else elem.content
            lines.append(f"{prefix}{content}")
        
        self.preview_text.setPlainText('\n\n'.join(lines))
    
    def refresh_template_list(self):
        self.template_list.clear()
        templates = self.main_window.template_manager.get_all_templates()
        
        for template in templates:
            item = QListWidgetItem(f"{template.template_name} ({template.template_type})")
            item.setData(Qt.UserRole, template.template_id)
            self.template_list.addItem(item)
    
    def filter_templates(self, filter_type: str):
        self.template_list.clear()
        
        if filter_type == '全部':
            templates = self.main_window.template_manager.get_all_templates()
        else:
            templates = self.main_window.template_manager.get_templates_by_type(filter_type)
        
        for template in templates:
            item = QListWidgetItem(f"{template.template_name} ({template.template_type})")
            item.setData(Qt.UserRole, template.template_id)
            self.template_list.addItem(item)
    
    def on_template_selected(self, item: QListWidgetItem):
        template_id = item.data(Qt.UserRole)
        self.selected_template = self.main_window.template_manager.get_template(template_id)
        self.update_rules_display()
    
    def update_rules_display(self):
        if not self.selected_template:
            return
        
        rules = self.selected_template.rules
        rule_names = {
            'font_family': '正文字体', 'font_size': '正文字号',
            'line_spacing': '行距', 'margin_top': '上边距',
            'margin_bottom': '下边距', 'margin_left': '左边距',
            'margin_right': '右边距', 'title_font': '标题字体',
            'title_size': '标题字号', 'heading1_size': '一级标题字号',
            'heading2_size': '二级标题字号', 'heading3_size': '三级标题字号',
            'reference_format': '参考文献格式'
        }
        
        self.rules_table.setRowCount(len(rules))
        
        for row, (key, value) in enumerate(rules.items()):
            name_item = QTableWidgetItem(rule_names.get(key, key))
            value_item = QTableWidgetItem(str(value))
            name_item.setFlags(name_item.flags() & ~Qt.ItemIsEditable)
            value_item.setFlags(value_item.flags() & ~Qt.ItemIsEditable)
            self.rules_table.setItem(row, 0, name_item)
            self.rules_table.setItem(row, 1, value_item)
    
    def auto_recognize(self):
        if not self.parsed_document:
            QMessageBox.warning(self, '提示', '请先打开文档')
            return
        
        self.ai_result_text.setPlainText('正在识别...')
        
        document_info = {
            'file_type': self.parsed_document.file_type,
            'metadata': self.parsed_document.metadata,
            'elements': [{'element_type': e.element_type, 'content': e.content, 'level': e.level} 
                        for e in self.parsed_document.elements]
        }
        
        result = self.ai_recognizer.recognize_template(document_info)
        self.recognized_template = result
        
        text = f"""模板: {result.template_name}
类型: {result.template_type}
置信度: {result.confidence:.0%}

字体: {result.rules.get('font_family', '-')}
字号: {result.rules.get('font_size', '-')}pt
行距: {result.rules.get('line_spacing', '-')}"""
        
        self.ai_result_text.setPlainText(text)
    
    def import_ai_data(self):
        dialog = AIDataImportDialog(self)
        if dialog.exec_() == QDialog.Accepted and dialog.parsed_data:
            if not dialog.parsed_data.is_valid:
                QMessageBox.warning(self, '错误', dialog.parsed_data.error_message)
                return
            
            template = Template(
                template_id=f"ai_imported_{datetime.now().strftime('%Y%m%d%H%M%S')}",
                template_name=dialog.parsed_data.template_name,
                template_type=dialog.parsed_data.template_type,
                description=dialog.parsed_data.description,
                rules=dialog.parsed_data.rules
            )
            
            self.selected_template = template
            self.update_rules_display()
            self.ai_result_text.setPlainText(f"已导入: {template.template_name}")
            QMessageBox.information(self, '成功', 'AI数据已导入')
    
    def apply_recognition(self):
        if not self.recognized_template:
            QMessageBox.warning(self, '提示', '没有识别结果')
            return
        
        template = Template(
            template_id=f"ai_recognized_{datetime.now().strftime('%Y%m%d%H%M%S')}",
            template_name=self.recognized_template.template_name,
            template_type=self.recognized_template.template_type,
            description=self.recognized_template.description,
            rules=self.recognized_template.rules
        )
        
        self.selected_template = template
        self.update_rules_display()
        QMessageBox.information(self, '成功', '已应用识别结果')
    
    def format_document(self):
        if not self.current_file:
            QMessageBox.warning(self, '提示', '请先打开文档')
            return
        
        if not self.selected_template:
            QMessageBox.warning(self, '提示', '请选择或识别模板')
            return
        
        output_dir = config.OUTPUT_DIR
        os.makedirs(output_dir, exist_ok=True)
        
        base_name = os.path.splitext(os.path.basename(self.current_file))[0]
        ext = os.path.splitext(self.current_file)[1]
        output_path = os.path.join(output_dir, f"{base_name}_formatted{ext}")
        
        self.result_text.setPlainText('正在处理...')
        
        try:
            result = DocumentFormatter.format(self.current_file, output_path, self.selected_template)
            
            if result.success:
                text = f"✓ 格式化完成!\n\n输出: {result.output_path}\n\n修改记录:\n"
                for change in result.changes:
                    text += f"• {change['type']}: {change['detail']}\n"
                self.result_text.setPlainText(text)
                self.save_as_btn.setEnabled(True)
            else:
                self.result_text.setPlainText(f"✗ 格式化失败\n\n错误:\n" + '\n'.join(result.errors))
        except Exception as e:
            self.result_text.setPlainText(f"处理出错: {str(e)}")
    
    def save_formatted(self):
        if not self.current_file:
            return
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '保存文档', '', 'Word文档 (*.docx)'
        )
        
        if save_path:
            output_dir = config.OUTPUT_DIR
            base_name = os.path.splitext(os.path.basename(self.current_file))[0]
            ext = os.path.splitext(self.current_file)[1]
            source_path = os.path.join(output_dir, f"{base_name}_formatted{ext}")
            
            if os.path.exists(source_path):
                import shutil
                shutil.copy2(source_path, save_path)
                QMessageBox.information(self, '成功', f'已保存到: {save_path}')
    
    def open_output_folder(self):
        os.makedirs(config.OUTPUT_DIR, exist_ok=True)
        os.startfile(config.OUTPUT_DIR)


class TemplateSettingsPage(GlassEffectPage):
    def __init__(self, parent):
        super().__init__(parent)
        self.main_window = parent
        self.init_ui()
    
    def init_ui(self):
        self.init_page_layout('模板设定', '#9b59b6', [700, 300])
    
    def create_left_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        list_group = QGroupBox('模板列表')
        list_layout = QVBoxLayout(list_group)
        
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel('筛选:'))
        self.filter_combo = QComboBox()
        self.filter_combo.addItems(['全部', '中文期刊', '国际期刊', '学位论文', '自定义'])
        self.filter_combo.currentTextChanged.connect(self.filter_templates)
        filter_layout.addWidget(self.filter_combo)
        list_layout.addLayout(filter_layout)
        
        self.template_list = QListWidget()
        self.template_list.itemClicked.connect(self.on_template_selected)
        self.refresh_template_list()
        list_layout.addWidget(self.template_list)
        
        btn_layout = QHBoxLayout()
        
        new_btn = QPushButton('新建')
        new_btn.clicked.connect(self.create_template)
        btn_layout.addWidget(new_btn)
        
        edit_btn = QPushButton('编辑')
        edit_btn.clicked.connect(self.edit_template)
        btn_layout.addWidget(edit_btn)
        
        delete_btn = QPushButton('删除')
        delete_btn.clicked.connect(self.delete_template)
        btn_layout.addWidget(delete_btn)
        
        list_layout.addLayout(btn_layout)
        
        layout.addWidget(list_group)
        
        create_group = QGroupBox('创建模板')
        create_layout = QVBoxLayout(create_group)
        
        create_btn = QPushButton('➕ 创建新模板')
        create_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 12px; font-weight: bold; font-size: 13px;')
        create_btn.clicked.connect(self.quick_parse_spec)
        create_layout.addWidget(create_btn)
        
        create_hint = QLabel('支持三种方式创建模板')
        create_hint.setStyleSheet('color: gray; font-size: 11px;')
        create_hint.setAlignment(Qt.AlignCenter)
        create_layout.addWidget(create_hint)
        
        layout.addWidget(create_group)
        
        ai_group = QGroupBox('AI数据导入')
        ai_layout = QVBoxLayout(ai_group)
        
        import_btn = QPushButton('🤖 导入AI数据')
        import_btn.setStyleSheet('background-color: #9b59b6; color: white; padding: 10px;')
        import_btn.clicked.connect(self.import_ai_data)
        ai_layout.addWidget(import_btn)
        
        ai_hint = QLabel('将豆包/千问等AI返回的\nJSON数据粘贴导入')
        ai_hint.setStyleSheet('color: gray; font-size: 11px;')
        ai_hint.setAlignment(Qt.AlignCenter)
        ai_layout.addWidget(ai_hint)
        
        layout.addWidget(ai_group)
        
        return panel
    
    def create_right_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        detail_group = QGroupBox('模板详情')
        detail_layout = QVBoxLayout(detail_group)
        
        self.detail_text = QTextEdit()
        self.detail_text.setReadOnly(True)
        self.detail_text.setPlaceholderText('选择模板查看详情...')
        detail_layout.addWidget(self.detail_text)
        
        layout.addWidget(detail_group)
        
        io_group = QGroupBox('导入/导出')
        io_layout = QHBoxLayout(io_group)
        
        export_btn = QPushButton('导出模板')
        export_btn.clicked.connect(self.export_template)
        io_layout.addWidget(export_btn)
        
        import_file_btn = QPushButton('导入模板文件')
        import_file_btn.clicked.connect(self.import_template_file)
        io_layout.addWidget(import_file_btn)
        
        layout.addWidget(io_group)
        
        return panel
    
    def refresh_template_list(self):
        self.template_list.clear()
        templates = self.main_window.template_manager.get_all_templates()
        
        for template in templates:
            item = QListWidgetItem(f"{template.template_name} ({template.template_type})")
            item.setData(Qt.UserRole, template.template_id)
            self.template_list.addItem(item)
    
    def filter_templates(self, filter_type: str):
        self.template_list.clear()
        
        if filter_type == '全部':
            templates = self.main_window.template_manager.get_all_templates()
        else:
            templates = self.main_window.template_manager.get_templates_by_type(filter_type)
        
        for template in templates:
            item = QListWidgetItem(f"{template.template_name} ({template.template_type})")
            item.setData(Qt.UserRole, template.template_id)
            self.template_list.addItem(item)
    
    def on_template_selected(self, item: QListWidgetItem):
        template_id = item.data(Qt.UserRole)
        template = self.main_window.template_manager.get_template(template_id)
        
        if template:
            text = f"""模板名称: {template.template_name}
模板类型: {template.template_type}
描述: {template.description}
创建时间: {template.created_at}

=== 格式规则 ===
"""
            rule_names = {
                'font_family': '正文字体', 'font_size': '正文字号',
                'line_spacing': '行距', 'margin_top': '上边距',
                'margin_bottom': '下边距', 'margin_left': '左边距',
                'margin_right': '右边距', 'title_font': '标题字体',
                'title_size': '标题字号', 'heading1_font': '一级标题字体',
                'heading1_size': '一级标题字号', 'heading2_font': '二级标题字体',
                'heading2_size': '二级标题字号', 'heading3_font': '三级标题字体',
                'heading3_size': '三级标题字号', 'abstract_font': '摘要字体',
                'abstract_size': '摘要字号', 'reference_format': '参考文献格式',
                'reference_font': '参考文献字体', 'reference_size': '参考文献字号'
            }
            
            for key, value in template.rules.items():
                name = rule_names.get(key, key)
                text += f"{name}: {value}\n"
            
            self.detail_text.setPlainText(text)
    
    def create_template(self):
        dialog = TemplateEditDialog(parent=self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_template_data()
            template = self.main_window.template_manager.create_custom_template(
                name=data['template_name'],
                custom_rules=data['rules']
            )
            template.template_type = data['template_type']
            template.description = data['description']
            self.main_window.template_manager.update_template(template.template_id, data['rules'])
            self.refresh_template_list()
            QMessageBox.information(self, '成功', '模板创建成功')
    
    def edit_template(self):
        current_item = self.template_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, '提示', '请先选择模板')
            return
        
        template_id = current_item.data(Qt.UserRole)
        template = self.main_window.template_manager.get_template(template_id)
        
        if not template:
            return
        
        dialog = TemplateEditDialog(template, parent=self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_template_data()
            self.main_window.template_manager.update_template(template_id, data['rules'])
            template.template_name = data['template_name']
            template.template_type = data['template_type']
            template.description = data['description']
            self.refresh_template_list()
            QMessageBox.information(self, '成功', '模板更新成功')
    
    def delete_template(self):
        current_item = self.template_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, '提示', '请先选择模板')
            return
        
        template_id = current_item.data(Qt.UserRole)
        
        if QMessageBox.question(self, '确认', '确定删除该模板?', 
                               QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
            if self.main_window.template_manager.delete_template(template_id):
                self.refresh_template_list()
                self.detail_text.clear()
                QMessageBox.information(self, '成功', '模板已删除')
            else:
                QMessageBox.warning(self, '提示', '无法删除内置模板')
    
    def import_ai_data(self):
        dialog = AIDataImportDialog(self)
        if dialog.exec_() == QDialog.Accepted and dialog.parsed_data:
            if not dialog.parsed_data.is_valid:
                QMessageBox.warning(self, '错误', dialog.parsed_data.error_message)
                return
            
            template = self.main_window.template_manager.create_custom_template(
                name=dialog.parsed_data.template_name,
                custom_rules=dialog.parsed_data.rules
            )
            template.template_type = dialog.parsed_data.template_type
            template.description = dialog.parsed_data.description
            
            self.refresh_template_list()
            QMessageBox.information(self, '成功', f'模板 "{template.template_name}" 创建成功')
    
    def quick_parse_spec(self):
        dialog = QuickSpecParseDialog(self)
        if dialog.exec_() == QDialog.Accepted and dialog.parsed_rules:
            data = dialog.get_template_data()
            
            template = self.main_window.template_manager.create_custom_template(
                name=data['template_name'],
                custom_rules=data['rules']
            )
            template.template_type = data['template_type']
            template.description = data['description']
            
            self.refresh_template_list()
            QMessageBox.information(self, '成功', f'模板 "{template.template_name}" 创建成功')
    
    def export_template(self):
        current_item = self.template_list.currentItem()
        if not current_item:
            QMessageBox.warning(self, '提示', '请先选择模板')
            return
        
        template_id = current_item.data(Qt.UserRole)
        template = self.main_window.template_manager.get_template(template_id)
        
        if not template:
            return
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '导出模板', f'{template.template_name}.json', 'JSON文件 (*.json)'
        )
        
        if save_path:
            if self.main_window.template_manager.export_template(template_id, save_path):
                QMessageBox.information(self, '成功', f'已导出到: {save_path}')
            else:
                QMessageBox.warning(self, '错误', '导出失败')
    
    def import_template_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '导入模板', '', 'JSON文件 (*.json);;YAML文件 (*.yaml *.yml)'
        )
        
        if file_path:
            template = self.main_window.template_manager.import_template(file_path)
            if template:
                self.refresh_template_list()
                QMessageBox.information(self, '成功', f'模板 "{template.template_name}" 导入成功')
            else:
                QMessageBox.warning(self, '错误', '导入失败')


class ReferenceFormatterPage(GlassEffectPage):
    def __init__(self, parent):
        super().__init__(parent)
        self.main_window = parent
        self.ref_manager = ReferenceManager()
        self.init_ui()
    
    def init_ui(self):
        self.init_page_layout('参考文献格式化', '#e67e22', [700, 300])
    
    def create_left_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        input_group = QGroupBox('输入参考文献')
        input_layout = QVBoxLayout(input_group)
        
        input_layout.addWidget(QLabel('请粘贴参考文献（每条一行）:'))
        
        self.input_text = QTextEdit()
        self.input_text.setPlaceholderText(
            '示例:\n'
            '[1] 张三, 李四. 论文标题[J]. 期刊名称, 2023, 10(2): 1-10.\n'
            '[2] Wang J, Smith M. Paper Title[J]. Journal Name, 2023, 5: 100-120.\n'
            '...'
        )
        input_layout.addWidget(self.input_text)
        
        parse_btn = QPushButton('解析参考文献')
        parse_btn.setStyleSheet('background-color: #e67e22; color: white; padding: 8px;')
        parse_btn.clicked.connect(self.parse_references)
        input_layout.addWidget(parse_btn)
        
        layout.addWidget(input_group)
        
        format_group = QGroupBox('格式设置')
        format_layout = QVBoxLayout(format_group)
        
        form_layout = QFormLayout()
        
        self.format_combo = QComboBox()
        self.format_combo.addItems([
            'GB/T 7714 (中文标准)',
            'IEEE',
            'APA',
            'MLA',
            'Springer',
        ])
        form_layout.addRow('输出格式:', self.format_combo)
        
        format_layout.addLayout(form_layout)
        
        convert_btn = QPushButton('转换格式')
        convert_btn.setStyleSheet('background-color: #27ae60; color: white; padding: 10px;')
        convert_btn.clicked.connect(self.convert_format)
        format_layout.addWidget(convert_btn)
        
        layout.addWidget(format_group)
        
        return panel
    
    def create_right_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        output_group = QGroupBox('输出结果')
        output_layout = QVBoxLayout(output_group)
        
        self.output_text = QTextEdit()
        self.output_text.setReadOnly(True)
        self.output_text.setPlaceholderText('转换后的参考文献将显示在这里...')
        output_layout.addWidget(self.output_text)
        
        btn_layout = QHBoxLayout()
        
        copy_btn = QPushButton('复制结果')
        copy_btn.clicked.connect(self.copy_result)
        btn_layout.addWidget(copy_btn)
        
        export_bib_btn = QPushButton('导出BibTeX')
        export_bib_btn.clicked.connect(self.export_bibtex)
        btn_layout.addWidget(export_bib_btn)
        
        save_btn = QPushButton('保存到文件')
        save_btn.clicked.connect(self.save_to_file)
        btn_layout.addWidget(save_btn)
        
        output_layout.addLayout(btn_layout)
        
        layout.addWidget(output_group)
        
        stats_group = QGroupBox('统计信息')
        stats_layout = QVBoxLayout(stats_group)
        
        self.stats_label = QLabel('解析后显示统计信息')
        self.stats_label.setStyleSheet('color: gray;')
        stats_layout.addWidget(self.stats_label)
        
        layout.addWidget(stats_group)
        
        return panel
    
    def parse_references(self):
        text = self.input_text.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, '提示', '请输入参考文献')
            return
        
        self.ref_manager.parse_from_text(text)
        
        count = len(self.ref_manager.references)
        self.stats_label.setText(f'已解析 {count} 条参考文献')
        
        self.output_text.setPlainText(self.ref_manager.get_formatted_output())
        
        issues = self.ref_manager.validate_references()
        if issues:
            issue_text = '\n'.join([f"参考文献 {i['number']}: {i['message']}" for i in issues[:5]])
            QMessageBox.warning(self, '解析警告', f'部分参考文献信息不完整:\n{issue_text}')
    
    def convert_format(self):
        if not self.ref_manager.references:
            QMessageBox.warning(self, '提示', '请先解析参考文献')
            return
        
        format_map = {
            'GB/T 7714 (中文标准)': ReferenceFormat.GB_T_7714,
            'IEEE': ReferenceFormat.IEEE,
            'APA': ReferenceFormat.APA,
            'MLA': ReferenceFormat.MLA,
            'Springer': ReferenceFormat.SPRINGER,
        }
        
        selected_format = self.format_combo.currentText()
        self.ref_manager.set_format(format_map.get(selected_format, ReferenceFormat.GB_T_7714))
        
        self.output_text.setPlainText(self.ref_manager.get_formatted_output())
        QMessageBox.information(self, '成功', f'已转换为 {selected_format} 格式')
    
    def copy_result(self):
        text = self.output_text.toPlainText()
        if text:
            QApplication.clipboard().setText(text)
            QMessageBox.information(self, '成功', '已复制到剪贴板')
    
    def export_bibtex(self):
        if not self.ref_manager.references:
            QMessageBox.warning(self, '提示', '请先解析参考文献')
            return
        
        bibtex = self.ref_manager.export_to_bibtex()
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '导出BibTeX', 'references.bib', 'BibTeX文件 (*.bib)'
        )
        
        if save_path:
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(bibtex)
            QMessageBox.information(self, '成功', f'已导出到: {save_path}')
    
    def save_to_file(self):
        text = self.output_text.toPlainText()
        if not text:
            QMessageBox.warning(self, '提示', '没有内容可保存')
            return
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '保存参考文献', 'references.txt', '文本文件 (*.txt)'
        )
        
        if save_path:
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(text)
            QMessageBox.information(self, '成功', f'已保存到: {save_path}')
    
    def show_settings(self):
        dialog = SettingsDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.ai_recognizer = AIRecognizer()
            QMessageBox.information(self, '成功', '设置已保存')
    
    def show_about(self):
        QMessageBox.about(
            self,
            '关于',
            '''论文排版优化工具 v1.0

功能特点:
• 支持 Word (.docx) 和 LaTeX (.tex) 文档
• 内置多种期刊和学位论文模板
• 支持 AI 自动识别文档模板
• 可自定义模板规则
• 一键优化文档排版格式
• 论文快速生成（填写式）

作者: 夕岸摇
'''
        )


class SectionFormatterPage(GlassEffectPage):
    def __init__(self, parent):
        super().__init__(parent)
        self.main_window = parent
        self.current_file = None
        self.document = None
        self.sections = []
        self.selected_section_index = -1
        self.init_ui()
    
    def init_ui(self):
        self.init_page_layout('分部分排版', '#00BCD4', [700, 300])
    
    def create_left_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        file_group = QGroupBox('文档操作')
        file_layout = QVBoxLayout(file_group)
        
        open_btn = QPushButton('📂 打开文档')
        open_btn.setStyleSheet('background-color: #00BCD4; color: white; padding: 8px;')
        open_btn.clicked.connect(self.open_document)
        file_layout.addWidget(open_btn)
        
        self.file_path_label = QLabel('未选择文件')
        self.file_path_label.setStyleSheet('color: gray;')
        self.file_path_label.setWordWrap(True)
        file_layout.addWidget(self.file_path_label)
        
        layout.addWidget(file_group)
        
        section_group = QGroupBox('章节列表')
        section_layout = QVBoxLayout(section_group)
        
        self.section_list = QListWidget()
        self.section_list.itemClicked.connect(self.on_section_selected)
        section_layout.addWidget(self.section_list)
        
        btn_layout = QHBoxLayout()
        
        refresh_btn = QPushButton('刷新章节')
        refresh_btn.clicked.connect(self.refresh_sections)
        btn_layout.addWidget(refresh_btn)
        
        section_layout.addLayout(btn_layout)
        
        layout.addWidget(section_group)
        
        return panel
    
    def create_right_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        preview_group = QGroupBox('内容预览')
        preview_layout = QVBoxLayout(preview_group)
        
        self.content_preview = QTextEdit()
        self.content_preview.setReadOnly(True)
        self.content_preview.setPlaceholderText('选择章节后显示内容...')
        preview_layout.addWidget(self.content_preview)
        
        layout.addWidget(preview_group)
        
        style_group = QGroupBox('标题样式设置')
        style_layout = QVBoxLayout(style_group)
        
        self.style_combo = QComboBox()
        self.style_combo.addItems([
            '正文',
            '标题 1',
            '标题 2',
            '标题 3',
            '标题 4',
            '标题 5',
            '标题 6',
            '标题 7',
            '标题 8',
            '标题 9',
            '标题 10'
        ])
        style_layout.addWidget(QLabel('设置为Word样式:'))
        style_layout.addWidget(self.style_combo)
        
        layout.addWidget(style_group)
        
        format_group = QGroupBox('段落排版设置')
        format_layout = QVBoxLayout(format_group)
        
        format_settings_layout = QFormLayout()
        
        self.indent_check = QCheckBox('自定义首行缩进')
        self.indent_check.setChecked(False)
        format_settings_layout.addRow(self.indent_check)
        
        self.indent_spin = QSpinBox()
        self.indent_spin.setRange(0, 10)
        self.indent_spin.setValue(2)
        self.indent_spin.setSuffix(' 字符')
        self.indent_spin.setEnabled(False)
        self.indent_check.stateChanged.connect(lambda s: self.indent_spin.setEnabled(s))
        format_settings_layout.addRow('缩进:', self.indent_spin)
        
        self.line_spacing_check = QCheckBox('自定义行距')
        self.line_spacing_check.setChecked(False)
        format_settings_layout.addRow(self.line_spacing_check)
        
        self.line_spacing_spin = QDoubleSpinBox()
        self.line_spacing_spin.setRange(0.5, 3.0)
        self.line_spacing_spin.setSingleStep(0.1)
        self.line_spacing_spin.setValue(1.5)
        self.line_spacing_spin.setSuffix(' 倍')
        self.line_spacing_spin.setEnabled(False)
        self.line_spacing_check.stateChanged.connect(lambda s: self.line_spacing_spin.setEnabled(s))
        format_settings_layout.addRow('行距:', self.line_spacing_spin)
        
        self.alignment_combo = QComboBox()
        self.alignment_combo.addItems(['默认', '两端对齐', '左对齐', '居中', '右对齐', '分散对齐'])
        format_settings_layout.addRow('对齐方式:', self.alignment_combo)
        
        self.font_name_edit = QLineEdit()
        self.font_name_edit.setPlaceholderText('留空使用默认')
        format_settings_layout.addRow('字体:', self.font_name_edit)
        
        self.font_size_check = QCheckBox('自定义字号')
        self.font_size_check.setChecked(False)
        format_settings_layout.addRow(self.font_size_check)
        
        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(5, 72)
        self.font_size_spin.setValue(12)
        self.font_size_spin.setSuffix(' pt')
        self.font_size_spin.setEnabled(False)
        self.font_size_check.stateChanged.connect(lambda s: self.font_size_spin.setEnabled(s))
        format_settings_layout.addRow('字号:', self.font_size_spin)
        
        format_layout.addLayout(format_settings_layout)
        
        btn_layout = QHBoxLayout()
        
        reset_btn = QPushButton('重置')
        reset_btn.clicked.connect(self.reset_format_settings)
        btn_layout.addWidget(reset_btn)
        
        apply_btn = QPushButton('应用到此章节')
        apply_btn.setStyleSheet('background-color: #00BCD4; color: white;')
        apply_btn.clicked.connect(self.apply_format_to_section)
        btn_layout.addWidget(apply_btn)
        
        format_layout.addLayout(btn_layout)
        
        layout.addWidget(format_group)
        
        save_layout = QHBoxLayout()
        
        save_btn = QPushButton('💾 保存文档')
        save_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 12px; font-weight: bold;')
        save_btn.clicked.connect(self.save_document)
        save_layout.addWidget(save_btn)
        
        layout.addLayout(save_layout)
        
        return panel
    
    def open_document(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, '选择文档', '',
            'Word文档 (*.docx)'
        )
        
        if file_path:
            self.current_file = file_path
            self.file_path_label.setText(file_path)
            self.file_path_label.setStyleSheet('color: black;')
            self.refresh_sections()
    
    def refresh_sections(self):
        if not self.current_file:
            QMessageBox.warning(self, '提示', '请先打开文档')
            return
        
        try:
            from docx import Document
            self.document = Document(self.current_file)
            self.sections = []
            self.section_list.clear()
            
            for i, para in enumerate(self.document.paragraphs):
                if para.style.name.startswith('Heading') or para.text.strip():
                    self.sections.append({
                        'index': i,
                        'text': para.text[:50] + ('...' if len(para.text) > 50 else '') or '(空段落)',
                        'paragraph': para,
                        'format_settings': {}
                    })
            
            for i, section in enumerate(self.sections):
                item_text = f"{i+1}. {section['text']}"
                item = QListWidgetItem(item_text)
                self.section_list.addItem(item)
            
            QMessageBox.information(self, '成功', f'已解析 {len(self.sections)} 个段落')
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'打开文档失败: {str(e)}')
    
    def on_section_selected(self, item):
        self.selected_section_index = self.section_list.row(item)
        section = self.sections[self.selected_section_index]
        
        self.content_preview.setPlainText(section['paragraph'].text)
        
        settings = section.get('format_settings', {})
        
        style = settings.get('style', '正文')
        style_index = self.style_combo.findText(style)
        self.style_combo.setCurrentIndex(style_index if style_index >= 0 else 0)
        
        self.indent_check.setChecked(bool(settings.get('first_line_indent')))
        self.indent_spin.setValue(settings.get('first_line_indent', 2))
        
        self.line_spacing_check.setChecked(bool(settings.get('line_spacing')))
        self.line_spacing_spin.setValue(settings.get('line_spacing', 1.5))
        
        alignment = settings.get('alignment', '默认')
        index = self.alignment_combo.findText(alignment)
        self.alignment_combo.setCurrentIndex(index if index >= 0 else 0)
        
        self.font_name_edit.setText(settings.get('font_name', ''))
        
        self.font_size_check.setChecked(bool(settings.get('font_size')))
        self.font_size_spin.setValue(settings.get('font_size', 12))
    
    def reset_format_settings(self):
        self.style_combo.setCurrentIndex(0)
        self.indent_check.setChecked(False)
        self.line_spacing_check.setChecked(False)
        self.alignment_combo.setCurrentIndex(0)
        self.font_name_edit.clear()
        self.font_size_check.setChecked(False)
    
    def apply_format_to_section(self):
        if self.selected_section_index < 0:
            QMessageBox.warning(self, '提示', '请先选择一个章节')
            return
        
        settings = {}
        
        style = self.style_combo.currentText()
        if style != '正文':
            settings['style'] = style
        
        if self.indent_check.isChecked():
            settings['first_line_indent'] = self.indent_spin.value()
        if self.line_spacing_check.isChecked():
            settings['line_spacing'] = self.line_spacing_spin.value()
        
        alignment = self.alignment_combo.currentText()
        if alignment != '默认':
            settings['alignment'] = alignment
        
        if self.font_name_edit.text().strip():
            settings['font_name'] = self.font_name_edit.text().strip()
        if self.font_size_check.isChecked():
            settings['font_size'] = self.font_size_spin.value()
        
        self.sections[self.selected_section_index]['format_settings'] = settings
        
        QMessageBox.information(self, '成功', '格式设置已保存')
    
    def save_document(self):
        if not self.document:
            QMessageBox.warning(self, '提示', '请先打开文档')
            return
        
        save_path, _ = QFileDialog.getSaveFileName(
            self, '保存文档',
            os.path.splitext(self.current_file)[0] + '_formatted.docx',
            'Word文档 (*.docx)'
        )
        
        if not save_path:
            return
        
        try:
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            align_map = {
                '左对齐': WD_PARAGRAPH_ALIGNMENT.LEFT,
                '居中': WD_PARAGRAPH_ALIGNMENT.CENTER,
                '右对齐': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                '两端对齐': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                '分散对齐': WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
            }
            
            style_map = {
                '标题 1': 'Heading 1',
                '标题 2': 'Heading 2',
                '标题 3': 'Heading 3',
                '标题 4': 'Heading 4',
                '标题 5': 'Heading 5',
                '标题 6': 'Heading 6',
                '标题 7': 'Heading 7',
                '标题 8': 'Heading 8',
                '标题 9': 'Heading 9',
                '标题 10': 'Heading 10'
            }
            
            for section in self.sections:
                para = section['paragraph']
                settings = section.get('format_settings', {})
                
                if settings.get('style') in style_map:
                    para.style = style_map[settings['style']]
                
                if settings.get('first_line_indent'):
                    para.paragraph_format.first_line_indent = Inches(settings['first_line_indent'] * 0.35)
                
                if settings.get('line_spacing'):
                    para.paragraph_format.line_spacing = settings['line_spacing']
                
                if settings.get('alignment') in align_map:
                    para.paragraph_format.alignment = align_map[settings['alignment']]
                
                if settings.get('font_name') or settings.get('font_size'):
                    for run in para.runs:
                        if settings.get('font_name'):
                            run.font.name = settings['font_name']
                        if settings.get('font_size'):
                            run.font.size = Pt(settings['font_size'])
            
            self.document.save(save_path)
            QMessageBox.information(self, '成功', f'文档已保存到: {save_path}')
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'保存文档失败: {str(e)}')


class StartupDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle('论文排版优化工具')
        self.setFixedSize(600, 400)
        self.setWindowFlags(Qt.WindowStaysOnTopHint)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        container = QWidget()
        container.setStyleSheet('''
            QWidget {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 1, y2: 1,
                    stop: 0 #667eea,
                    stop: 0.5 #764ba2,
                    stop: 1 #f093fb
                );
            }
        ''')
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(40, 40, 40, 40)
        container_layout.setSpacing(20)
        
        logo_label = QLabel('📄')
        logo_label.setStyleSheet('font-size: 80px;')
        logo_label.setAlignment(Qt.AlignCenter)
        container_layout.addWidget(logo_label)
        
        title_label = QLabel('论文排版优化工具')
        title_label.setStyleSheet('font-size: 32px; font-weight: bold; color: white;')
        title_label.setAlignment(Qt.AlignCenter)
        container_layout.addWidget(title_label)
        
        subtitle_label = QLabel('智能排版 · 一键优化')
        subtitle_label.setStyleSheet('font-size: 18px; color: rgba(255, 255, 255, 0.85);')
        subtitle_label.setAlignment(Qt.AlignCenter)
        container_layout.addWidget(subtitle_label)
        
        container_layout.addStretch()
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet('''
            QProgressBar {
                border: 2px solid rgba(255, 255, 255, 0.3);
                border-radius: 8px;
                text-align: center;
                color: white;
                background: rgba(255, 255, 255, 0.2);
                height: 20px;
            }
            QProgressBar::chunk {
                background: qlineargradient(
                    x1: 0, y1: 0, x2: 1, y2: 0,
                    stop: 0 #27ae60,
                    stop: 1 #2ecc71
                );
                border-radius: 6px;
            }
        ''')
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        container_layout.addWidget(self.progress_bar)
        
        self.status_label = QLabel('正在启动...')
        self.status_label.setStyleSheet('font-size: 14px; color: white;')
        self.status_label.setAlignment(Qt.AlignCenter)
        container_layout.addWidget(self.status_label)
        
        layout.addWidget(container)
        
        self.timer = QTimer()
        self.timer.timeout.connect(self.update_progress)
        self.timer.start(50)
        self.progress = 0
        self.steps = [
            (10, '加载界面组件...'),
            (30, '初始化模板管理器...'),
            (50, '配置AI识别模块...'),
            (70, '准备文档解析器...'),
            (90, '完成启动准备...'),
            (100, '欢迎使用！')
        ]
        self.current_step = 0
    
    def update_progress(self):
        if self.progress >= 100:
            self.timer.stop()
            QTimer.singleShot(500, self.accept)
            return
        
        if self.current_step < len(self.steps) and self.progress >= self.steps[self.current_step][0]:
            self.status_label.setText(self.steps[self.current_step][1])
            self.current_step += 1
        
        self.progress += 1
        self.progress_bar.setValue(self.progress)


def main():
    app = QApplication(sys.argv)
    app.setApplicationName('论文排版优化工具')
    app.setOrganizationName('PaperFormatter')
    
    startup_errors = []
    
    if FLUENT_WIDGETS_AVAILABLE:
        try:
            setTheme(Theme.DARK)
            setThemeColor('#009faa')
        except Exception as e:
            startup_errors.append(f'设置Fluent主题失败: {str(e)}')
            import traceback
            startup_errors.append(traceback.format_exc())
    else:
        startup_errors.append('PyQt-Fluent-Widgets 库未正确加载')
        if 'FLUENT_IMPORT_ERROR' in globals():
            startup_errors.append(f'错误详情: {FLUENT_IMPORT_ERROR}')
    
    font = QFont('Microsoft YaHei', 10)
    app.setFont(font)
    
    try:
        startup_dialog = StartupDialog()
        startup_dialog.exec_()
        
        window = MainWindow()
        
        if startup_errors:
            from PyQt5.QtWidgets import QMessageBox
            error_text = '\n\n'.join(startup_errors)
            QMessageBox.warning(
                window,
                '启动警告',
                f'程序启动时检测到以下问题：\n\n{error_text}\n\n程序将继续运行，但部分UI功能可能受限。'
            )
        
        if PYWINSTYLES_AVAILABLE:
            try:
                pywinstyles.apply_style(window, "mica")
            except Exception as e:
                startup_errors.append(f'应用Mica效果失败: {str(e)}')
                window.setStyleSheet('''
                    QMainWindow {
                        background: qlineargradient(
                            x1: 0, y1: 0, x2: 1, y2: 1,
                            stop: 0 #667eea,
                            stop: 0.5 #764ba2,
                            stop: 1 #f093fb
                        );
                    }
                ''')
        else:
            window.setStyleSheet('''
                QMainWindow {
                    background: qlineargradient(
                        x1: 0, y1: 0, x2: 1, y2: 1,
                        stop: 0 #667eea,
                        stop: 0.5 #764ba2,
                        stop: 1 #f093fb
                    );
                }
            ''')
        
        window.show()
        sys.exit(app.exec_())
    except Exception as e:
        import traceback
        error_msg = f'启动失败: {str(e)}\n\n{traceback.format_exc()}'
        
        try:
            from PyQt5.QtWidgets import QMessageBox, QWidget
            msg_box = QMessageBox()
            msg_box.setIcon(QMessageBox.Critical)
            msg_box.setWindowTitle('启动失败')
            msg_box.setText('程序启动失败！')
            msg_box.setDetailedText(error_msg)
            msg_box.exec_()
        except:
            print(f'启动失败: {error_msg}')
            with open('startup_error.log', 'w', encoding='utf-8') as f:
                f.write(error_msg)
        sys.exit(1)


if __name__ == '__main__':
    main()
