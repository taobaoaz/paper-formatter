import os
import re
from abc import ABC, abstractmethod
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
import chardet


@dataclass
class DocumentElement:
    element_type: str
    content: str
    style: Dict[str, Any] = field(default_factory=dict)
    position: int = 0
    level: int = 0


@dataclass
class ParsedDocument:
    file_path: str
    file_type: str
    elements: List[DocumentElement] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    raw_content: str = ""


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        pass
    
    def detect_encoding(self, file_path: str) -> str:
        with open(file_path, 'rb') as f:
            raw_data = f.read(10000)
            result = chardet.detect(raw_data)
            return result.get('encoding', 'utf-8')


class DocxParser(BaseParser):
    def get_supported_extensions(self) -> List[str]:
        return ['.docx']
    
    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(file_path)
        parsed_doc = ParsedDocument(
            file_path=file_path,
            file_type='docx',
            metadata=self._extract_metadata(doc)
        )
        
        position = 0
        for para in doc.paragraphs:
            element = self._parse_paragraph(para, position)
            if element:
                parsed_doc.elements.append(element)
                position += 1
        
        for table_idx, table in enumerate(doc.tables):
            table_element = self._parse_table(table, position, table_idx)
            parsed_doc.elements.append(table_element)
            position += 1
        
        return parsed_doc
    
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        core_props = doc.core_properties
        return {
            'author': core_props.author or '',
            'title': core_props.title or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else ''
        }
    
    def _parse_paragraph(self, para, position: int) -> Optional[DocumentElement]:
        text = para.text.strip()
        if not text:
            return None
        
        style_info = {
            'style_name': para.style.name if para.style else '',
            'alignment': str(para.alignment) if para.alignment else '',
            'font_name': '',
            'font_size': None,
            'bold': False,
            'italic': False
        }
        
        if para.runs:
            first_run = para.runs[0]
            if first_run.font.name:
                style_info['font_name'] = first_run.font.name
            if first_run.font.size:
                style_info['font_size'] = first_run.font.size.pt
            style_info['bold'] = first_run.bold or False
            style_info['italic'] = first_run.italic or False
        
        element_type = self._determine_element_type(para, text)
        level = self._determine_heading_level(para, text)
        
        return DocumentElement(
            element_type=element_type,
            content=text,
            style=style_info,
            position=position,
            level=level
        )
    
    def _determine_element_type(self, para, text: str) -> str:
        style_name = para.style.name if para.style else ''
        
        if 'Title' in style_name or '标题' in style_name:
            return 'title'
        elif 'Heading' in style_name or '标题' in style_name:
            return 'heading'
        elif text.startswith(('摘要', 'Abstract', 'ABSTRACT')):
            return 'abstract'
        elif text.startswith(('关键词', 'Keywords', 'KEYWORDS', '关键字')):
            return 'keywords'
        elif text.startswith(('参考文献', 'References', 'REFERENCES')):
            return 'references'
        elif re.match(r'^\[\d+\]', text) or re.match(r'^\d+\.', text):
            return 'reference_item'
        else:
            return 'paragraph'
    
    def _determine_heading_level(self, para, text: str) -> int:
        style_name = para.style.name if para.style else ''
        
        if 'Heading 1' in style_name or '标题 1' in style_name:
            return 1
        elif 'Heading 2' in style_name or '标题 2' in style_name:
            return 2
        elif 'Heading 3' in style_name or '标题 3' in style_name:
            return 3
        
        patterns = [
            (r'^第[一二三四五六七八九十]+[章节部分]', 1),
            (r'^[1-9]\s', 1),
            (r'^[1-9]\.[0-9]+\s', 2),
            (r'^[1-9]\.[0-9]+\.[0-9]+\s', 3),
        ]
        
        for pattern, level in patterns:
            if re.match(pattern, text):
                return level
        
        return 0
    
    def _parse_table(self, table, position: int, table_idx: int) -> DocumentElement:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        return DocumentElement(
            element_type='table',
            content=str(table_data),
            style={'table_index': table_idx, 'rows': len(table.rows), 'cols': len(table.columns)},
            position=position
        )


class TexParser(BaseParser):
    def get_supported_extensions(self) -> List[str]:
        return ['.tex']
    
    def parse(self, file_path: str) -> ParsedDocument:
        encoding = self.detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
            content = f.read()
        
        parsed_doc = ParsedDocument(
            file_path=file_path,
            file_type='tex',
            raw_content=content,
            metadata=self._extract_metadata(content)
        )
        
        self._parse_content(content, parsed_doc)
        
        return parsed_doc
    
    def _extract_metadata(self, content: str) -> Dict[str, Any]:
        metadata = {}
        
        title_match = re.search(r'\\title\{([^}]+)\}', content)
        if title_match:
            metadata['title'] = title_match.group(1)
        
        author_match = re.search(r'\\author\{([^}]+)\}', content)
        if author_match:
            metadata['author'] = author_match.group(1)
        
        date_match = re.search(r'\\date\{([^}]+)\}', content)
        if date_match:
            metadata['date'] = date_match.group(1)
        
        return metadata
    
    def _parse_content(self, content: str, parsed_doc: ParsedDocument):
        position = 0
        
        patterns = [
            (r'\\title\{([^}]+)\}', 'title', 0),
            (r'\\begin\{abstract\}(.*?)\\end\{abstract\}', 'abstract', 0, re.DOTALL),
            (r'\\section\{([^}]+)\}', 'heading', 1),
            (r'\\subsection\{([^}]+)\}', 'heading', 2),
            (r'\\subsubsection\{([^}]+)\}', 'heading', 3),
            (r'\\begin\{thebibliography\}.*?\\end\{thebibliography\}', 'references', 0, re.DOTALL),
        ]
        
        for pattern_data in patterns:
            if len(pattern_data) == 3:
                pattern, elem_type, level = pattern_data
                flags = 0
            else:
                pattern, elem_type, level, flags = pattern_data
            
            for match in re.finditer(pattern, content, flags):
                element = DocumentElement(
                    element_type=elem_type,
                    content=match.group(1) if match.lastindex else match.group(0),
                    style={'raw_match': match.group(0)},
                    position=position,
                    level=level
                )
                parsed_doc.elements.append(element)
                position += 1
        
        paragraphs = re.split(r'\n\s*\n', content)
        for para in paragraphs:
            para = para.strip()
            if para and not para.startswith('%') and not para.startswith('\\'):
                if len(para) > 20:
                    element = DocumentElement(
                        element_type='paragraph',
                        content=para[:500],
                        style={},
                        position=position
                    )
                    parsed_doc.elements.append(element)
                    position += 1


class DocumentParserFactory:
    _parsers: Dict[str, BaseParser] = {}
    
    @classmethod
    def register_parser(cls, parser: BaseParser):
        for ext in parser.get_supported_extensions():
            cls._parsers[ext.lower()] = parser
    
    @classmethod
    def parse(cls, file_path: str) -> ParsedDocument:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in cls._parsers:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        return cls._parsers[ext].parse(file_path)
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        return list(cls._parsers.keys())


DocumentParserFactory.register_parser(DocxParser())
DocumentParserFactory.register_parser(TexParser())
