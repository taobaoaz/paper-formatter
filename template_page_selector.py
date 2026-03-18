"""
模板页面选择器
Template Page Selector
允许用户上传 Word 模板后，选择哪两页是封面和声明页
"""

from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QGridLayout,
    QLabel, QPushButton, QComboBox, QGroupBox,
    QFileDialog, QMessageBox, QListWidget, QListWidgetItem,
    QScrollArea, QWidget, QRadioButton, QButtonGroup
)
from PyQt5.QtCore import Qt, QSize
from PyQt5.QtGui import QPixmap, QImage
from docx import Document
import os
import tempfile
import win32com.client


class TemplatePageSelectorDialog(QDialog):
    """模板页面选择对话框"""
    
    def __init__(self, template_path=None, parent=None):
        super().__init__(parent)
        self.template_path = template_path
        self.page_count = 0
        self.cover_page_index = 0  # 封面页索引（从 0 开始）
        self.declaration_page_index = 1  # 声明页索引（从 0 开始）
        
        self.init_ui()
        
        if template_path:
            self.load_template(template_path)
    
    def init_ui(self):
        """初始化界面"""
        self.setWindowTitle('选择封面和声明页')
        self.setMinimumSize(800, 600)
        
        layout = QVBoxLayout(self)
        
        # 1. 说明
        info_label = QLabel(
            '📄 请从上传的模板中选择封面和声明页的位置\n\n'
            '提示：通常封面是第 1 页，声明页是第 2 页'
        )
        info_label.setStyleSheet('font-size: 13px; padding: 10px; background-color: #f0f0f0;')
        layout.addWidget(info_label)
        
        # 2. 页面选择区域
        select_group = QGroupBox('📑 页面选择')
        select_layout = QGridLayout()
        
        # 封面页选择
        select_layout.addWidget(QLabel('封面页:'), 0, 0)
        self.cover_page_combo = QComboBox()
        self.cover_page_combo.currentIndexChanged.connect(self.on_cover_page_changed)
        select_layout.addWidget(self.cover_page_combo, 0, 1)
        
        # 声明页选择
        select_layout.addWidget(QLabel('声明页:'), 1, 0)
        self.declaration_page_combo = QComboBox()
        self.declaration_page_combo.currentIndexChanged.connect(self.on_declaration_page_changed)
        select_layout.addWidget(self.declaration_page_combo, 1, 1)
        
        select_group.setLayout(select_layout)
        layout.addWidget(select_group)
        
        # 3. 页面预览区域
        preview_group = QGroupBox('👁️ 页面预览')
        preview_layout = QVBoxLayout()
        
        # 创建滚动区域来显示页面列表
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        
        preview_widget = QWidget()
        self.preview_layout = QVBoxLayout(preview_widget)
        
        # 页面列表
        self.page_list = QListWidget()
        self.page_list.setViewMode(QListWidget.ListMode)
        self.page_list.setSpacing(5)
        self.page_list.itemClicked.connect(self.on_page_selected)
        
        self.preview_layout.addWidget(self.page_list)
        scroll.setWidget(preview_widget)
        preview_layout.addWidget(scroll)
        
        preview_group.setLayout(preview_layout)
        layout.addWidget(preview_group)
        
        # 4. 按钮
        button_layout = QHBoxLayout()
        
        self.upload_btn = QPushButton('📂 上传模板')
        self.upload_btn.clicked.connect(self.upload_template)
        
        self.ok_btn = QPushButton('✅ 确定')
        self.ok_btn.clicked.connect(self.accept_selection)
        self.ok_btn.setStyleSheet('background-color: #4CAF50; color: white; padding: 8px;')
        self.ok_btn.setEnabled(False)
        
        self.cancel_btn = QPushButton('❌ 取消')
        self.cancel_btn.clicked.connect(self.reject)
        
        button_layout.addWidget(self.upload_btn)
        button_layout.addStretch()
        button_layout.addWidget(self.ok_btn)
        button_layout.addWidget(self.cancel_btn)
        
        layout.addLayout(button_layout)
    
    def load_template(self, template_path):
        """加载模板文件"""
        try:
            self.template_path = template_path
            
            # 获取页数
            self.page_count = self.get_word_page_count(template_path)
            
            # 填充下拉框
            self.cover_page_combo.clear()
            self.declaration_page_combo.clear()
            
            for i in range(self.page_count):
                page_num = i + 1
                self.cover_page_combo.addItem(f'第 {page_num} 页', i)
                self.declaration_page_combo.addItem(f'第 {page_num} 页', i)
            
            # 默认选择
            self.cover_page_combo.setCurrentIndex(0)
            if self.page_count > 1:
                self.declaration_page_combo.setCurrentIndex(1)
            
            # 填充页面列表
            self.page_list.clear()
            for i in range(self.page_count):
                page_num = i + 1
                item_text = f'第 {page_num} 页'
                
                if i == 0:
                    item_text += ' - [封面]'
                elif i == 1:
                    item_text += ' - [声明页]'
                
                item = QListWidgetItem(item_text)
                item.setData(Qt.UserRole, i)  # 存储页索引
                self.page_list.addItem(item)
            
            self.ok_btn.setEnabled(True)
            
            QMessageBox.information(
                self,
                '成功',
                f'模板加载成功！\n共 {self.page_count} 页'
            )
            
        except Exception as e:
            QMessageBox.critical(self, '错误', f'加载模板失败：{str(e)}')
    
    def get_word_page_count(self, file_path):
        """获取 Word 文档页数"""
        try:
            # 方法 1：使用 win32com（需要安装 pywin32）
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            page_count = doc.ComputeStatistics(2)  # wdStatisticPages = 2
            doc.Close(False)
            word.Quit()
            return page_count
            
        except Exception:
            # 方法 2：估算页数（不精确）
            doc = Document(file_path)
            para_count = len(doc.paragraphs)
            # 粗略估算：每页约 30 段
            estimated_pages = max(1, (para_count // 30) + 1)
            return estimated_pages
    
    def on_cover_page_changed(self, index):
        """封面页改变时的处理"""
        cover_index = self.cover_page_combo.currentData()
        self.update_page_list()
    
    def on_declaration_page_changed(self, index):
        """声明页改变时的处理"""
        declaration_index = self.declaration_page_combo.currentData()
        self.update_page_list()
    
    def update_page_list(self):
        """更新页面列表显示"""
        cover_index = self.cover_page_combo.currentData()
        declaration_index = self.declaration_page_combo.currentData()
        
        for i in range(self.page_list.count()):
            item = self.page_list.item(i)
            page_num = i + 1
            
            # 更新标记
            text = f'第 {page_num} 页'
            
            if i == cover_index:
                text += ' - [封面 ✓]'
            elif i == declaration_index:
                text += ' - [声明页 ✓]'
            
            item.setText(text)
    
    def on_page_selected(self, item):
        """页面列表项被点击"""
        page_index = item.data(Qt.UserRole)
        
        # 如果点击的是封面页
        if page_index == self.cover_page_combo.currentData():
            QMessageBox.information(
                self,
                '封面页',
                f'已选择第 {page_index + 1} 页作为封面页'
            )
        # 如果点击的是声明页
        elif page_index == self.declaration_page_combo.currentData():
            QMessageBox.information(
                self,
                '声明页',
                f'已选择第 {page_index + 1} 页作为声明页'
            )
    
    def upload_template(self):
        """上传模板文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            '选择模板文件',
            '',
            'Word 文档 (*.docx)'
        )
        
        if file_path:
            self.load_template(file_path)
    
    def accept_selection(self):
        """确认选择"""
        cover_index = self.cover_page_combo.currentData()
        declaration_index = self.declaration_page_combo.currentData()
        
        # 检查是否选择了相同的页面
        if cover_index == declaration_index:
            QMessageBox.warning(
                self,
                '警告',
                '封面页和声明页不能是同一页！'
            )
            return
        
        # 保存选择结果
        self.selected_cover_index = cover_index
        self.selected_declaration_index = declaration_index
        
        self.accept()
    
    def get_selection(self):
        """获取选择结果"""
        return {
            'cover_page_index': self.selected_cover_index,
            'declaration_page_index': self.selected_declaration_index,
            'template_path': self.template_path
        }


class TemplatePageManager:
    """模板页面管理器"""
    
    def __init__(self):
        self.template_path = None
        self.cover_page_index = 0
        self.declaration_page_index = 1
    
    def extract_pages(self, template_path, cover_index, declaration_index, output_dir=None):
        """从模板中提取封面和声明页"""
        from docx import Document
        
        if output_dir is None:
            output_dir = os.path.dirname(template_path)
        
        # 加载模板
        doc = Document(template_path)
        
        # 创建封面文档
        cover_doc = Document()
        self.copy_page_to_doc(doc, cover_index, cover_doc)
        cover_path = os.path.join(output_dir, 'template_cover.docx')
        cover_doc.save(cover_path)
        
        # 创建声明页文档
        declaration_doc = Document()
        self.copy_page_to_doc(doc, declaration_index, declaration_doc)
        declaration_path = os.path.join(output_dir, 'template_declaration.docx')
        declaration_doc.save(declaration_path)
        
        return {
            'cover_path': cover_path,
            'declaration_path': declaration_path
        }
    
    def copy_page_to_doc(self, source_doc, page_index, target_doc):
        """将指定页复制到目标文档"""
        # 注意：python-docx 不直接支持按页操作
        # 这里简化处理：复制所有段落，用户手动删除不需要的内容
        
        for para in source_doc.paragraphs:
            new_para = target_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                if run.font.size:
                    new_run.font.size = run.font.size
    
    def merge_with_thesis(self, thesis_path, cover_path, declaration_path, output_path=None):
        """合并封面、声明页和论文"""
        from docx import Document
        
        if output_path is None:
            base, ext = os.path.splitext(thesis_path)
            output_path = f"{base}_完整版{ext}"
        
        # 创建新文档
        output_doc = Document()
        
        # 添加封面
        cover_doc = Document(cover_path)
        for para in cover_doc.paragraphs:
            new_para = output_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
        
        # 插入分节符
        output_doc.add_section()
        
        # 添加声明页
        declaration_doc = Document(declaration_path)
        for para in declaration_doc.paragraphs:
            new_para = output_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
        
        # 插入分节符
        output_doc.add_section()
        
        # 添加论文正文
        thesis_doc = Document(thesis_path)
        for para in thesis_doc.paragraphs:
            new_para = output_doc.add_paragraph()
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
        
        # 保存
        output_doc.save(output_path)
        
        return output_path


def select_template_pages(template_path=None, parent=None):
    """便捷函数：打开页面选择对话框"""
    dialog = TemplatePageSelectorDialog(template_path, parent)
    
    if dialog.exec_() == QDialog.Accepted:
        return dialog.get_selection()
    else:
        return None


if __name__ == '__main__':
    import sys
    from PyQt5.QtWidgets import QApplication
    
    app = QApplication(sys.argv)
    dialog = TemplatePageSelectorDialog()
    dialog.show()
    sys.exit(app.exec_())
