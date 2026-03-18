import os
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config


@dataclass
class DocumentField:
    field_id: str
    field_name: str
    field_type: str
    required: bool
    default_value: str = ""
    placeholder: str = ""
    order: int = 0
    group: str = "基本信息"


@dataclass
class DocumentTemplate:
    template_id: str
    template_name: str
    description: str
    fields: List[DocumentField]
    sections: List[Dict[str, Any]] = field(default_factory=list)


DEFAULT_FIELDS = [
    DocumentField("title", "论文标题", "text", True, "", "请输入论文标题", 1, "基本信息"),
    DocumentField("title_en", "英文标题", "text", False, "", "English Title", 2, "基本信息"),
    DocumentField("author", "作者姓名", "text", True, "", "请输入作者姓名", 3, "基本信息"),
    DocumentField("author_en", "作者英文名", "text", False, "", "Author Name", 4, "基本信息"),
    DocumentField("affiliation", "作者单位", "text", True, "", "请输入作者单位", 5, "基本信息"),
    DocumentField("affiliation_en", "作者单位英文名", "text", False, "", "Affiliation", 6, "基本信息"),
    DocumentField("email", "通讯邮箱", "text", False, "", "email@example.com", 7, "基本信息"),
    
    DocumentField("abstract", "中文摘要", "textarea", True, "", "请输入中文摘要（200-300字）", 10, "摘要关键词"),
    DocumentField("abstract_en", "英文摘要", "textarea", False, "", "English Abstract", 11, "摘要关键词"),
    DocumentField("keywords", "中文关键词", "text", True, "", "关键词1; 关键词2; 关键词3", 12, "摘要关键词"),
    DocumentField("keywords_en", "英文关键词", "text", False, "", "keyword1; keyword2; keyword3", 13, "摘要关键词"),
    
    DocumentField("introduction", "引言", "textarea", False, "", "请输入引言内容", 20, "正文内容"),
    DocumentField("method", "方法", "textarea", False, "", "请输入研究方法", 21, "正文内容"),
    DocumentField("result", "结果", "textarea", False, "", "请输入研究结果", 22, "正文内容"),
    DocumentField("discussion", "讨论", "textarea", False, "", "请输入讨论内容", 23, "正文内容"),
    DocumentField("conclusion", "结论", "textarea", False, "", "请输入结论", 24, "正文内容"),
    
    DocumentField("references", "参考文献", "textarea", False, "", "请输入参考文献（每行一条）", 30, "参考文献"),
    DocumentField("acknowledgement", "致谢", "textarea", False, "", "请输入致谢内容", 31, "参考文献"),
]

SECTION_TEMPLATES = {
    "学位论文": [
        {"title": "摘要", "field": "abstract", "level": 0},
        {"title": "Abstract", "field": "abstract_en", "level": 0},
        {"title": "目录", "type": "toc", "level": 0},
        {"title": "第一章 绪论", "field": "introduction", "level": 1},
        {"title": "第二章 研究方法", "field": "method", "level": 1},
        {"title": "第三章 研究结果", "field": "result", "level": 1},
        {"title": "第四章 讨论", "field": "discussion", "level": 1},
        {"title": "第五章 结论", "field": "conclusion", "level": 1},
        {"title": "参考文献", "field": "references", "level": 1},
        {"title": "致谢", "field": "acknowledgement", "level": 1},
    ],
    "中文期刊": [
        {"title": "摘要", "field": "abstract", "level": 0},
        {"title": "关键词", "field": "keywords", "level": 0},
        {"title": "Abstract", "field": "abstract_en", "level": 0},
        {"title": "Keywords", "field": "keywords_en", "level": 0},
        {"title": "1 引言", "field": "introduction", "level": 1},
        {"title": "2 方法", "field": "method", "level": 1},
        {"title": "3 结果", "field": "result", "level": 1},
        {"title": "4 讨论", "field": "discussion", "level": 1},
        {"title": "5 结论", "field": "conclusion", "level": 1},
        {"title": "参考文献", "field": "references", "level": 1},
    ],
    "国际期刊": [
        {"title": "Abstract", "field": "abstract_en", "level": 0},
        {"title": "Keywords", "field": "keywords_en", "level": 0},
        {"title": "1. Introduction", "field": "introduction", "level": 1},
        {"title": "2. Method", "field": "method", "level": 1},
        {"title": "3. Results", "field": "result", "level": 1},
        {"title": "4. Discussion", "field": "discussion", "level": 1},
        {"title": "5. Conclusion", "field": "conclusion", "level": 1},
        {"title": "References", "field": "references", "level": 1},
    ],
    "自定义": [
        {"title": "摘要", "field": "abstract", "level": 0},
        {"title": "关键词", "field": "keywords", "level": 0},
        {"title": "正文", "field": "introduction", "level": 1},
        {"title": "参考文献", "field": "references", "level": 1},
    ]
}


class DocumentGenerator:
    def __init__(self):
        pass
    
    def generate_document_with_sections(self, basic_values: Dict[str, str],
                                        sections: List[Dict[str, Any]],
                                        format_rules: Dict[str, Any],
                                        doc_type: str = "中文期刊",
                                        output_path: str = None,
                                        attachments: List[str] = None) -> str:
        doc = Document()
        
        self._setup_page_settings(doc, format_rules)
        
        self._add_title_page(doc, basic_values, format_rules)
        
        self._add_abstract_section(doc, basic_values, format_rules)
        
        for section in sections:
            self._add_dynamic_section(doc, section, format_rules)
        
        if basic_values.get('references'):
            self._add_references_section(doc, basic_values['references'], format_rules)
        
        if basic_values.get('acknowledgement'):
            self._add_acknowledgement_section(doc, basic_values['acknowledgement'], format_rules)
        
        if attachments:
            self._add_attachments_section(doc, attachments, format_rules)
        
        if output_path is None:
            output_dir = config.OUTPUT_DIR
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"generated_{basic_values.get('title', 'document')[:20]}.docx")
        
        doc.save(output_path)
        return output_path
    
    def _add_abstract_section(self, doc: Document, values: Dict[str, str], rules: Dict[str, Any]):
        if values.get('abstract'):
            abstract_title = doc.add_paragraph()
            abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = abstract_title.add_run('摘要')
            self._set_font(run, rules.get('title_font', '黑体'), rules.get('abstract_size', 12))
            run.bold = True
            
            abstract_para = doc.add_paragraph()
            run = abstract_para.add_run(values['abstract'])
            self._set_font(run, rules.get('abstract_font', '宋体'), rules.get('abstract_size', 10.5))
            abstract_para.paragraph_format.first_line_indent = Cm(0.74)
            doc.add_paragraph()
        
        if values.get('keywords'):
            kw_para = doc.add_paragraph()
            run = kw_para.add_run('关键词: ')
            self._set_font(run, rules.get('font_family', '宋体'), rules.get('font_size', 10.5))
            run.bold = True
            run = kw_para.add_run(values['keywords'])
            self._set_font(run, rules.get('font_family', '宋体'), rules.get('font_size', 10.5))
            doc.add_paragraph()
        
        if values.get('abstract_en'):
            abstract_en_title = doc.add_paragraph()
            abstract_en_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = abstract_en_title.add_run('Abstract')
            self._set_font(run, 'Times New Roman', rules.get('abstract_size', 12))
            run.bold = True
            
            abstract_en_para = doc.add_paragraph()
            run = abstract_en_para.add_run(values['abstract_en'])
            self._set_font(run, 'Times New Roman', rules.get('abstract_size', 10.5))
            abstract_en_para.paragraph_format.first_line_indent = Cm(0.74)
            doc.add_paragraph()
        
        if values.get('keywords_en'):
            kw_en_para = doc.add_paragraph()
            run = kw_en_para.add_run('Keywords: ')
            self._set_font(run, 'Times New Roman', rules.get('font_size', 10.5))
            run.bold = True
            run = kw_en_para.add_run(values['keywords_en'])
            self._set_font(run, 'Times New Roman', rules.get('font_size', 10.5))
            doc.add_paragraph()
    
    def _add_dynamic_section(self, doc: Document, section: Dict[str, Any], rules: Dict[str, Any]):
        level = section.get('level', 1)
        title = section.get('title', '')
        content = section.get('content', '')
        format_settings = section.get('format_settings', {})
        
        if title:
            heading_para = doc.add_paragraph()
            run = heading_para.add_run(title)
            
            heading_sizes = {
                1: rules.get('heading1_size', 16),
                2: rules.get('heading2_size', 14),
                3: rules.get('heading3_size', 12),
                4: rules.get('heading4_size', 12),
                5: rules.get('heading5_size', 11),
                6: rules.get('heading6_size', 11),
                7: rules.get('heading7_size', 10.5),
                8: rules.get('heading8_size', 10.5),
                9: rules.get('heading9_size', 10),
                10: rules.get('heading10_size', 10),
            }
            
            heading_fonts = {
                1: rules.get('heading1_font', '黑体'),
                2: rules.get('heading2_font', '黑体'),
                3: rules.get('heading3_font', '黑体'),
            }
            
            font_name = heading_fonts.get(level, rules.get('font_family', '宋体'))
            font_size = heading_sizes.get(level, 12)
            
            self._set_font(run, font_name, font_size)
            run.bold = True
        
        if content:
            paragraphs = content.strip().split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    content_para = doc.add_paragraph()
                    run = content_para.add_run(para_text.strip())
                    
                    para_font = format_settings.get('font_name', rules.get('font_family', '宋体'))
                    para_font_size = format_settings.get('font_size', rules.get('font_size', 12))
                    self._set_font(run, para_font, para_font_size)
                    
                    pf = content_para.paragraph_format
                    
                    indent = format_settings.get('first_line_indent', rules.get('first_line_indent', 2))
                    if indent:
                        pf.first_line_indent = Cm(indent * 0.35)
                    
                    line_spacing = format_settings.get('line_spacing', rules.get('line_spacing', 1.5))
                    pf.line_spacing = line_spacing
                    
                    if format_settings.get('space_before'):
                        pf.space_before = Pt(format_settings['space_before'])
                    if format_settings.get('space_after'):
                        pf.space_after = Pt(format_settings['space_after'])
                    
                    alignment = format_settings.get('alignment')
                    if alignment:
                        alignment_map = {
                            '两端对齐': WD_ALIGN_PARAGRAPH.JUSTIFY,
                            '左对齐': WD_ALIGN_PARAGRAPH.LEFT,
                            '居中': WD_ALIGN_PARAGRAPH.CENTER,
                            '右对齐': WD_ALIGN_PARAGRAPH.RIGHT,
                            '分散对齐': WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                        }
                        if alignment in alignment_map:
                            content_para.alignment = alignment_map[alignment]
        
        doc.add_paragraph()
    
    def _add_references_section(self, doc: Document, references: str, rules: Dict[str, Any]):
        ref_title = doc.add_paragraph()
        ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_title.add_run('参考文献')
        self._set_font(run, rules.get('title_font', '黑体'), rules.get('heading1_size', 16))
        run.bold = True
        doc.add_paragraph()
        
        ref_lines = references.strip().split('\n')
        for line in ref_lines:
            if line.strip():
                ref_para = doc.add_paragraph()
                run = ref_para.add_run(line.strip())
                self._set_font(run, rules.get('reference_font', '宋体'), rules.get('reference_size', 9))
                ref_para.paragraph_format.line_spacing = 1.25
        
        doc.add_paragraph()
    
    def _add_acknowledgement_section(self, doc: Document, acknowledgement: str, rules: Dict[str, Any]):
        ack_title = doc.add_paragraph()
        ack_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ack_title.add_run('致谢')
        self._set_font(run, rules.get('title_font', '黑体'), rules.get('heading1_size', 16))
        run.bold = True
        doc.add_paragraph()
        
        paragraphs = acknowledgement.strip().split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                ack_para = doc.add_paragraph()
                run = ack_para.add_run(para_text.strip())
                self._set_font(run, rules.get('font_family', '宋体'), rules.get('font_size', 12))
                ack_para.paragraph_format.first_line_indent = Cm(0.74)
                ack_para.paragraph_format.line_spacing = rules.get('line_spacing', 1.5)
        
        doc.add_paragraph()
    
    def _add_attachments_section(self, doc: Document, attachments: List[str], rules: Dict[str, Any]):
        attach_title = doc.add_paragraph()
        attach_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = attach_title.add_run('附件')
        self._set_font(run, rules.get('title_font', '黑体'), rules.get('heading1_size', 16))
        run.bold = True
        doc.add_paragraph()
        
        for i, attach_path in enumerate(attachments, 1):
            if os.path.exists(attach_path):
                file_name = os.path.basename(attach_path)
                file_ext = os.path.splitext(file_name)[1].lower()
                
                attach_para = doc.add_paragraph()
                run = attach_para.add_run(f'附件{i}: {file_name}')
                self._set_font(run, rules.get('font_family', '宋体'), rules.get('font_size', 12))
                run.bold = True
                
                if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                    try:
                        doc.add_picture(attach_path)
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        note_para = doc.add_paragraph()
                        run = note_para.add_run(f'[图片文件: {file_name}]')
                        self._set_font(run, rules.get('font_family', '宋体'), 10)
                        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    note_para = doc.add_paragraph()
                    run = note_para.add_run(f'[附件文件: {file_name}]')
                    self._set_font(run, rules.get('font_family', '宋体'), 10)
                    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
    
    def generate_document(self, field_values: Dict[str, str], 
                          format_rules: Dict[str, Any],
                          doc_type: str = "中文期刊",
                          output_path: str = None) -> str:
        doc = Document()
        
        self._setup_page_settings(doc, format_rules)
        
        sections = SECTION_TEMPLATES.get(doc_type, SECTION_TEMPLATES["自定义"])
        
        self._add_title_page(doc, field_values, format_rules)
        
        for section in sections:
            self._add_section(doc, section, field_values, format_rules)
        
        if output_path is None:
            output_dir = config.OUTPUT_DIR
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"generated_{doc_type}_{field_values.get('title', 'document')[:20]}.docx")
        
        doc.save(output_path)
        return output_path
    
    def _setup_page_settings(self, doc: Document, rules: Dict[str, Any]):
        for section in doc.sections:
            if 'margin_top' in rules:
                section.top_margin = Cm(rules['margin_top'])
            if 'margin_bottom' in rules:
                section.bottom_margin = Cm(rules['margin_bottom'])
            if 'margin_left' in rules:
                section.left_margin = Cm(rules['margin_left'])
            if 'margin_right' in rules:
                section.right_margin = Cm(rules['margin_right'])
    
    def _add_title_page(self, doc: Document, values: Dict[str, str], rules: Dict[str, Any]):
        if values.get('title'):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(values['title'])
            self._set_font(run, rules.get('title_font', '黑体'), rules.get('title_size', 22))
            run.bold = rules.get('title_bold', True)
            doc.add_paragraph()
        
        if values.get('title_en'):
            title_en_para = doc.add_paragraph()
            title_en_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_en_para.add_run(values['title_en'])
            self._set_font(run, 'Times New Roman', rules.get('title_size', 22) - 2)
            doc.add_paragraph()
        
        if values.get('author'):
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(values['author'])
            self._set_font(run, rules.get('font_family', '宋体'), 14)
            doc.add_paragraph()
        
        if values.get('affiliation'):
            aff_para = doc.add_paragraph()
            aff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = aff_para.add_run(values['affiliation'])
            self._set_font(run, rules.get('font_family', '宋体'), 12)
            doc.add_paragraph()
        
        if values.get('email'):
            email_para = doc.add_paragraph()
            email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = email_para.add_run(values['email'])
            self._set_font(run, 'Times New Roman', 10)
            doc.add_paragraph()
    
    def _add_section(self, doc: Document, section: Dict[str, Any], 
                     values: Dict[str, str], rules: Dict[str, Any]):
        section_type = section.get('type', 'content')
        
        if section_type == 'toc':
            doc.add_paragraph("目录")
            doc.add_paragraph("[此处插入自动生成的目录]")
            doc.add_paragraph()
            return
        
        title = section.get('title', '')
        level = section.get('level', 1)
        field_name = section.get('field', '')
        
        if title:
            heading_para = doc.add_paragraph()
            run = heading_para.add_run(title)
            
            if level == 0:
                self._set_font(run, rules.get('title_font', '黑体'), rules.get('title_size', 18))
                run.bold = True
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                heading_key = f'heading{level}_font'
                size_key = f'heading{level}_size'
                self._set_font(run, rules.get(heading_key, '黑体'), rules.get(size_key, 14))
                run.bold = rules.get(f'heading{level}_bold', True)
        
        if field_name and values.get(field_name):
            content = values[field_name]
            
            if field_name in ['keywords', 'keywords_en']:
                content_para = doc.add_paragraph()
                label = "关键词: " if 'keywords' == field_name else "Keywords: "
                run = content_para.add_run(label)
                self._set_font(run, rules.get('keywords_font', '宋体'), rules.get('keywords_size', 10.5))
                run.bold = True
                
                run = content_para.add_run(content)
                self._set_font(run, rules.get('keywords_font', '宋体'), rules.get('keywords_size', 10.5))
            
            elif field_name == 'references':
                ref_lines = content.strip().split('\n')
                for line in ref_lines:
                    if line.strip():
                        ref_para = doc.add_paragraph()
                        run = ref_para.add_run(line.strip())
                        self._set_font(run, rules.get('reference_font', '宋体'), rules.get('reference_size', 9))
            
            elif field_name in ['abstract', 'abstract_en']:
                content_para = doc.add_paragraph()
                run = content_para.add_run(content)
                self._set_font(run, rules.get('abstract_font', '宋体'), rules.get('abstract_size', 10.5))
                content_para.paragraph_format.first_line_indent = Cm(0.74)
            
            else:
                paragraphs = content.strip().split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        content_para = doc.add_paragraph()
                        run = content_para.add_run(para_text.strip())
                        self._set_font(run, rules.get('font_family', '宋体'), rules.get('font_size', 12))
                        content_para.paragraph_format.first_line_indent = Cm(0.74)
                        content_para.paragraph_format.line_spacing = rules.get('line_spacing', 1.5)
        
        doc.add_paragraph()
    
    def _set_font(self, run, font_name: str, font_size: float):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r = run._element.rPr
        if r is None:
            r = OxmlElement('w:rPr')
            run._element.insert(0, r)
        rFonts = r.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            r.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


def get_fields_by_group() -> Dict[str, List[DocumentField]]:
    groups = {}
    for field in DEFAULT_FIELDS:
        if field.group not in groups:
            groups[field.group] = []
        groups[field.group].append(field)
    return groups


def get_all_fields() -> List[DocumentField]:
    return sorted(DEFAULT_FIELDS, key=lambda x: x.order)
