"""
文件预览模块
File Preview Module

功能：
- Word 文档预览
- 文本内容提取
- 文档结构分析
- 简单渲染预览
"""

import os
from docx import Document
from typing import Dict, List, Any


class FilePreview:
    """文件预览器"""
    
    def __init__(self):
        """初始化预览器"""
        self.document = None
        self.file_path = None
    
    def load_file(self, file_path: str) -> bool:
        """
        加载文件
        
        参数：
        - file_path: 文件路径
        
        返回：
        - bool: 是否成功
        """
        try:
            if not os.path.exists(file_path):
                print(f'✗ 文件不存在：{file_path}')
                return False
            
            if not file_path.endswith('.docx'):
                print(f'✗ 不支持的文件格式：{file_path}')
                return False
            
            self.document = Document(file_path)
            self.file_path = file_path
            
            print(f'✓ 文件已加载：{file_path}')
            return True
            
        except Exception as e:
            print(f'✗ 加载文件失败：{e}')
            return False
    
    def get_text_content(self, max_length: int = 10000) -> str:
        """
        获取文本内容
        
        参数：
        - max_length: 最大长度
        
        返回：
        - str: 文本内容
        """
        if not self.document:
            return ''
        
        paragraphs = []
        for para in self.document.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        text = '\n'.join(paragraphs)
        
        # 截断
        if len(text) > max_length:
            text = text[:max_length] + '...'
        
        return text
    
    def get_document_info(self) -> Dict[str, Any]:
        """
        获取文档信息
        
        返回：
        - dict: 文档信息
        """
        if not self.document:
            return {}
        
        # 统计信息
        para_count = len(self.document.paragraphs)
        text_length = len(self.get_text_content())
        
        # 估算页数（每页约 30 段）
        estimated_pages = max(1, para_count // 30 + 1)
        
        return {
            'file_name': os.path.basename(self.file_path) if self.file_path else '',
            'file_path': self.file_path,
            'paragraph_count': para_count,
            'text_length': text_length,
            'estimated_pages': estimated_pages,
        }
    
    def get_structure(self) -> List[Dict[str, Any]]:
        """
        获取文档结构
        
        返回：
        - list: 结构列表
        """
        if not self.document:
            return []
        
        structure = []
        
        for i, para in enumerate(self.document.paragraphs):
            if para.text.strip():
                # 判断标题级别
                style_name = para.style.name if para.style else ''
                level = self._get_heading_level(style_name, para.text)
                
                structure.append({
                    'index': i,
                    'text': para.text[:100],  # 只显示前 100 字符
                    'style': style_name,
                    'level': level,
                    'is_heading': level > 0,
                })
        
        return structure
    
    def _get_heading_level(self, style_name: str, text: str) -> int:
        """
        判断标题级别
        
        参数：
        - style_name: 样式名
        - text: 文本
        
        返回：
        - int: 标题级别（0 表示不是标题）
        """
        # 检查样式名
        if 'Heading' in style_name or '标题' in style_name:
            if '1' in style_name or '一' in style_name:
                return 1
            elif '2' in style_name or '二' in style_name:
                return 2
            elif '3' in style_name or '三' in style_name:
                return 3
        
        # 检查文本模式
        import re
        
        # 第 1 章 格式
        if re.match(r'^第 [一二三四五六七八九十]+章', text):
            return 1
        
        # 1.1 格式
        if re.match(r'^\d+\.\d+', text):
            return 2
        
        # 1.1.1 格式
        if re.match(r'^\d+\.\d+\.\d+', text):
            return 3
        
        # 1 格式（一级标题）
        if re.match(r'^\d+[、.．]', text) and len(text) < 50:
            return 1
        
        return 0
    
    def get_preview_html(self, max_paragraphs: int = 50) -> str:
        """
        获取预览 HTML
        
        参数：
        - max_paragraphs: 最大段落数
        
        返回：
        - str: HTML 内容
        """
        if not self.document:
            return '<p>文档未加载</p>'
        
        html = ['<div style="font-family: Arial, sans-serif; padding: 20px;">']
        
        for i, para in enumerate(self.document.paragraphs[:max_paragraphs]):
            if not para.text.strip():
                continue
            
            style_name = para.style.name if para.style else ''
            level = self._get_heading_level(style_name, para.text)
            
            # 根据级别设置样式
            if level == 1:
                html.append(f'<h1 style="font-size: 24px; font-weight: bold; margin: 20px 0 10px 0;">{para.text}</h1>')
            elif level == 2:
                html.append(f'<h2 style="font-size: 20px; font-weight: bold; margin: 15px 0 8px 0;">{para.text}</h2>')
            elif level == 3:
                html.append(f'<h3 style="font-size: 16px; font-weight: bold; margin: 12px 0 6px 0;">{para.text}</h3>')
            else:
                html.append(f'<p style="margin: 5px 0; text-indent: 2em;">{para.text}</p>')
        
        if len(self.document.paragraphs) > max_paragraphs:
            html.append('<p style="color: #999; text-align: center; margin-top: 20px;">... 还有更多内容 ...</p>')
        
        html.append('</div>')
        
        return '\n'.join(html)
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        获取统计信息
        
        返回：
        - dict: 统计信息
        """
        if not self.document:
            return {}
        
        # 统计字数
        total_chars = 0
        chinese_chars = 0
        english_words = 0
        
        for para in self.document.paragraphs:
            text = para.text
            total_chars += len(text)
            
            # 中文字符
            chinese_chars += sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
            
            # 英文单词（简单统计）
            english_words += len(text.split())
        
        # 统计标题数
        headings = {'level_1': 0, 'level_2': 0, 'level_3': 0}
        for item in self.get_structure():
            if item['level'] == 1:
                headings['level_1'] += 1
            elif item['level'] == 2:
                headings['level_2'] += 1
            elif item['level'] == 3:
                headings['level_3'] += 1
        
        return {
            'total_chars': total_chars,
            'chinese_chars': chinese_chars,
            'english_words': english_words,
            'paragraph_count': len(self.document.paragraphs),
            'headings': headings,
        }
    
    def close(self):
        """关闭文档"""
        self.document = None
        self.file_path = None
        print('✓ 文档已关闭')


# 快捷函数
def preview_file(file_path: str) -> FilePreview:
    """快速预览文件"""
    preview = FilePreview()
    preview.load_file(file_path)
    return preview


if __name__ == '__main__':
    # 测试
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        preview = FilePreview()
        
        if preview.load_file(file_path):
            print('\n文档信息:')
            info = preview.get_document_info()
            for k, v in info.items():
                print(f'  {k}: {v}')
            
            print('\n文档结构:')
            structure = preview.get_structure()
            for item in structure[:10]:  # 只显示前 10 个
                if item['is_heading']:
                    print(f'  [{"H" * item["level"]}] {item["text"]}')
            
            print('\n统计信息:')
            stats = preview.get_statistics()
            for k, v in stats.items():
                print(f'  {k}: {v}')
    else:
        print('用法：python file_preview.py <docx 文件>')
