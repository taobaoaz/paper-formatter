"""
PDF 导出模块
PDF Exporter Module

功能：
- 将 Word 文档导出为 PDF
- 支持多种导出选项
- 批量导出支持

版本：v2.2.0
"""

import os
import sys
from typing import List, Optional, Dict, Any
from pathlib import Path


class PDFExporter:
    """PDF 导出器"""
    
    def __init__(self):
        """初始化 PDF 导出器"""
        self.reportlab_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查依赖"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            self.reportlab_available = True
            print("✅ reportlab 可用")
        except ImportError:
            print("⚠️ reportlab 未安装，PDF 导出功能受限")
            self.reportlab_available = False
    
    def export_docx_to_pdf(self, docx_path: str, pdf_path: str, 
                          options: Dict[str, Any] = None) -> bool:
        """
        导出 Word 文档为 PDF
        
        Args:
            docx_path: Word 文档路径
            pdf_path: PDF 输出路径
            options: 导出选项
            
        Returns:
            是否成功
        """
        if not os.path.exists(docx_path):
            print(f"❌ 文件不存在：{docx_path}")
            return False
        
        if not self.reportlab_available:
            print("❌ reportlab 未安装，无法导出 PDF")
            return False
        
        try:
            # 使用 reportlab 创建 PDF
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, letter, legal
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
            
            # 读取 Word 文档
            from docx import Document
            doc = Document(docx_path)
            
            # 创建 PDF 文档
            page_size = self._get_page_size(options.get('page_size', 'A4'))
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=page_size,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )
            
            # 存储 PDF 内容
            story = []
            styles = getSampleStyleSheet()
            
            # 自定义样式
            styles.add(ParagraphStyle(
                name='Chinese',
                parent=styles['Normal'],
                fontName='Helvetica',  # 需要中文字体支持
                fontSize=12,
                leading=14
            ))
            
            # 处理文档内容
            for para in doc.paragraphs:
                if para.text.strip():
                    # 根据样式选择 PDF 样式
                    style = self._get_paragraph_style(para.style.name, styles)
                    p = Paragraph(para.text, style)
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
            
            # 处理表格
            for table in doc.tables:
                from reportlab.platypus import Table, TableStyle
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 0.2*inch))
            
            # 构建 PDF
            pdf_doc.build(story)
            
            print(f"✅ PDF 导出成功：{pdf_path}")
            return True
            
        except Exception as e:
            print(f"❌ PDF 导出失败：{e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _get_page_size(self, size_name: str):
        """获取页面大小"""
        from reportlab.lib.pagesizes import A4, A3, B5, letter, legal
        
        sizes = {
            'A4': A4,
            'A3': A3,
            'B5': B5,
            'Letter': letter,
            'Legal': legal
        }
        
        return sizes.get(size_name, A4)
    
    def _get_paragraph_style(self, style_name: str, styles):
        """获取段落样式"""
        style_name_lower = style_name.lower()
        
        if 'heading 1' in style_name_lower or '标题 1' in style_name_lower:
            return styles['Heading1']
        elif 'heading 2' in style_name_lower or '标题 2' in style_name_lower:
            return styles['Heading2']
        elif 'heading 3' in style_name_lower or '标题 3' in style_name_lower:
            return styles['Heading3']
        elif 'title' in style_name_lower or '标题' in style_name_lower:
            return styles['Title']
        else:
            return styles['Normal']
    
    def batch_export(self, docx_files: List[str], output_dir: str,
                    options: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        批量导出 PDF
        
        Args:
            docx_files: Word 文件列表
            output_dir: 输出目录
            options: 导出选项
            
        Returns:
            导出结果统计
        """
        os.makedirs(output_dir, exist_ok=True)
        
        results = {
            'total': len(docx_files),
            'success': 0,
            'failed': 0,
            'files': []
        }
        
        for docx_path in docx_files:
            try:
                # 生成 PDF 文件名
                base_name = os.path.splitext(os.path.basename(docx_path))[0]
                pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
                
                # 导出
                success = self.export_docx_to_pdf(docx_path, pdf_path, options)
                
                if success:
                    results['success'] += 1
                    results['files'].append({
                        'docx': docx_path,
                        'pdf': pdf_path,
                        'status': 'success'
                    })
                else:
                    results['failed'] += 1
                    results['files'].append({
                        'docx': docx_path,
                        'pdf': pdf_path,
                        'status': 'failed',
                        'error': '导出失败'
                    })
                    
            except Exception as e:
                results['failed'] += 1
                results['files'].append({
                    'docx': docx_path,
                    'status': 'failed',
                    'error': str(e)
                })
        
        return results


class PDFExportOptions:
    """PDF 导出选项"""
    
    def __init__(self):
        """初始化导出选项"""
        self.page_size = 'A4'  # A4, A3, B5, Letter, Legal
        self.quality = 'standard'  # standard, high, minimum
        self.compress = True  # 是否压缩
        self.embed_fonts = True  # 嵌入字体
        self.margin_top = 2.54  # cm
        self.margin_bottom = 2.54  # cm
        self.margin_left = 3.17  # cm
        self.margin_right = 3.17  # cm
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            'page_size': self.page_size,
            'quality': self.quality,
            'compress': self.compress,
            'embed_fonts': self.embed_fonts,
            'margin_top': self.margin_top,
            'margin_bottom': self.margin_bottom,
            'margin_left': self.margin_left,
            'margin_right': self.margin_right
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'PDFExportOptions':
        """从字典创建"""
        options = cls()
        options.page_size = data.get('page_size', 'A4')
        options.quality = data.get('quality', 'standard')
        options.compress = data.get('compress', True)
        options.embed_fonts = data.get('embed_fonts', True)
        options.margin_top = data.get('margin_top', 2.54)
        options.margin_bottom = data.get('margin_bottom', 2.54)
        options.margin_left = data.get('margin_left', 3.17)
        options.margin_right = data.get('margin_right', 3.17)
        return options
